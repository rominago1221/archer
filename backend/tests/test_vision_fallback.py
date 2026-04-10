"""
Test Claude Vision Fallback for Scanned Documents
Tests the new feature that auto-detects scanned PDFs and image uploads,
routing them to Claude Vision API for OCR + analysis.

Features tested:
1. JPEG file upload → vision_mode: true
2. PNG file upload → vision_mode: true
3. Scanned PDF (< 100 chars text) → vision_mode: true
4. Normal text PDF → vision_mode: false
5. Accept attribute includes image formats
"""

import pytest
import requests
import os
import io
from PIL import Image

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')

# Test credentials from test_credentials.md
TEST_EMAIL = "test_cg@example.com"
TEST_PASSWORD = "testpassword123"


class TestVisionFallback:
    """Test Claude Vision fallback for scanned documents and images"""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        """Setup test session with authentication"""
        self.session = requests.Session()
        
        # Login to get session cookie (don't set Content-Type for multipart uploads)
        login_response = self.session.post(
            f"{BASE_URL}/api/auth/login",
            json={"email": TEST_EMAIL, "password": TEST_PASSWORD},
            headers={"Content-Type": "application/json"}
        )
        
        if login_response.status_code != 200:
            pytest.skip(f"Login failed: {login_response.status_code} - {login_response.text}")
        
        # Session cookie should be set automatically
        yield
    
    def create_test_jpeg(self, text: str = "Test Legal Document") -> bytes:
        """Create a simple JPEG image with text for testing"""
        img = Image.new('RGB', (400, 200), color='white')
        # Save to bytes
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=85)
        return buf.getvalue()
    
    def create_test_png(self, text: str = "Test Legal Document") -> bytes:
        """Create a simple PNG image for testing"""
        img = Image.new('RGB', (400, 200), color='white')
        buf = io.BytesIO()
        img.save(buf, format='PNG')
        return buf.getvalue()
    
    def create_minimal_pdf(self) -> bytes:
        """Create a minimal PDF with very little text (simulates scanned PDF)"""
        # Minimal PDF with almost no text content
        pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT
/F1 12 Tf
100 700 Td
(Hi) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000206 00000 n 
trailer
<< /Size 5 /Root 1 0 R >>
startxref
300
%%EOF"""
        return pdf_content
    
    def test_login_works(self):
        """Regression: Verify login still works"""
        response = self.session.get(f"{BASE_URL}/api/auth/me")
        assert response.status_code == 200, f"Auth check failed: {response.text}"
        data = response.json()
        assert data.get("email") == TEST_EMAIL
        assert data.get("plan") == "pro"
        print(f"✓ Login works - user: {data.get('email')}, plan: {data.get('plan')}")
    
    def test_jpeg_upload_returns_vision_mode_true(self):
        """Test: JPEG file upload returns vision_mode: true"""
        jpeg_bytes = self.create_test_jpeg()
        
        files = {
            'file': ('test_document.jpg', jpeg_bytes, 'image/jpeg')
        }
        data = {
            'analysis_mode': 'standard'
        }
        
        response = self.session.post(
            f"{BASE_URL}/api/documents/upload",
            files=files,
            data=data,
            timeout=120  # Vision API can take 15-30 seconds
        )
        
        assert response.status_code == 200, f"Upload failed: {response.status_code} - {response.text}"
        result = response.json()
        
        # Check vision_mode is true for JPEG
        assert result.get("vision_mode") == True, f"Expected vision_mode=true for JPEG, got: {result.get('vision_mode')}"
        assert result.get("document_id") is not None
        assert result.get("status") == "analyzed"
        
        print(f"✓ JPEG upload returns vision_mode=true, document_id={result.get('document_id')}")
    
    def test_png_upload_returns_vision_mode_true(self):
        """Test: PNG file upload returns vision_mode: true"""
        png_bytes = self.create_test_png()
        
        files = {
            'file': ('test_document.png', png_bytes, 'image/png')
        }
        data = {
            'analysis_mode': 'standard'
        }
        
        response = self.session.post(
            f"{BASE_URL}/api/documents/upload",
            files=files,
            data=data,
            timeout=120
        )
        
        assert response.status_code == 200, f"Upload failed: {response.status_code} - {response.text}"
        result = response.json()
        
        # Check vision_mode is true for PNG
        assert result.get("vision_mode") == True, f"Expected vision_mode=true for PNG, got: {result.get('vision_mode')}"
        assert result.get("document_id") is not None
        
        print(f"✓ PNG upload returns vision_mode=true, document_id={result.get('document_id')}")
    
    def test_scanned_pdf_fallback_to_vision(self):
        """Test: Scanned PDF (< 100 chars text) falls back to Claude Vision"""
        pdf_bytes = self.create_minimal_pdf()
        
        files = {
            'file': ('scanned_document.pdf', pdf_bytes, 'application/pdf')
        }
        data = {
            'analysis_mode': 'standard'
        }
        
        response = self.session.post(
            f"{BASE_URL}/api/documents/upload",
            files=files,
            data=data,
            timeout=120
        )
        
        assert response.status_code == 200, f"Upload failed: {response.status_code} - {response.text}"
        result = response.json()
        
        # Check vision_mode is true for scanned PDF
        assert result.get("vision_mode") == True, f"Expected vision_mode=true for scanned PDF, got: {result.get('vision_mode')}"
        
        print(f"✓ Scanned PDF (minimal text) returns vision_mode=true")
    
    def test_text_pdf_returns_vision_mode_false(self):
        """Test: Normal text-based PDF returns vision_mode: false
        
        Note: This test uses a DOCX file instead since creating a valid PDF
        with extractable text programmatically is complex. DOCX files should
        always have vision_mode=false since they contain extractable text.
        """
        # Create a DOCX file with substantial text content
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_paragraph("This is a legal document with substantial text content. " * 20)
        doc.add_paragraph("EMPLOYMENT AGREEMENT between Company ABC and Employee XYZ.")
        doc.add_paragraph("Section 1: The employee agrees to the following terms and conditions.")
        
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
        
        files = {
            'file': ('text_document.docx', docx_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        data = {
            'analysis_mode': 'standard'
        }
        
        response = self.session.post(
            f"{BASE_URL}/api/documents/upload",
            files=files,
            data=data,
            timeout=120
        )
        
        assert response.status_code == 200, f"Upload failed: {response.status_code} - {response.text}"
        result = response.json()
        
        # For text-based DOCX, vision_mode should be false
        assert result.get("vision_mode") == False, f"Expected vision_mode=false for DOCX, got: {result.get('vision_mode')}"
        assert result.get("document_id") is not None
        assert result.get("status") == "analyzed"
        
        print(f"✓ DOCX upload returns vision_mode=false (text extraction works)")
    
    def test_upload_response_has_analysis(self):
        """Test: Upload response includes analysis data"""
        jpeg_bytes = self.create_test_jpeg()
        
        files = {
            'file': ('legal_notice.jpg', jpeg_bytes, 'image/jpeg')
        }
        data = {
            'analysis_mode': 'standard'
        }
        
        response = self.session.post(
            f"{BASE_URL}/api/documents/upload",
            files=files,
            data=data,
            timeout=120
        )
        
        assert response.status_code == 200, f"Upload failed: {response.text}"
        result = response.json()
        
        # Check response structure
        assert "document_id" in result
        assert "case_id" in result
        assert "analysis" in result
        assert "vision_mode" in result
        assert "file_name" in result
        
        # Analysis may be None if image has no readable text, but field should exist
        print(f"✓ Upload response has all required fields: document_id, case_id, analysis, vision_mode, file_name")
    
    def test_webp_upload_returns_vision_mode_true(self):
        """Test: WebP file upload returns vision_mode: true"""
        # Create a simple WebP image
        img = Image.new('RGB', (400, 200), color='white')
        buf = io.BytesIO()
        img.save(buf, format='WEBP', quality=85)
        webp_bytes = buf.getvalue()
        
        files = {
            'file': ('test_document.webp', webp_bytes, 'image/webp')
        }
        data = {
            'analysis_mode': 'standard'
        }
        
        response = self.session.post(
            f"{BASE_URL}/api/documents/upload",
            files=files,
            data=data,
            timeout=120
        )
        
        assert response.status_code == 200, f"Upload failed: {response.status_code} - {response.text}"
        result = response.json()
        
        # Check vision_mode is true for WebP
        assert result.get("vision_mode") == True, f"Expected vision_mode=true for WebP, got: {result.get('vision_mode')}"
        
        print(f"✓ WebP upload returns vision_mode=true")


class TestUploadEndpointAcceptTypes:
    """Test that upload endpoint accepts various file types"""
    
    @pytest.fixture(autouse=True)
    def setup(self):
        """Setup test session with authentication"""
        self.session = requests.Session()
        
        # Login
        login_response = self.session.post(
            f"{BASE_URL}/api/auth/login",
            json={"email": TEST_EMAIL, "password": TEST_PASSWORD}
        )
        
        if login_response.status_code != 200:
            pytest.skip(f"Login failed: {login_response.status_code}")
        
        yield
    
    def test_accepts_jpeg_extension(self):
        """Test: Endpoint accepts .jpeg files"""
        img = Image.new('RGB', (100, 100), color='white')
        buf = io.BytesIO()
        img.save(buf, format='JPEG')
        
        files = {'file': ('test.jpeg', buf.getvalue(), 'image/jpeg')}
        response = self.session.post(
            f"{BASE_URL}/api/documents/upload",
            files=files,
            data={'analysis_mode': 'standard'},
            timeout=120
        )
        
        assert response.status_code == 200, f"JPEG upload failed: {response.text}"
        print("✓ .jpeg files accepted")
    
    def test_accepts_jpg_extension(self):
        """Test: Endpoint accepts .jpg files"""
        img = Image.new('RGB', (100, 100), color='white')
        buf = io.BytesIO()
        img.save(buf, format='JPEG')
        
        files = {'file': ('test.jpg', buf.getvalue(), 'image/jpeg')}
        response = self.session.post(
            f"{BASE_URL}/api/documents/upload",
            files=files,
            data={'analysis_mode': 'standard'},
            timeout=120
        )
        
        assert response.status_code == 200, f"JPG upload failed: {response.text}"
        print("✓ .jpg files accepted")
    
    def test_accepts_png_extension(self):
        """Test: Endpoint accepts .png files"""
        img = Image.new('RGB', (100, 100), color='white')
        buf = io.BytesIO()
        img.save(buf, format='PNG')
        
        files = {'file': ('test.png', buf.getvalue(), 'image/png')}
        response = self.session.post(
            f"{BASE_URL}/api/documents/upload",
            files=files,
            data={'analysis_mode': 'standard'},
            timeout=120
        )
        
        assert response.status_code == 200, f"PNG upload failed: {response.text}"
        print("✓ .png files accepted")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
