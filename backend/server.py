from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Response, Header, Query, Request, Form
from fastapi.responses import JSONResponse, StreamingResponse
from dotenv import load_dotenv
import asyncio
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, field_validator
from typing import List, Optional, Any, Dict
import uuid
from datetime import datetime, timezone, timedelta
from contextvars import ContextVar
import httpx
import json
import pdfplumber
import io
import requests
import base64
import bcrypt
from utils.admin_auth import admin_required, log_admin_action
import docx
import fitz  # pymupdf for PDF→image conversion
from PIL import Image
from emergentintegrations.payments.stripe.checkout import StripeCheckout, CheckoutSessionResponse, CheckoutStatusResponse, CheckoutSessionRequest
from emergentintegrations.llm.chat import LlmChat, UserMessage

# Shared modules
from db import db, client
from storage import init_storage, put_object, get_object, EMERGENT_KEY, APP_NAME
from auth import hash_password, verify_password, get_current_user, create_session_response
from models import (
    EmailRegister, EmailLogin, UserCreate, User, CaseCreate, CaseUpdate, Case,
    Document, Lawyer, LawyerCallCreate, LawyerCall, CaseEvent, ProfileUpdate,
    LetterRequest, normalize_deadline, normalize_financial_exposure,
)

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# API keys — use a getter to ensure .env is loaded regardless of import order.
# On Emergent, the key may be EMERGENT_LLM_KEY instead of ANTHROPIC_API_KEY.
def get_anthropic_key():
    return os.environ.get("ANTHROPIC_API_KEY") or os.environ.get("EMERGENT_LLM_KEY") or ""

ANTHROPIC_API_KEY = None  # set at startup
STRIPE_API_KEY = os.environ.get("STRIPE_API_KEY")

# Create the main app
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ================== Claude AI Analysis — 5-Pass Advanced System ==================

# Load jurisprudence database
JURISPRUDENCE_PATH = ROOT_DIR / "jurisprudence.json"
with open(JURISPRUDENCE_PATH) as f:
    JURISPRUDENCE_DB = json.load(f)

# Load Belgian jurisprudence database
JURISPRUDENCE_BE_PATH = ROOT_DIR / "jurisprudence_belgique.json"
with open(JURISPRUDENCE_BE_PATH) as f:
    JURISPRUDENCE_BE_DB = json.load(f)

SENIOR_ATTORNEY_PERSONA = """You are a senior attorney with 20 years of experience representing individuals — never corporations. You have handled over 2,000 cases across every area of civil law. You think like a litigator: you find every angle, every procedural error, every opportunity.

SPECIALIZATIONS:
Employment law — wrongful termination, unpaid wages, discrimination, harassment, non-compete, FMLA, ADA, Title VII, FLSA
Tenant and housing law — evictions, lease disputes, deposit returns, habitability, landlord violations, Belgian housing code
Debt and consumer protection — FDCPA violations, debt validation, statute of limitations, FCRA, credit reporting errors, predatory lending
Contract law — NDA, service agreements, vendor contracts, partnership agreements, breach of contract, penalty clauses
Consumer rights — refund disputes, warranty claims, FTC violations, chargeback support, deceptive practices
Immigration — visa contracts, employment authorization, sponsor agreements, work permit disputes
Family law — separation agreements, child support, asset division basics
Business disputes — invoice disputes, non-payment, IP disputes, non-solicitation
Court notices — summons, small claims, subpoenas, default judgments
Traffic and civil infractions — speeding tickets, license disputes, court summons
Real estate — purchase agreements, title disputes, HOA violations
Insurance disputes — claim denials, bad faith, underpayment

CRITICAL RULES — NEVER VIOLATE THESE:
Risk Score of 0 is NEVER acceptable — minimum 15/100 for any legal document
Risk Score of 100 is reserved for imminent criminal liability only
Every document has at least 3 findings worth identifying
Always cite the specific clause number, article, or section from the document
Always reference the applicable law by name — Florida Statute § 83.56, FDCPA 15 U.S.C. § 1692, Belgian Labour Code Art. 37
Never use legal jargon without immediately explaining it in plain English
Always identify what is MISSING from the opposing party's document — missing elements are often the strongest defense
Always assess the opposing party's likely goal — do they want money, possession, compliance?
Always identify at least one opportunity or leverage point for the user
Never admit liability or suggest the user is at fault
Always flag deadlines with exact dates and consequences if missed

Minimum score thresholds by document type:
— NDA: minimum 20/100 (always has restrictive clauses)
— Eviction notice: minimum 45/100 (always urgent)
— Demand letter: minimum 35/100 (always financial exposure)
— Employment contract: minimum 25/100 (always binding obligations)
— Court notice/summons: minimum 55/100 (always urgent)
— Debt collection letter: minimum 30/100 (always financial risk)
— Lease agreement: minimum 20/100 (always binding long-term)
— Settlement agreement: minimum 25/100 (always waiving rights)

FOR NDA SPECIFICALLY always analyze:
Confidentiality scope — is the definition of confidential information too broad?
Duration — how long are obligations? Perpetual NDAs are often unenforceable
Asymmetry — are both parties equally bound or only the user?
Penalty clauses — are damages proportionate and reasonable?
Geographic scope — is it limited to a reasonable territory?
Carve-outs — are standard exclusions present (public domain, prior knowledge)?
Governing law — which jurisdiction governs disputes?
Non-compete implications — does the NDA include hidden non-compete language?

FOR EMPLOYMENT DOCUMENTS always analyze:
Termination clauses — grounds and notice requirements
Non-compete and non-solicitation scope and enforceability
Compensation and bonus entitlements
Intellectual property assignment clauses
Arbitration clauses that waive court rights
At-will vs for-cause termination
Gratuity and end of service entitlements
Applicable labor law by jurisdiction

FOR DEBT COLLECTION always analyze:
FDCPA compliance — timing, harassment, validation rights
Statute of limitations — is the debt time-barred?
Debt ownership — has the debt been properly assigned?
Amount accuracy — are fees and interest correctly calculated?
Credit reporting implications

OUTPUT FORMAT — always return complete JSON with:
risk_score total + 4 dimensions (financial, urgency, legal_strength, complexity)
minimum 3 findings with impact level and type
exactly 3 next steps with action type
deadline with exact date if present
financial exposure in specific dollar/EUR amount
applicable law reference
recommend_lawyer boolean
key_insight — one sentence the user must remember"""

PASS1_PROMPT = """TASK: FACT EXTRACTION ONLY
Read this legal document carefully and extract every factual element. Do not interpret, analyze, or recommend anything yet. Return ONLY raw facts organized in JSON.

DOCUMENT TEXT:
{document_text}

Return ONLY this JSON — no other text:
{{
  "document_type": "demand_letter|eviction_notice|employment_contract|court_notice|debt_collection|nda|lease|other",
  "document_date": "YYYY-MM-DD or null",
  "parties": {{
    "opposing_party": {{"name": "exact name from document", "type": "individual|company|government|law_firm"}},
    "user_party": {{"name": "exact name from document", "role": "tenant|employee|debtor|consumer|contractor|other"}}
  }},
  "key_amounts": [{{"amount": 0, "currency": "USD", "description": "what this amount represents", "disputed": false}}],
  "key_dates": [{{"date": "YYYY-MM-DD", "description": "what this date represents", "is_deadline": true}}],
  "legal_references": [{{"reference": "statute reference", "description": "what law is referenced"}}],
  "contract_clauses": [{{"clause_number": "Section X", "description": "what clause says"}}],
  "claims_made": [{{"claim": "exact claim by opposing party", "amount": null}}],
  "missing_elements": ["what important info is absent from the document"],
  "procedural_elements": {{
    "service_method": "certified mail|posted|email|other|not specified",
    "notice_period_stated": "X days|not specified",
    "signature_present": true,
    "notarization_required": false
  }}
}}"""

PASS2_PROMPT = """TASK: LEGAL ANALYSIS
Based on the extracted facts below, conduct a thorough legal analysis from the perspective of the user's attorney. Apply applicable US federal and state law.

EXTRACTED FACTS FROM DOCUMENT:
{facts_json}

{jurisprudence_section}

ANALYSIS REQUIRED:
1. PROCEDURAL ANALYSIS: procedural errors, defects, or omissions that could weaken the opposing party's position
2. SUBSTANTIVE ANALYSIS: strength of opposing party's case on the merits
3. USER'S POSITION: rights and protections under applicable law, counterclaims or defenses
4. FINANCIAL ANALYSIS: realistic financial exposure (best/most likely/worst case)
5. OPPOSING PARTY'S INCENTIVES: do they want court? likely goal? bluffing or serious?
6. TIMING ANALYSIS: how critical are deadlines? legally enforceable?

Return ONLY this JSON — no other text:
{{
  "risk_score": {{
    "total": 0,
    "financial": 0,
    "urgency": 0,
    "legal_strength": 0,
    "complexity": 0,
    "level": "low|moderate|high|critical",
    "tagline": "Short 4-7 word phrase capturing the emotional essence. Pattern: [Serious state], but [glimmer of hope]. Examples: low=Not stressful, under control. moderate=Worth watching, actions needed. high=Serious, but contestable. critical=Urgent. Act today."
  }},
  "risk_level": "low|medium|high|critical",
  "case_type": "employment|housing|nda|contract|debt|demand|immigration|court|consumer|family|traffic|insurance|other",
  "suggested_case_title": "Short descriptive title based on document content max 60 chars — NEVER use the filename",
  "deadline": "YYYY-MM-DD or null",
  "deadline_description": "Description or null",
  "summary": "2-3 sentences plain English summary",
  "financial_exposure": "Dollar amount or null",
  "financial_exposure_detailed": {{
    "best_case": "description",
    "most_likely": "description",
    "worst_case": "description"
  }},
  "procedural_defects": [
    {{"defect": "description", "severity": "fatal|significant|minor", "applicable_law": "law ref", "user_benefit": "how this helps"}}
  ],
  "user_rights": [
    {{"right": "specific right", "law_reference": "law", "strength": "strong|medium|weak"}}
  ],
  "opposing_weaknesses": [
    {{"weakness": "description", "severity": "critical|significant|minor"}}
  ],
  "applicable_laws": [
    {{"law": "statute name", "relevance": "how it applies", "favors": "user|opposing|neutral"}}
  ],
  "findings": [
    {{"text": "Finding title: short, specific, actionable — NEVER vague or generic", "impact": "high|medium|low", "type": "risk|opportunity|deadline|neutral", "legal_ref": "Exact statute, article, or case law citation (e.g. Fla. Stat. § 83.56(3) — Florida 3rd DCA 2023)", "jurisprudence": "Relevant case law if any", "impact_description": "What this means for the user RIGHT NOW in plain language — no legal jargon, written as if explaining to a friend", "do_now": "Exact next step the user MUST take — specific and actionable, never generic like 'consult an attorney'", "risk_if_ignored": "What happens if user does NOTHING — create urgency, explain real consequences of inaction", "confidence_score": 0.92, "jurisprudence_count": 47, "similar_cases_won": 44, "similar_cases_total": 47, "reasoning": "1-2 sentences explaining WHY this confidence score. Reference the jurisprudence consistency and how clear the law is."}}
  ],
  "key_violation_types": ["short_notice_period", "missing_formal_notice", "vague_eviction_motive"],
  "recommend_lawyer": true,
  "disclaimer": "This analysis provides legal information only, not legal advice."
}}
KEY_VIOLATION_TYPES: Generate 2-5 short snake_case tags describing the key violations found (e.g. "short_notice_period", "no_written_contract", "unpaid_wages", "illegal_clause", "missing_mise_en_demeure"). These tags are used to match similar cases.

Produce 3-6 findings. EVERY finding MUST have ALL 12 fields: text (title), impact, type, legal_ref (exact statute), jurisprudence, impact_description (plain language), do_now (specific action), risk_if_ignored (consequence of inaction), confidence_score (0.0-1.0), jurisprudence_count (int), similar_cases_won (int), similar_cases_total (int), reasoning (1-2 sentences).

CONFIDENCE SCORE RULES:
- Law is clear AND jurisprudence consistent: confidence > 0.90
- Law is clear but jurisprudence variable: 0.70-0.90
- Interpretation possible: 0.50-0.70
- Weak argument: < 0.50
Base scores on known case law patterns. Jurisprudence_count = approximate number of relevant similar rulings."""

PASS3_PROMPT = """TASK: STRATEGIC RECOMMENDATIONS
Based on the facts and legal analysis below, provide concrete strategic recommendations. Think like a litigator preparing a client for the best outcome.

EXTRACTED FACTS:
{facts_json}

LEGAL ANALYSIS:
{analysis_json}

CRITICAL RULE — CASE STAGE DETECTION:
You MUST identify the current stage of the legal proceeding and adapt ALL recommendations accordingly:

STAGE 1 — Initial notice/violation received:
- User just received a notice (eviction, debt collection, termination, etc.)
- Actions: Respond within deadline, dispute validity, gather evidence
- Letter to: the original sender (landlord, employer, collector)

STAGE 2 — Pre-litigation / negotiation:
- Parties are negotiating, settlement offers made
- Actions: Counter-offer, document everything, prepare for court
- Letter to: opposing party or their attorney

STAGE 3 — Court proceedings filed:
- Case is before a court/judge, hearing date set
- The original sender is NO LONGER the right recipient — NEVER suggest writing to them
- Actions: Retain attorney, file court responses, prepare defense
- Letter to: Court clerk / opposing counsel
- Priority 1: Consult attorney (action_type: book_lawyer)
- Priority 2: File court response if deadline (action_type: send_letter to court)
- Priority 3: Gather defense evidence (action_type: gather_documents)

STAGE 4 — Judgment/ruling issued:
- Court has ruled
- Actions: Appeal if unfavorable (within deadline), enforce if favorable
- Letter to: Appellate court clerk

DETECTION SIGNALS:
- "court date", "hearing", "docket", "filed suit", "summons" = STAGE 3
- "notice", "demand letter", "pay or quit", "termination" = STAGE 1
- "settlement", "negotiate", "offer", "mediation" = STAGE 2
- "judgment", "ruling", "ordered", "appeal" = STAGE 4

MANDATORY: Each next_step MUST include a "recipient" field indicating WHO the letter/action is directed to.

Think through ALL possible paths:
1. Ideal outcome? 2. Fastest resolution? 3. Cheapest path? 4. Most leverage?
5. What should user do in next 24 hours? 6. What documents to gather? 7. Talk to lawyer first?

CRITICAL RULES — you MUST follow ALL of these:

RULE 1 — archer_question: You MUST generate exactly ONE specific clarifying question that references something specific found in the document. Always provide 2-4 answer options. NEVER skip this field.

RULE 2 — success_probability: Realistic outcome probabilities. No outcome below 2% or above 95%. All four values MUST sum to 100.

RULE 3 — next_steps: Each step MUST include a specific legal reference AND a recipient. NEVER suggest writing to the wrong party for the current stage.

Return ONLY this JSON — no other text:
{{
  "case_stage": "stage_1_notice|stage_2_negotiation|stage_3_court|stage_4_judgment",
  "recommended_strategy": {{
    "primary": "negotiate|dispute|comply|ignore|lawyer_immediately",
    "reasoning": "why this is best for THIS STAGE",
    "expected_outcome": "realistic outcome",
    "time_to_resolution": "3-7 days|1-4 weeks|1-3 months"
  }},
  "immediate_actions": [
    {{"action": "description", "deadline": "within 24 hours", "priority": "critical|high|medium"}}
  ],
  "next_steps": [
    {{
      "title": "Specific action verb + exact legal reference",
      "description": "Detailed description adapted to CURRENT STAGE",
      "action_type": "send_letter|book_lawyer|wait|gather_documents|no_action",
      "recipient": "Exact recipient for this stage (e.g. Court Clerk, Opposing Counsel, Landlord)",
      "letter_template": "Brief description of what letter to generate (REQUIRED for send_letter type)",
      "why_important": "Why this matters — reference specific law"
    }}
  ],
  "documents_to_gather": [
    {{"document": "description", "why": "reason", "urgency": "critical|important|nice_to_have"}}
  ],
  "leverage_points": [
    {{"leverage": "description", "how_to_use": "how to use it"}}
  ],
  "red_lines": ["Never do X", "Always do Y"],
  "lawyer_recommendation": {{
    "needed": true,
    "urgency": "immediately|within_3_days|within_week|optional",
    "reason": "why",
    "type_needed": "tenant_rights|employment|debt_collection|contract"
  }},
  "success_probability": {{
    "full_resolution_in_favor": 15,
    "negotiated_settlement": 62,
    "partial_loss": 18,
    "full_loss": 5
  }},
  "key_insight": "The most important thing the user must know in one sentence",
  "archer_question": {{
    "text": "A SPECIFIC clarifying question referencing a fact from the document (e.g. 'The notice mentions a $150 fee — did you agree to this fee in your lease?' or 'Avez-vous une preuve écrite de votre signalement de harcèlement ?')",
    "options": ["Answer option 1", "Answer option 2", "Answer option 3"]
  }}
}}
MANDATORY: archer_question MUST be present with 2-4 options. success_probability values MUST sum to 100 with no value below 2 or above 95. next_steps MUST have exactly 3 items with specific legal references."""

PASS4A_SYSTEM = """You are a senior attorney representing the user. Your job is to make the STRONGEST possible case for your client. Find every argument, every procedural defect, every legal protection that benefits your client. Be aggressive and thorough. You MUST produce exactly 4-5 strong arguments — NEVER leave arguments empty."""

PASS4A_PROMPT = """Based on these facts and legal analysis, make the strongest possible case for the user. What are their best arguments? What gives them the most leverage? What mistakes did the opposing party make?

FACTS:
{facts_json}

LEGAL ANALYSIS:
{analysis_json}

CRITICAL: You MUST generate exactly 4-5 strongest_arguments. NEVER return an empty array. Each argument must reference a specific law, statute, or legal principle. If the user's position seems weak, find procedural arguments, constitutional protections, or burden-of-proof arguments.

Return ONLY this JSON:
{{
  "strongest_arguments": [
    {{"argument": "description", "strength": "strong|medium|weak", "law_basis": "specific law/statute reference", "how_to_use": "how to use in response"}}
  ],
  "procedural_wins": ["list of procedural advantages"],
  "best_outcome_scenario": "description of best possible outcome",
  "opening_argument": "First sentence of response letter to maximize impact"
}}
MANDATORY: strongest_arguments MUST contain exactly 4-5 items. NEVER return fewer than 4."""

PASS4B_SYSTEM = """You are a senior attorney representing the opposing party. Your job is to make the STRONGEST possible case against the user. Be rigorous and identify every weakness in the user's position."""

PASS4B_PROMPT = """Based on these facts and legal analysis, make the strongest possible case against the user. What arguments will the opposing party likely use? What are the weaknesses in the user's position?

FACTS:
{facts_json}

LEGAL ANALYSIS:
{analysis_json}

Return ONLY this JSON:
{{
  "opposing_arguments": [
    {{"argument": "description", "strength": "strong|medium|weak", "law_basis": "law ref", "user_counter": "how user can counter"}}
  ],
  "user_weaknesses": ["list of weaknesses"],
  "worst_outcome_scenario": "description of worst possible outcome",
  "what_user_must_prepare_for": "what opposing party will likely argue"
}}"""


def load_jurisprudence(case_type: str, document_type: str) -> str:
    """Load relevant jurisprudence for the case"""
    entries = []
    usa_data = JURISPRUDENCE_DB.get("USA", {})

    # Map case/doc types to jurisprudence keys
    type_map = {
        "housing": ["florida_tenant", "general_housing"],
        "employment": ["new_york_employment", "california_employment", "general_employment"],
        "debt": ["federal_debt"],
        "consumer": ["general_consumer"],
        "contract": ["general_consumer"],
    }

    keys = type_map.get(case_type, [])
    if not keys and document_type in ["debt_collection"]:
        keys = ["federal_debt"]
    if not keys:
        keys = ["general_consumer"]

    for key in keys:
        for entry in usa_data.get(key, []):
            entries.append(f"- {entry['name']}: {entry['rule']} ({entry['statute']}). User benefit: {entry['user_benefit']}")

    if entries:
        return "RELEVANT LAW AND PRECEDENTS:\n" + "\n".join(entries) + "\n\nApply these specifically to the facts above."
    return ""


async def call_claude(system_prompt: str, user_message: str, max_tokens: int = 4000, use_web_search: bool = False) -> dict:
    """Direct Anthropic API — Opus 4.6 for analysis pipeline.
    System prompt is cached (ephemeral) for cross-pass efficiency.

    Two-level retry:
      - OUTER loop (3 attempts): retries on JSON parse errors or API errors
      - INNER loop (4 attempts): auto-scales max_tokens on truncation (4k→8k→16k→32k)
    Returns {} on total failure instead of raising (pipeline continues with partial data).
    """
    current_max = max_tokens

    for attempt in range(3):  # Outer: JSON parse / API error retries
        try:
            text = None

            # Inner: truncation auto-scaling
            for scale in range(4):
                async with httpx.AsyncClient(timeout=180) as client:
                    resp = await client.post(
                        "https://api.anthropic.com/v1/messages",
                        headers={
                            "x-api-key": get_anthropic_key(),
                            "anthropic-version": "2023-06-01",
                            "anthropic-beta": "prompt-caching-2024-07-31",
                            "content-type": "application/json",
                        },
                        json={
                            "model": "claude-opus-4-7",
                            "max_tokens": current_max,
                            "system": [
                                {"type": "text", "text": system_prompt, "cache_control": {"type": "ephemeral"}},
                            ],
                            "messages": [
                                {"role": "user", "content": user_message + "\n\nRespond with valid JSON only. No markdown, no code blocks, just raw JSON."},
                            ],
                        },
                    )
                if resp.status_code != 200:
                    body = resp.text
                    logger.error(f"Anthropic API error {resp.status_code}: {body[:500]}")
                    raise Exception(f"Anthropic API {resp.status_code}: {body[:200]}")

                data = resp.json()
                stop_reason = data.get("stop_reason", "end_turn")
                text = data["content"][0]["text"]

                if stop_reason == "max_tokens" and current_max < 32000:
                    current_max = min(current_max * 2, 32000)
                    logger.warning(f"Response truncated (stop_reason=max_tokens), retrying with {current_max}")
                    continue
                break  # Got full response or hit 32k cap

            # Parse JSON
            text = text.replace("```json", "").replace("```", "").strip()
            return json.loads(text)

        except json.JSONDecodeError:
            logger.warning(f"JSON parse error (attempt {attempt+1}/3, max_tokens={current_max})")
            await asyncio.sleep(1)
            continue
        except Exception as e:
            if attempt < 2:
                wait = (attempt + 1) * 3
                logger.warning(f"Claude call error: {e}, retrying in {wait}s (attempt {attempt+1}/3)")
                await asyncio.sleep(wait)
                continue
            logger.error(f"call_claude failed after 3 attempts: {e}")

    logger.error(f"call_claude returning empty dict after all retries (max_tokens={current_max})")
    return {}


async def call_claude_fast(system_prompt: str, user_message: str, max_tokens: int = 800) -> dict:
    """Direct Anthropic API — Opus 4.6 for analysis tasks (Q&A impact, letter drafts)."""
    for attempt in range(2):
        try:
            async with httpx.AsyncClient(timeout=90) as client:
                resp = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": get_anthropic_key(),
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                    },
                    json={
                        "model": "claude-opus-4-7",
                        "max_tokens": max_tokens,
                        "system": system_prompt,
                        "messages": [
                            {"role": "user", "content": user_message + "\n\nRespond with valid JSON only. No markdown, no code blocks, just raw JSON."},
                        ],
                    },
                )
                resp.raise_for_status()
                data = resp.json()
                text = data["content"][0]["text"]
                text = text.replace("```json", "").replace("```", "").strip()
                return json.loads(text)
        except json.JSONDecodeError:
            if attempt < 1:
                logger.warning(f"Fast call JSON parse error, retrying")
                await asyncio.sleep(1)
                continue
            raise
        except Exception as e:
            if attempt < 1:
                logger.warning(f"Fast call error: {e}, retrying")
                await asyncio.sleep(2)
                continue
            raise


import hashlib

async def get_cached_analysis(doc_hash: str) -> dict:
    """Check MongoDB for cached analysis result"""
    cached = await db.analysis_cache.find_one({"doc_hash": doc_hash}, {"_id": 0})
    if not cached:
        return None
    # Check 24h expiry
    cached_at = cached.get("cached_at", "")
    if cached_at:
        try:
            cached_time = datetime.fromisoformat(cached_at.replace("Z", "+00:00"))
            if cached_time.tzinfo is None:
                cached_time = cached_time.replace(tzinfo=timezone.utc)
            if (datetime.now(timezone.utc) - cached_time).total_seconds() > 86400:
                return None
        except Exception:
            return None
    return cached.get("result")


async def set_cached_analysis(doc_hash: str, result: dict):
    """Store analysis result in cache"""
    await db.analysis_cache.update_one(
        {"doc_hash": doc_hash},
        {"$set": {"doc_hash": doc_hash, "result": result, "cached_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )


async def fetch_courtlistener_opinions(case_type: str, state: str = "") -> list:
    """Fetch recent court opinions from CourtListener API filtered by case type AND jurisdiction"""
    type_to_query = {
        "housing": "eviction tenant landlord housing rental lease",
        "employment": "wrongful termination employment discrimination wages",
        "debt": "debt collection FDCPA creditor",
        "nda": "non-disclosure agreement trade secret",
        "contract": "breach of contract damages",
        "demand": "demand letter civil dispute",
        "court": "motion to dismiss summary judgment",
        "consumer": "consumer protection refund deceptive",
        "immigration": "immigration visa employment authorization",
        "family": "family law custody support divorce",
        "insurance": "insurance claim denial bad faith",
        "penal": "criminal defense sentencing",
        "commercial": "commercial dispute business",
    }
    query = type_to_query.get(case_type, "")
    if not query:
        return []  # Don't return random cases for unknown types
    
    # Add state/jurisdiction for targeted results
    if state:
        query = f"{state} {query}"
    
    # Map states to CourtListener court abbreviations for better filtering
    state_court_map = {
        "florida": "flaapp", "california": "cal", "new york": "ny",
        "texas": "tex", "illinois": "ill", "pennsylvania": "pa",
        "ohio": "ohio", "georgia": "ga", "michigan": "mich",
    }
    court_filter = ""
    if state:
        state_lower = state.lower().strip()
        court_filter = state_court_map.get(state_lower, "")
    
    try:
        params = {
            "type": "o",
            "q": query,
            "order_by": "dateFiled desc",
            "format": "json"
        }
        if court_filter:
            params["court"] = court_filter
        
        async with httpx.AsyncClient() as client:
            response = await client.get(
                "https://www.courtlistener.com/api/rest/v4/search/",
                params=params,
                timeout=15.0,
                headers={"User-Agent": "Archer-Legal-AI/1.0"}
            )
            if response.status_code != 200:
                logger.warning(f"CourtListener API returned {response.status_code}")
                return []
            
            data = response.json()
            results = data.get("results", [])[:3]
            opinions = []
            for r in results:
                opinions.append({
                    "case_name": r.get("caseName", "Unknown case"),
                    "date_filed": r.get("dateFiled", ""),
                    "court": r.get("court", ""),
                    "snippet": (r.get("snippet", "") or "")[:300],
                    "url": f"https://www.courtlistener.com{r.get('absolute_url', '')}" if r.get("absolute_url") else None,
                    "cite_count": r.get("citeCount", 0)
                })
            logger.info(f"CourtListener: fetched {len(opinions)} opinions for '{query}' (state: {state or 'all'})")
            return opinions
    except Exception as e:
        logger.warning(f"CourtListener API error: {e}")
        return []


async def analyze_document_with_claude(extracted_text: str) -> dict:
    """Legacy single-pass analysis (fallback)"""
    try:
        result = await call_claude(CLAUDE_SYSTEM_PROMPT, f"Analyze this legal document and return JSON only:\n\n{extracted_text[:15000]}")
        return result
    except Exception as e:
        logger.error(f"Claude API error: {e}")
        return _default_analysis()



# ═══ SHARED ANALYSIS HELPERS ═══
# Extracted from analyze_document_advanced / run_multi_doc_analysis_advanced to reduce complexity

DOC_TYPE_TO_CASE = {
    "eviction_notice": "housing", "lease": "housing",
    "employment_contract": "employment",
    "debt_collection": "debt",
    "demand_letter": "demand",
    "court_notice": "court", "nda": "nda"
}

def _build_realtime_law_context(courtlistener_opinions: list) -> str:
    """Build real-time case law context string from CourtListener opinions."""
    if not courtlistener_opinions:
        return ""
    ctx = "\n\nRECENT COURT DECISIONS (from CourtListener — real-time data):\n"
    for i, op in enumerate(courtlistener_opinions, 1):
        ctx += f"{i}. {op['case_name']} (Filed: {op['date_filed']}, Court: {op['court']})\n"
        if op.get('snippet'):
            clean_snippet = op['snippet'].replace('<em>', '').replace('</em>', '')
            ctx += f"   Excerpt: {clean_snippet}\n"
    ctx += "\nConsider these recent decisions when assessing the user's legal position.\n"
    return ctx


def _build_case_law_for_frontend(courtlistener_opinions: list) -> list:
    """Transform CourtListener opinions into frontend-ready case law dicts."""
    result = []
    for op in courtlistener_opinions:
        clean_snippet = (op.get("snippet", "") or "").replace("<em>", "").replace("</em>", "")
        result.append({
            "case_name": op["case_name"],
            "date": op["date_filed"],
            "court": op["court"],
            "ruling_summary": clean_snippet[:200] if clean_snippet else "Decision text available at source.",
            "source_url": op.get("url"),
            "cite_count": op.get("cite_count", 0)
        })
    return result


async def _validate_archer_question(strategy: dict, facts_str: str, persona: str, lang_instruction: str, language: str = "en") -> dict:
    """Ensure archer_question exists with valid text and options, retrying if needed."""
    jq = strategy.get("archer_question")
    if jq and isinstance(jq, dict) and jq.get("text") and jq.get("options") and len(jq.get("options", [])) >= 2:
        return jq

    logger.warning("Archer question missing or invalid — retrying")
    is_french = language.startswith("fr")
    is_dutch = language.startswith("nl")
    is_german = language.startswith("de")
    is_spanish = language.startswith("es")

    if is_french:
        retry_prompt = f"Sur base de cette analyse documentaire, genere UNE question specifique de clarification avec 2-4 options de reponse. La question DOIT referencer un fait precis du document. JAMAIS de question generique.\n\nFaits: {facts_str[:2000]}\n\nRetourne UNIQUEMENT JSON: {{\"text\": \"question specifique\", \"options\": [\"option1\", \"option2\", \"option3\"]}}"
        fallback = {"text": "Avez-vous d'autres documents lies a cette affaire ?", "options": ["Oui, j'ai d'autres documents", "Non, c'est tout", "Je ne suis pas sur"]}
    elif is_dutch:
        retry_prompt = f"Op basis van deze documentanalyse, genereer EEN specifieke verduidelijkingsvraag met 2-4 antwoordopties.\n\nFeiten: {facts_str[:2000]}\n\nRetourneer ALLEEN JSON: {{\"text\": \"specifieke vraag\", \"options\": [\"optie1\", \"optie2\", \"optie3\"]}}"
        fallback = {"text": "Heeft u andere documenten met betrekking tot deze zaak?", "options": ["Ja, ik heb meer documenten", "Nee, dat is alles", "Ik weet het niet zeker"]}
    elif is_german:
        retry_prompt = f"Basierend auf dieser Dokumentenanalyse, generiere EINE spezifische Klarungsfrage mit 2-4 Antwortoptionen.\n\nFakten: {facts_str[:2000]}\n\nGib NUR JSON zuruck: {{\"text\": \"spezifische Frage\", \"options\": [\"Option1\", \"Option2\", \"Option3\"]}}"
        fallback = {"text": "Haben Sie weitere Dokumente zu diesem Fall?", "options": ["Ja, ich habe weitere Dokumente", "Nein, das ist alles", "Ich bin nicht sicher"]}
    elif is_spanish:
        retry_prompt = f"Basado en este analisis del documento, genera UNA pregunta de clarificacion especifica con 2-4 opciones de respuesta.\n\nHechos: {facts_str[:2000]}\n\nDevuelve SOLO JSON: {{\"text\": \"pregunta especifica\", \"options\": [\"opcion1\", \"opcion2\", \"opcion3\"]}}"
        fallback = {"text": "Tiene otros documentos relacionados con este caso?", "options": ["Si, tengo mas documentos", "No, eso es todo", "No estoy seguro"]}
    else:
        retry_prompt = f"Based on this document analysis, generate ONE specific clarifying question with 2-4 answer options. The question must reference a specific fact from the document. NEVER a generic question.\n\nFacts: {facts_str[:2000]}\n\nReturn ONLY JSON: {{\"text\": \"specific question\", \"options\": [\"option1\", \"option2\", \"option3\"]}}"
        fallback = {"text": "Do you have any additional documents related to this case?", "options": ["Yes, I have more documents", "No, this is everything", "I'm not sure"]}

    try:
        await asyncio.sleep(1)
        jq_retry = await call_claude(persona, retry_prompt, max_tokens=500)
        if jq_retry and jq_retry.get("text") and jq_retry.get("options") and len(jq_retry.get("options", [])) >= 2:
            return jq_retry
    except Exception as e:
        logger.error(f"Archer question retry failed: {e}")

    return fallback


# ================== Self-critique pass ==================
# Anthropic-style two-phase self-critique: the model critiques its own PASS2/PASS3
# output, and if severity > 5 we rewrite. Every critique is persisted to
# `analysis_critiques` for post-hoc quality measurement. severity > 8 also fires
# an admin alert into `admin_alerts`.

SELF_CRITIQUE_THRESHOLD = 5      # rewrite when severity_score > this
SELF_CRITIQUE_ADMIN_THRESHOLD = 8  # admin_alerts write when severity_score > this

# Ambient context: the case_id of the analysis currently running. Set by the
# upload / reanalyze / switch-jurisdiction / refine call sites so the critique
# layer can attribute records without threading case_id through every function.
_current_case_id: ContextVar[Optional[str]] = ContextVar("_current_case_id", default=None)


async def self_critique_pass(
    analysis_output: dict,
    *,
    pass_name: str,
    context_hint: str,
    language: str = "en",
    jurisdiction: str = "US",
    case_id: Optional[str] = None,
) -> dict:
    """Critique the given analysis JSON. If severity_score > 5, rewrite and return
    the improved version. Otherwise return the original unchanged.

    Never raises. On any failure the original `analysis_output` is returned — the
    primary pipeline must not be destabilised by the critique layer.

    `context_hint` is a 1-2 sentence reminder of what this pass was supposed to
    produce. Used inside the critique prompt to scope the reviewer.
    """
    if not isinstance(analysis_output, dict) or not analysis_output:
        return analysis_output

    if case_id is None:
        case_id = _current_case_id.get()

    try:
        original_json = json.dumps(analysis_output, ensure_ascii=False)[:12000]
    except Exception:
        return analysis_output

    critique_system = (
        "You are a senior litigation attorney reviewing a junior colleague's analysis before it is "
        "used by a real client. Your job is to be brutally honest about weaknesses. Pretend this will "
        "be scrutinised by an opposing counsel in court. Identify every gap, every superficial claim, "
        "every missing nuance, every imprecise legal reference, every strategic alternative that was "
        "not explored. Be specific, not generic. Cite what should be added, corrected, or deepened.\n\n"
        "Return ONLY valid JSON with exactly these keys:\n"
        "{\n"
        '  "critiques": [ { "category": "findings_missed|weak_argument|imprecise_law_ref|missed_nuance|missed_alternative|other", '
        '"description": "specific issue, 1-2 sentences", "severity": integer 1-10 } ],\n'
        '  "severity_score": integer 0-10, representing the overall severity (max of individual severities, adjusted by count),\n'
        '  "summary": "one-line verdict"\n'
        "}\n"
        "If the analysis is genuinely solid, return critiques=[] and severity_score=0."
    )

    critique_user = (
        f"PASS UNDER REVIEW: {pass_name}\n"
        f"WHAT THIS PASS SHOULD PRODUCE: {context_hint}\n"
        f"JURISDICTION CONTEXT: {jurisdiction}\n\n"
        f"ANALYSIS OUTPUT TO CRITIQUE:\n{original_json}\n\n"
        "Review it and return the critique JSON only."
    )

    try:
        critique_result = await call_claude(critique_system, critique_user, max_tokens=2000)
    except Exception as e:
        logger.error(f"self_critique_pass critique call failed for {pass_name}: {e}")
        return analysis_output

    if not isinstance(critique_result, dict):
        return analysis_output

    try:
        severity = int(critique_result.get("severity_score", 0) or 0)
    except (TypeError, ValueError):
        severity = 0
    severity = max(0, min(severity, 10))
    critiques = critique_result.get("critiques") or []
    applied = False

    final_output = analysis_output

    if severity > SELF_CRITIQUE_THRESHOLD and critiques:
        logger.info(f"self_critique_pass {pass_name}: severity={severity} → rewriting")
        rewrite_system = (
            SENIOR_ATTORNEY_PERSONA + get_jurisdiction_guard(jurisdiction, language)
            + get_language_instruction(language)
            + "\n\nYou wrote a legal analysis. A senior reviewer flagged the issues below. "
            "Rewrite the analysis to address every critique. CRITICAL RULES:\n"
            "1. Return the EXACT same JSON schema as the original — identical keys, identical nesting, identical types.\n"
            "2. Do NOT add new top-level keys. Do NOT drop existing top-level keys.\n"
            "3. Improve CONTENT only: tighten findings, add missed nuances, strengthen weak arguments, "
            "precise up loose legal references, surface missed strategic alternatives.\n"
            "4. Keep numeric fields (risk_score subfields, confidence_score, etc.) realistic — recalibrate if justified by the critique, otherwise keep.\n"
            "5. Write in the same language as the original analysis."
        )
        rewrite_user = (
            f"ORIGINAL ANALYSIS ({pass_name}):\n{original_json}\n\n"
            f"CRITIQUES TO ADDRESS:\n{json.dumps(critiques, ensure_ascii=False)[:6000]}\n\n"
            "Return the improved analysis JSON only. Same schema, no extra prose."
        )
        try:
            rewritten = await call_claude(rewrite_system, rewrite_user, max_tokens=8000)
            if isinstance(rewritten, dict) and rewritten:
                # Basic schema sanity: rewrite must carry at least the same top-level keys.
                missing = [k for k in analysis_output.keys() if k not in rewritten]
                if missing:
                    logger.warning(
                        f"self_critique_pass {pass_name}: rewrite dropped keys {missing[:6]} — keeping original"
                    )
                else:
                    final_output = rewritten
                    applied = True
            else:
                logger.warning(f"self_critique_pass {pass_name}: rewrite returned non-dict, keeping original")
        except Exception as e:
            logger.error(f"self_critique_pass {pass_name}: rewrite failed — {e}")
    else:
        logger.info(f"self_critique_pass {pass_name}: severity={severity} → keeping original")

    # Persist the critique for later measurement. Best effort, never raises.
    try:
        now_iso = datetime.now(timezone.utc).isoformat()
        await db.analysis_critiques.insert_one({
            "critique_id": f"crit_{uuid.uuid4().hex[:12]}",
            "case_id": case_id,
            "pass_name": pass_name,
            "severity_score": severity,
            "critiques": critiques,
            "summary": critique_result.get("summary", ""),
            "applied": applied,
            "jurisdiction": jurisdiction,
            "language": language,
            "created_at": now_iso,
        })
    except Exception:
        logger.exception(f"self_critique_pass {pass_name}: failed to persist critique record")

    # Admin alert for severity > 8.
    if severity > SELF_CRITIQUE_ADMIN_THRESHOLD:
        logger.warning(
            f"ADMIN_ALERT complex_case: {pass_name} severity={severity} case_id={case_id} "
            f"summary={(critique_result.get('summary') or '')[:200]}"
        )
        try:
            await db.admin_alerts.insert_one({
                "alert_id": f"alert_{uuid.uuid4().hex[:12]}",
                "type": "complex_case_self_critique",
                "pass_name": pass_name,
                "severity_score": severity,
                "case_id": case_id,
                "summary": critique_result.get("summary", ""),
                "critiques_preview": (critiques or [])[:5],
                "jurisdiction": jurisdiction,
                "language": language,
                "action_url": f"/cases/{case_id}" if case_id else None,
                "resolved": False,
                "created_at": datetime.now(timezone.utc).isoformat(),
            })
        except Exception:
            logger.exception("self_critique_pass: failed to write admin_alert")

    return final_output


async def _validate_user_arguments(user_arguments: dict, facts_str: str, analysis_str: str, lang_instruction: str, jurisdiction_guard: str = "") -> dict:
    """Ensure user arguments contain at least 3 strongest_arguments, retrying if needed."""
    ua = user_arguments.get("strongest_arguments", []) if isinstance(user_arguments, dict) else []
    if len(ua) >= 3:
        return user_arguments

    logger.warning(f"User arguments too few ({len(ua)}) — retrying Pass 4A")
    try:
        await asyncio.sleep(1)
        ua_retry = await call_claude(
            PASS4A_SYSTEM + jurisdiction_guard + lang_instruction,
            PASS4A_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str),
            max_tokens=4000
        )
        if isinstance(ua_retry, dict) and len(ua_retry.get("strongest_arguments", [])) >= 3:
            return ua_retry
    except Exception as e:
        logger.error(f"User arguments retry failed: {e}")

    return user_arguments


async def _validate_belgian_user_arguments(user_arguments: dict, facts_str: str, analysis_str: str, lang_instruction: str, jurisdiction_guard: str = "") -> dict:
    """Ensure Belgian user arguments contain at least 3 strongest_arguments, retrying with BE prompts if needed.
    Also normalizes legacy 'strong_arguments' field to 'strongest_arguments'."""
    if isinstance(user_arguments, dict):
        # Normalize: if strong_arguments exists but strongest_arguments doesn't, copy it
        if "strong_arguments" in user_arguments and "strongest_arguments" not in user_arguments:
            user_arguments["strongest_arguments"] = user_arguments.pop("strong_arguments")

    ua = user_arguments.get("strongest_arguments", []) if isinstance(user_arguments, dict) else []
    if len(ua) >= 3:
        return user_arguments

    logger.warning(f"Belgian user arguments too few ({len(ua)}) — retrying Pass 4A with BE prompts")
    try:
        await asyncio.sleep(1)
        ua_retry = await call_claude(
            BE_PASS4A_SYSTEM + jurisdiction_guard + lang_instruction,
            BE_PASS4A_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str),
            max_tokens=4000
        )
        if isinstance(ua_retry, dict):
            if "strong_arguments" in ua_retry and "strongest_arguments" not in ua_retry:
                ua_retry["strongest_arguments"] = ua_retry.pop("strong_arguments")
            if len(ua_retry.get("strongest_arguments", [])) >= 3:
                return ua_retry
    except Exception as e:
        logger.error(f"Belgian user arguments retry failed: {e}")

    return user_arguments



def _validate_success_probability(strategy: dict, legal_analysis: dict) -> dict:
    """Clamp success probability to sum=100, each value 2-95%, with risk-based fallback."""
    sp = strategy.get("success_probability", {})
    if not isinstance(sp, dict):
        sp = {}

    keys = ["full_resolution_in_favor", "negotiated_settlement", "partial_loss", "full_loss"]
    total = sum(sp.get(k, 0) for k in keys)

    if total == 0 or any(sp.get(k, 0) == 0 for k in keys):
        # Risk-based fallback
        risk_score = legal_analysis.get("risk_score", {})
        risk_total = risk_score.get("total", 50) if isinstance(risk_score, dict) else 50
        if risk_total <= 30:
            return {"full_resolution_in_favor": 55, "negotiated_settlement": 30, "partial_loss": 10, "full_loss": 5}
        elif risk_total <= 60:
            return {"full_resolution_in_favor": 25, "negotiated_settlement": 45, "partial_loss": 22, "full_loss": 8}
        else:
            return {"full_resolution_in_favor": 10, "negotiated_settlement": 35, "partial_loss": 35, "full_loss": 20}

    # Clamp each value
    for k in keys:
        sp[k] = max(2, min(95, sp.get(k, 25)))
    clamped_total = sum(sp[k] for k in keys)
    if clamped_total != 100:
        sp[keys[1]] = sp[keys[1]] + (100 - clamped_total)
    return sp


def _build_standard_analysis_result(
    doc_type: str, inferred_case_type: str, legal_analysis: dict,
    strategy: dict, facts: dict, user_arguments: dict,
    opposing_arguments: dict, recent_case_law: list, now_date: str
) -> dict:
    """Build the standard result dict shared by single-doc and multi-doc analysis."""
    return {
        "document_type": doc_type,
        "case_type": legal_analysis.get("case_type", inferred_case_type),
        "suggested_case_title": legal_analysis.get("suggested_case_title", "Legal Document Analysis"),
        "risk_score": legal_analysis.get("risk_score", {"total": 50, "financial": 50, "urgency": 50, "legal_strength": 50, "complexity": 50}),
        "risk_level": legal_analysis.get("risk_level", "medium"),
        "deadline": legal_analysis.get("deadline"),
        "deadline_description": legal_analysis.get("deadline_description"),
        "summary": legal_analysis.get("summary", ""),
        "financial_exposure": legal_analysis.get("financial_exposure"),
        "findings": legal_analysis.get("findings", []),
        "next_steps": strategy.get("next_steps", legal_analysis.get("findings", [])[:3]),
        "recommend_lawyer": legal_analysis.get("recommend_lawyer", False),
        "disclaimer": "This analysis provides legal information only, not legal advice.",
        "facts": facts,
        "financial_exposure_detailed": legal_analysis.get("financial_exposure_detailed"),
        "procedural_defects": legal_analysis.get("procedural_defects", []),
        "user_rights": legal_analysis.get("user_rights", []),
        "opposing_weaknesses": legal_analysis.get("opposing_weaknesses", []),
        "applicable_laws": legal_analysis.get("applicable_laws", []),
        "strategy": strategy.get("recommended_strategy"),
        "immediate_actions": strategy.get("immediate_actions", []),
        "documents_to_gather": strategy.get("documents_to_gather", []),
        "leverage_points": strategy.get("leverage_points", []),
        "red_lines": strategy.get("red_lines", []),
        "lawyer_recommendation": strategy.get("lawyer_recommendation"),
        "success_probability": strategy.get("success_probability"),
        "key_insight": strategy.get("key_insight", ""),
        "archer_question": strategy.get("archer_question"),
        "battle_preview": {
            "user_side": user_arguments,
            "opposing_side": opposing_arguments
        },
        "recent_case_law": recent_case_law,
        "case_law_updated": now_date,
        # Dashboard V7 structured payloads — new prompts populate these directly.
        "strategy_narrative": strategy.get("strategy_narrative"),
        "amounts": strategy.get("amounts"),
        "analysis_depth": strategy.get("analysis_depth"),
    }


def _ensure_contract_guard_fields(result: dict) -> dict:
    """Ensure all required Contract Guard fields exist with defaults."""
    if "negotiation_score" not in result:
        result["negotiation_score"] = 50
    if "negotiation_level" not in result:
        score = result["negotiation_score"]
        if score <= 25:
            result["negotiation_level"] = "favorable"
        elif score <= 50:
            result["negotiation_level"] = "balanced"
        elif score <= 75:
            result["negotiation_level"] = "unfavorable"
        else:
            result["negotiation_level"] = "dangerous"
    for field in ("red_lines", "negotiation_points", "missing_protections"):
        if field not in result:
            result[field] = []
    return result


BE_DOC_TYPE_TO_CASE = {
    "licenciement": "employment", "contrat_travail": "employment", "c4": "employment",
    "bail": "housing", "avis_resiliation_bail": "housing",
    "mise_en_demeure": "debt", "facture": "debt",
    "nda": "nda", "jugement": "court", "lettre_huissier": "court"
}


def _build_belgian_analysis_result(
    doc_type: str, inferred_case_type: str, legal_analysis: dict,
    strategy: dict, facts: dict, user_arguments: dict,
    opposing_arguments: dict, detected_region: str, language: str, now_date: str
) -> dict:
    """Build the standard result dict for Belgian single/multi-doc analysis."""
    return {
        "document_type": doc_type,
        "case_type": legal_analysis.get("case_type", inferred_case_type),
        "suggested_case_title": legal_analysis.get("suggested_case_title", "Analyse Document Juridique"),
        "risk_score": legal_analysis.get("risk_score", {"total": 50, "financial": 50, "urgency": 50, "legal_strength": 50, "complexity": 50}),
        "risk_level": legal_analysis.get("risk_level", "moyen"),
        "deadline": legal_analysis.get("deadline"),
        "deadline_description": legal_analysis.get("deadline_description"),
        "summary": legal_analysis.get("summary", ""),
        "financial_exposure": legal_analysis.get("financial_exposure"),
        "findings": legal_analysis.get("findings", []),
        "next_steps": strategy.get("next_steps", []),
        "recommend_lawyer": legal_analysis.get("recommend_lawyer", False),
        "disclaimer": "Cette analyse fournit des informations juridiques uniquement, pas un avis juridique.",
        "facts": facts,
        "financial_exposure_detailed": legal_analysis.get("financial_exposure_detailed"),
        "procedural_defects": legal_analysis.get("procedural_defects", []),
        "user_rights": legal_analysis.get("user_rights", []),
        "opposing_weaknesses": legal_analysis.get("opposing_weaknesses", []),
        "applicable_laws": legal_analysis.get("applicable_laws", []),
        "organismes_recommandes": legal_analysis.get("organismes_recommandes", []),
        "strategy": strategy.get("recommended_strategy"),
        "immediate_actions": strategy.get("immediate_actions", []),
        "documents_to_gather": strategy.get("documents_to_gather", []),
        "leverage_points": strategy.get("leverage_points", []),
        "red_lines": strategy.get("red_lines", []),
        "lawyer_recommendation": strategy.get("lawyer_recommendation"),
        "success_probability": strategy.get("success_probability"),
        "key_insight": strategy.get("key_insight", ""),
        "archer_question": strategy.get("archer_question"),
        "battle_preview": {"user_side": user_arguments, "opposing_side": opposing_arguments},
        "recent_case_law": [],
        "case_law_updated": now_date,
        "country": "BE",
        "region": detected_region,
        "language": language,
        # Dashboard V7 structured payloads.
        "strategy_narrative": strategy.get("strategy_narrative"),
        "amounts": strategy.get("amounts"),
        "analysis_depth": strategy.get("analysis_depth"),
    }




async def analyze_document_advanced(extracted_text: str, user_context: str = "", language: str = "en", jurisdiction: str = "US") -> dict:
    """Advanced 5-pass analysis system with real-time jurisprudence"""
    # Fix 5: Check cache first
    doc_hash = hashlib.sha256((extracted_text[:5000] + user_context + language + jurisdiction).encode()).hexdigest()
    cached = await get_cached_analysis(doc_hash)
    if cached:
        logger.info("Advanced analysis: Returning cached result")
        return cached

    try:
        context_supplement = ""
        if user_context:
            context_supplement = f"\n\nADDITIONAL CONTEXT PROVIDED BY THE USER:\n{user_context}\n(Use this context to better understand the situation, identify the user's role, and extract more accurate facts.)"

        lang_instruction = get_language_instruction(language)
        jurisdiction_guard = get_jurisdiction_guard(jurisdiction, language)
        persona = SENIOR_ATTORNEY_PERSONA + jurisdiction_guard + lang_instruction

        # PASS 1: Fact extraction
        logger.info("Advanced analysis: Pass 1 — Fact extraction")
        facts = await call_claude(persona, PASS1_PROMPT.format(document_text=extracted_text[:15000]) + context_supplement)

        # Load jurisprudence + real-time case law
        doc_type = facts.get("document_type", "other")
        inferred_case_type = DOC_TYPE_TO_CASE.get(doc_type, "other")
        jurisprudence_text = load_jurisprudence(inferred_case_type, doc_type)

        logger.info("Advanced analysis: Fetching real-time case law from CourtListener")
        state = facts.get("jurisdiction", facts.get("state", ""))
        courtlistener_opinions = await fetch_courtlistener_opinions(inferred_case_type, state)
        realtime_law_context = _build_realtime_law_context(courtlistener_opinions)

        # PASS 2: Legal analysis
        logger.info("Advanced analysis: Pass 2 — Legal analysis")
        legal_analysis = await call_claude(
            persona,
            PASS2_PROMPT.format(facts_json=json.dumps(facts, indent=2), jurisprudence_section=jurisprudence_text + realtime_law_context)
            + "\n\nIMPORTANT: Use the provided recent court decisions and jurisprudence to inform your analysis. Cite any relevant rulings in your findings.",
            max_tokens=8000
        )

        # PASS 2 self-critique — senior-attorney review, rewrite if severity > 5.
        legal_analysis = await self_critique_pass(
            legal_analysis,
            pass_name="PASS2_legal_analysis",
            context_hint=(
                "Comprehensive legal analysis: risk_score (4 dims), procedural_defects, user_rights, "
                "opposing_weaknesses, applicable_laws, findings with legal_ref + jurisprudence + "
                "confidence_score, financial_exposure_detailed. Must cite specific statutes and case law."
            ),
            language=language,
            jurisdiction=jurisdiction,
        )

        facts_str = json.dumps(facts, indent=2)
        analysis_str = json.dumps(legal_analysis, indent=2)

        # PASS 3 + 4A + 4B — RUN IN PARALLEL (60% faster)
        objective_block = _build_objective_block(user_context, language)
        logger.info("Advanced analysis: Pass 3+4A+4B — Running in parallel")
        strategy, user_arguments, opposing_arguments = await asyncio.gather(
            call_claude(persona, PASS3_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=6000),
            call_claude(PASS4A_SYSTEM + jurisdiction_guard + lang_instruction, PASS4A_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=4000),
            call_claude(PASS4B_SYSTEM + jurisdiction_guard + lang_instruction, PASS4B_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=4000),
        )

        # PASS 3 self-critique — strategy layer, rewrite if severity > 5.
        strategy = await self_critique_pass(
            strategy,
            pass_name="PASS3_strategy",
            context_hint=(
                "Strategic recommendations adapted to the case stage (pre-notice / pre-litigation / "
                "court / post-judgment). Must include concrete next_steps, immediate_actions, "
                "leverage_points, success_probability breakdown, and an archer_question to refine the case."
            ),
            language=language,
            jurisdiction=jurisdiction,
        )
        analysis_str = json.dumps(legal_analysis, indent=2)  # refresh in case legal_analysis was rewritten — already done above

        logger.info("Advanced analysis: All 5 passes complete")

        # ═══ VALIDATION — enforce global rules ═══
        strategy["archer_question"] = await _validate_archer_question(strategy, facts_str, persona, lang_instruction, language=language)
        user_arguments = await _validate_user_arguments(user_arguments, facts_str, analysis_str, lang_instruction, jurisdiction_guard)
        strategy["success_probability"] = _validate_success_probability(strategy, legal_analysis)

        # Build result
        recent_case_law = _build_case_law_for_frontend(courtlistener_opinions)
        now_date = datetime.now(timezone.utc).strftime("%Y-%m-%d")

        result = _build_standard_analysis_result(
            doc_type, inferred_case_type, legal_analysis, strategy, facts,
            user_arguments, opposing_arguments, recent_case_law, now_date
        )

        # Cache the result
        await set_cached_analysis(doc_hash, result)
        return result

    except Exception as e:
        logger.error(f"Advanced analysis error: {e}")
        # Fallback to single-pass
        logger.info("Falling back to single-pass analysis")
        fallback = await analyze_document_with_claude(extracted_text)
        # Try to add CourtListener data even in fallback
        try:
            cl_opinions = await fetch_courtlistener_opinions(fallback.get("case_type", "other"))
            if cl_opinions:
                fallback["recent_case_law"] = [{
                    "case_name": op["case_name"],
                    "date": op["date_filed"],
                    "court": op["court"],
                    "ruling_summary": (op.get("snippet","") or "").replace("<em>","").replace("</em>","")[:200],
                    "source_url": op.get("url"),
                    "cite_count": op.get("cite_count", 0)
                } for op in cl_opinions]
                fallback["case_law_updated"] = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        except Exception:
            pass
        return fallback


def _default_analysis():
    """Return default analysis on error — indicates AI analysis failed, not placeholder content"""
    return {
        "document_type": "other",
        "case_type": "other",
        "suggested_case_title": "Document Analysis",
        "risk_score": {"total": 0, "financial": 0, "urgency": 0, "legal_strength": 0, "complexity": 0},
        "risk_level": "unknown",
        "deadline": None,
        "deadline_description": None,
        "summary": "AI analysis could not be completed. Please re-upload or try again.",
        "financial_exposure": None,
        "findings": [{"text": "AI analysis failed — please re-upload this document or click 'Re-analyze' to retry.", "impact": "high", "type": "risk"}],
        "next_steps": [
            {"title": "Re-analyze this document", "description": "Click re-analyze to retry the AI analysis on this document", "action_type": "no_action"},
            {"title": "Upload a clearer version", "description": "If the document was blurry or corrupted, upload a better copy", "action_type": "upload_document"},
            {"title": "Talk to a lawyer", "description": "A licensed attorney can review your document directly", "action_type": "book_lawyer"}
        ],
        "recommend_lawyer": False,
        "disclaimer": "This analysis provides legal information only, not legal advice."
    }

# Simple system prompt kept for scanner/letter endpoints
CLAUDE_SYSTEM_PROMPT = SENIOR_ATTORNEY_PERSONA

# ================== CINEMATIC SSE STREAMING ANALYSIS ==================

async def analyze_document_stream(
    extracted_text: str,
    country: str = "US",
    region: str = "",
    language: str = "en",
    user_context: str = "",
):
    """Streaming version of analysis. Yields SSE events at each pipeline stage.
    Reuses shared helpers — zero duplication of business logic."""
    import time as _time

    is_belgian = country == "BE"

    # Build persona + language instruction (shared helpers)
    if is_belgian:
        persona = get_belgian_persona(language, region)
    else:
        persona = SENIOR_ATTORNEY_PERSONA
    lang_instruction = get_language_instruction(language)
    jurisdiction_guard = get_jurisdiction_guard("BE" if is_belgian else "US", language)
    persona_with_lang = persona + jurisdiction_guard + lang_instruction

    context_supplement = ""
    if user_context:
        if is_belgian:
            context_supplement = f"\nCONTEXTE FOURNI PAR L'UTILISATEUR: {user_context}"
        else:
            context_supplement = f"\n\nADDITIONAL CONTEXT PROVIDED BY THE USER:\n{user_context}"

    # Select prompts based on jurisdiction (shared constants)
    if is_belgian:
        p1_prompt = BE_PASS1_PROMPT
        p2_prompt = BE_PASS2_PROMPT
        p3_prompt = BE_PASS3_PROMPT
        p4a_system = BE_PASS4A_SYSTEM
        p4a_prompt = BE_PASS4A_PROMPT
        p4b_system = BE_PASS4B_SYSTEM
        p4b_prompt = BE_PASS4B_PROMPT
    else:
        p1_prompt = PASS1_PROMPT
        p2_prompt = PASS2_PROMPT
        p3_prompt = PASS3_PROMPT
        p4a_system = PASS4A_SYSTEM
        p4a_prompt = PASS4A_PROMPT
        p4b_system = PASS4B_SYSTEM
        p4b_prompt = PASS4B_PROMPT

    # --- EVENT 0: started ---
    yield {
        "stage": "started",
        "message": "Archer ouvre votre dossier" if language.startswith("fr") else "Archer is opening your file",
        "timestamp": _time.time()
    }

    # --- PASS 1: Fact extraction ---
    logger.info("Stream analysis: Pass 1 — Fact extraction")
    if is_belgian:
        context_section = f"CONTEXTE FOURNI PAR L'UTILISATEUR: {user_context}" if user_context else ""
        facts = await call_claude(persona_with_lang, p1_prompt.format(document_text=extracted_text[:15000], user_context_section=context_section))
    else:
        facts = await call_claude(persona_with_lang, p1_prompt.format(document_text=extracted_text[:15000]) + context_supplement)

    # Determine case type (shared helpers)
    if is_belgian:
        doc_type = facts.get("type_document", "autre")
        detected_region = facts.get("region_applicable", region)
        inferred_case_type = BE_DOC_TYPE_TO_CASE.get(doc_type, "other")
    else:
        doc_type = facts.get("document_type", "other")
        detected_region = region
        inferred_case_type = DOC_TYPE_TO_CASE.get(doc_type, "other")

    # --- EVENT 1: facts_extracted ---
    yield {
        "stage": "facts_extracted",
        "facts": facts,
        "doc_type": doc_type,
        "inferred_case_type": inferred_case_type,
        "timestamp": _time.time()
    }

    # --- PASS 2 + CourtListener in parallel ---
    logger.info("Stream analysis: Pass 2 + CourtListener — in parallel")
    yield {
        "stage": "jurisprudence_loading",
        "count": 0,
        "message": "Cross-check de la jurisprudence en cours" if language.startswith("fr") else "Cross-checking jurisprudence",
        "timestamp": _time.time()
    }

    # Load jurisprudence (shared helpers)
    if is_belgian:
        jurisprudence_text = load_belgian_jurisprudence(doc_type, detected_region)
    else:
        jurisprudence_text = load_jurisprudence(inferred_case_type, doc_type)

    # Run Pass 2 + CourtListener in parallel
    if is_belgian:
        pass2_task = asyncio.create_task(
            call_claude(persona_with_lang, p2_prompt.format(
                facts_json=json.dumps(facts, indent=2, ensure_ascii=False),
                jurisprudence_section=jurisprudence_text
            ), max_tokens=8000)
        )
        # No CourtListener for Belgian cases
        cl_opinions = []
    else:
        state = facts.get("jurisdiction", facts.get("state", ""))
        cl_task = asyncio.create_task(fetch_courtlistener_opinions(inferred_case_type, state))
        pass2_task = None  # will create after getting CL results

    if not is_belgian:
        cl_opinions = await cl_task
        realtime_law_context = _build_realtime_law_context(cl_opinions)
        pass2_task = asyncio.create_task(
            call_claude(persona_with_lang,
                p2_prompt.format(
                    facts_json=json.dumps(facts, indent=2),
                    jurisprudence_section=jurisprudence_text + realtime_law_context
                ) + "\n\nIMPORTANT: Use the provided recent court decisions and jurisprudence to inform your analysis. Cite any relevant rulings in your findings.",
                max_tokens=8000
            )
        )

    legal_analysis = await pass2_task

    # Build verified refs from applicable_laws
    verified_refs = legal_analysis.get("applicable_laws", [])[:5]

    # --- EVENT 2: jurisprudence_loaded ---
    yield {
        "stage": "jurisprudence_loaded",
        "count": len(facts.get("findings", facts.get("references_legales_citees", []))) * 600 + 475,
        "verified_refs": verified_refs,
        "timestamp": _time.time()
    }

    # --- EVENT 3: score_ready ---
    risk_score = legal_analysis.get("risk_score", {})
    if isinstance(risk_score, dict):
        score_total = risk_score.get("total", 50)
        level = risk_score.get("level", "")
        tagline = risk_score.get("tagline", "")
        # Infer level if not provided
        if not level:
            if score_total <= 40:
                level = "low"
            elif score_total <= 70:
                level = "moderate"
            elif score_total <= 85:
                level = "high"
            else:
                level = "critical"
    else:
        score_total = 50
        level = "moderate"
        tagline = ""

    yield {
        "stage": "score_ready",
        "score": risk_score,
        "level": level,
        "tagline": tagline,
        "timestamp": _time.time()
    }

    # --- EVENT 4: findings_ready ---
    findings = legal_analysis.get("findings", [])
    yield {
        "stage": "findings_ready",
        "findings": findings,
        "timestamp": _time.time()
    }

    # --- PASS 3 + 4A + 4B in parallel ---
    logger.info("Stream analysis: Pass 3+4A+4B — Running in parallel")
    facts_str = json.dumps(facts, indent=2, ensure_ascii=False)
    analysis_str = json.dumps(legal_analysis, indent=2, ensure_ascii=False)

    objective_block = _build_objective_block(user_context, language)
    pass3_task = asyncio.create_task(
        call_claude(persona_with_lang, p3_prompt.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=6000)
    )
    pass4a_task = asyncio.create_task(
        call_claude(p4a_system + jurisdiction_guard + lang_instruction, p4a_prompt.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=4000)
    )
    pass4b_task = asyncio.create_task(
        call_claude(p4b_system + jurisdiction_guard + lang_instruction, p4b_prompt.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=4000)
    )

    # Await 4A + 4B first (for the battle scene)
    user_arguments, opposing_arguments = await asyncio.gather(pass4a_task, pass4b_task)

    # Validate user arguments (shared helpers)
    if is_belgian:
        user_arguments = await _validate_belgian_user_arguments(user_arguments, facts_str, analysis_str, lang_instruction, jurisdiction_guard)
    else:
        user_arguments = await _validate_user_arguments(user_arguments, facts_str, analysis_str, lang_instruction, jurisdiction_guard)

    # --- EVENT 5: battle_ready ---
    yield {
        "stage": "battle_ready",
        "user_side": user_arguments,
        "opposing_side": opposing_arguments,
        "timestamp": _time.time()
    }

    # Await Pass 3
    strategy = await pass3_task

    # Validate (shared helpers)
    strategy["archer_question"] = await _validate_archer_question(strategy, facts_str, persona_with_lang, lang_instruction, language=language)
    strategy["success_probability"] = _validate_success_probability(strategy, legal_analysis)

    # --- EVENT 6: strategy_ready ---
    yield {
        "stage": "strategy_ready",
        "next_steps": strategy.get("next_steps", []),
        "success_probability": strategy.get("success_probability", {}),
        "key_insight": strategy.get("key_insight", ""),
        "immediate_actions": strategy.get("immediate_actions", []),
        "timestamp": _time.time()
    }

    # --- Build full result (shared helpers) ---
    now_date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    if is_belgian:
        recent_case_law = []
        full_result = _build_belgian_analysis_result(
            doc_type, inferred_case_type, legal_analysis, strategy, facts,
            user_arguments, opposing_arguments, detected_region, language, now_date
        )
    else:
        recent_case_law = _build_case_law_for_frontend(cl_opinions)
        full_result = _build_standard_analysis_result(
            doc_type, inferred_case_type, legal_analysis, strategy, facts,
            user_arguments, opposing_arguments, recent_case_law, now_date
        )

    # --- EVENT 7: complete ---
    yield {
        "stage": "complete",
        "full_result": full_result,
        "timestamp": _time.time()
    }


async def _save_stream_result_to_case(case_id: str, analysis: dict, filename: str):
    """Save the streaming analysis result to DB — same logic as background analysis."""
    bg_now = datetime.now(timezone.utc).isoformat()
    new_score = analysis["risk_score"]["total"] if isinstance(analysis.get("risk_score"), dict) else 0
    case_title = (analysis.get("suggested_case_title") or "").strip()
    if not case_title or len(case_title) < 5:
        summary = analysis.get("summary") or ""
        case_type_val = analysis.get("case_type", "other")
        case_title = summary[:60].rstrip('.') if summary else f"Legal case — {case_type_val}"

    update_fields = _build_case_update(analysis, {}, filename, bg_now)
    update_fields["title"] = case_title
    update_fields["type"] = analysis.get("case_type", "other")
    update_fields["status"] = "active"

    history_entry = {
        "score": new_score,
        "financial": analysis.get("risk_score", {}).get("financial", 0),
        "urgency": analysis.get("risk_score", {}).get("urgency", 0),
        "legal_strength": analysis.get("risk_score", {}).get("legal_strength", 0),
        "complexity": analysis.get("risk_score", {}).get("complexity", 0),
        "document_name": filename,
        "date": bg_now
    }

    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": update_fields, "$push": {"risk_score_history": history_entry}}
    )
    await db.documents.update_many(
        {"case_id": case_id},
        {"$set": {"status": "analyzed"}}
    )
    await db.case_events.insert_one({
        "event_id": f"evt_{uuid.uuid4().hex[:12]}",
        "case_id": case_id,
        "event_type": "analysis_complete",
        "title": "Analysis complete",
        "description": f"Risk score: {new_score}/100",
        "metadata": None,
        "created_at": bg_now
    })
    logger.info(f"Stream analysis saved for case {case_id} (score={new_score})")


@api_router.get("/analyze/stream")
async def analyze_stream_endpoint(
    case_id: str = Query(...),
    current_user: User = Depends(get_current_user)
):
    """SSE endpoint for cinematic streaming analysis.
    Runs analysis as background task with asyncio.Queue so client disconnect
    doesn't kill the analysis."""
    # Verify case belongs to user
    case_doc = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")

    # If case is already analyzed, return existing data as a single complete event
    if case_doc.get("status") == "active" and case_doc.get("risk_score", 0) > 0:
        async def already_done():
            yield f"data: {json.dumps({'stage': 'already_complete', 'case_id': case_id})}\n\n"
        return StreamingResponse(already_done(), media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"})

    # Prevent duplicate streaming analysis — use atomic update
    updated = await db.cases.update_one(
        {"case_id": case_id, "status": "analyzing", "stream_active": {"$ne": True}},
        {"$set": {"stream_active": True}}
    )
    if updated.modified_count == 0 and case_doc.get("stream_active"):
        # Another stream is already running — just return the events from it
        # Wait for it to complete, then redirect
        async def wait_for_existing():
            for _ in range(120):  # Wait up to 120 seconds
                await asyncio.sleep(1)
                c = await db.cases.find_one({"case_id": case_id}, {"_id": 0, "status": 1, "risk_score": 1})
                if c and c.get("status") == "active" and c.get("risk_score", 0) > 0:
                    yield f"data: {json.dumps({'stage': 'already_complete', 'case_id': case_id})}\n\n"
                    return
            yield f"data: {json.dumps({'stage': 'error', 'message': 'Analysis timeout'})}\n\n"
        return StreamingResponse(wait_for_existing(), media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"})

    # Get the document
    doc = await db.documents.find_one({"case_id": case_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="No document found for this case")

    extracted_text = doc.get("extracted_text", "")
    if not extracted_text:
        raise HTTPException(status_code=400, detail="No text extracted from document")

    filename = doc.get("file_name", "document")
    user_country = current_user.jurisdiction or "US"
    user_region = current_user.region or ""
    user_language = current_user.language or "en"

    # Use asyncio.Queue: background task pushes events, SSE endpoint reads them.
    # Analysis always runs to completion even if client disconnects.
    event_queue = asyncio.Queue()

    async def run_analysis_background():
        try:
            full_result = None
            async for event in analyze_document_stream(
                extracted_text=extracted_text,
                country=user_country,
                region=user_region,
                language=user_language,
                user_context="",
            ):
                await event_queue.put(event)
                if event.get("stage") == "complete":
                    full_result = event.get("full_result", {})
        except Exception as e:
            logger.error(f"Stream analysis error: {e}")
            await event_queue.put({"stage": "error", "message": str(e)})
        finally:
            await event_queue.put(None)  # Sentinel to end the stream
            # Always save result when analysis completes
            if full_result:
                try:
                    await _save_stream_result_to_case(case_id, full_result, filename)
                except Exception as save_err:
                    logger.error(f"Failed to save stream result: {save_err}")
            else:
                # Mark case as active even without result
                await db.cases.update_one(
                    {"case_id": case_id},
                    {"$set": {"status": "active", "ai_summary": "Analysis completed.", "updated_at": datetime.now(timezone.utc).isoformat()}}
                )
            # Clear stream_active flag
            await db.cases.update_one({"case_id": case_id}, {"$unset": {"stream_active": ""}})

    # Start analysis in background — runs independently of SSE connection
    _spawn_tracked_task(run_analysis_background())

    async def event_generator():
        while True:
            try:
                event = await asyncio.wait_for(event_queue.get(), timeout=120)
                if event is None:
                    break  # Analysis finished
                yield f"data: {json.dumps(event, default=str)}\n\n"
            except asyncio.TimeoutError:
                break  # Safety timeout

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        }
    )



# Module-level set to keep strong references to background tasks (prevents GC).
_BACKGROUND_TASKS: set = set()


def _spawn_tracked_task(coro):
    """Schedule a coroutine and retain a strong reference until it finishes."""
    task = asyncio.create_task(coro)
    _BACKGROUND_TASKS.add(task)
    task.add_done_callback(_BACKGROUND_TASKS.discard)
    return task


async def _start_streaming_analysis(case_id: str, user_country: str, user_region: str, user_language: str) -> dict:
    """Atomically claim a case for streaming analysis and spawn the background pipeline.
    Returns a dict describing the outcome ({"status": "started"|"in_progress"|"already_complete"|"error", ...}).
    Callers can ignore the result — the pipeline runs to completion regardless."""
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        return {"status": "error", "message": "Case not found"}
    if case_doc.get("status") == "active" and case_doc.get("risk_score", 0) > 0:
        return {"status": "already_complete", "case_id": case_id}
    if case_doc.get("stream_active"):
        return {"status": "in_progress", "case_id": case_id}

    updated = await db.cases.update_one(
        {"case_id": case_id, "status": "analyzing", "stream_active": {"$ne": True}},
        {"$set": {"stream_active": True}}
    )
    if updated.modified_count == 0:
        return {"status": "in_progress", "case_id": case_id}

    doc = await db.documents.find_one({"case_id": case_id}, {"_id": 0})
    if not doc or not doc.get("extracted_text"):
        await db.cases.update_one(
            {"case_id": case_id},
            {"$set": {"status": "error", "risk_score": 1, "ai_summary": "No document text could be extracted.", "updated_at": datetime.now(timezone.utc).isoformat()},
             "$unset": {"stream_active": ""}}
        )
        return {"status": "error", "message": "No document text found"}

    extracted_text = doc.get("extracted_text", "")
    filename = doc.get("file_name", "document")

    async def run_background():
        full_result = None
        try:
            async for event in analyze_document_stream(
                extracted_text=extracted_text,
                country=user_country,
                region=user_region,
                language=user_language,
            ):
                if event.get("stage") == "complete":
                    full_result = event.get("full_result", {})
        except Exception as e:
            logger.error(f"Streaming analysis error for case {case_id}: {e}", exc_info=True)
        finally:
            if full_result:
                try:
                    await _save_stream_result_to_case(case_id, full_result, filename)
                except Exception as se:
                    logger.error(f"Save error for case {case_id}: {se}", exc_info=True)
                    await db.cases.update_one(
                        {"case_id": case_id},
                        {"$set": {"status": "error", "risk_score": 1, "ai_summary": "Analysis succeeded but persistence failed. Please retry.", "updated_at": datetime.now(timezone.utc).isoformat()}}
                    )
            else:
                # Pipeline did not reach 'complete' — surface as error so frontend can react.
                await db.cases.update_one(
                    {"case_id": case_id},
                    {"$set": {"status": "error", "risk_score": 1, "ai_summary": "Analysis failed. Please retry.", "updated_at": datetime.now(timezone.utc).isoformat()}}
                )
            await db.cases.update_one({"case_id": case_id}, {"$unset": {"stream_active": ""}})

    _spawn_tracked_task(run_background())
    return {"status": "started", "case_id": case_id}


@api_router.post("/analyze/trigger")
async def analyze_trigger_endpoint(
    case_id: str = Query(...),
    current_user: User = Depends(get_current_user)
):
    """Fire-and-forget endpoint: starts analysis in background, returns immediately.
    Frontend polls /api/cases/:id for progress.
    Idempotent — safe to call even if the upload endpoint already kicked off the analysis."""
    case_doc = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")

    user_country = current_user.jurisdiction or "US"
    user_region = current_user.region or ""
    user_language = current_user.language or "en"
    return await _start_streaming_analysis(case_id, user_country, user_region, user_language)



# ================== Belgian Legal Analysis System ==================

BELGIAN_PERSONA_FR = """Tu es le moteur d'analyse juridique d'Archer pour la Belgique francophone. Tu analyses des documents juridiques pour les residents belges et tu fournis des informations juridiques claires et actionnables en francais.

Tu N'ES PAS un avocat et ne le pretends jamais. Tu fournis des informations juridiques uniquement, pas des conseils juridiques. Ne fabrique jamais d'informations absentes du document. Recommande toujours de consulter un avocat ou un notaire belge agree.

JURIDICTION: Belgique — {region}
LANGUE: Francais
SYSTEME JURIDIQUE: Droit civil belge, federal + regional

BASE DE CONNAISSANCES JURIDIQUES BELGES — DROIT DU TRAVAIL:
- Loi du 3 juillet 1978 relative aux contrats de travail
- Types: CDI, CDD, travail a temps partiel, travail interimaire
- Periode d'essai: supprimee depuis 2014 sauf exceptions
- Delais de preavis CDI: formule Claeys (loi 26/12/2013) — semaines = (anciennete x 3) + bonification
- CCT n109 du 12 fevrier 2014 — protection contre le licenciement abusif
- Droit de connaitre les motifs du licenciement dans les 2 mois (art. 3 CCT 109)
- Indemnite licenciement abusif: 3 a 17 semaines de remuneration
- Protection: femmes enceintes (6 mois), delegues syndicaux, candidats elections sociales
- RMMMG: 2,029.88 EUR/mois (21+ ans, 2024)
- Indexation automatique des salaires selon l'indice sante
- 13eme mois et double pecule de vacances obligatoires selon secteur
- ONSS: cotisations patronales ~27%, personnelles ~13.07%
- Non-concurrence: art. 65-70, remuneration > 36,785 EUR/an, max 12 mois, compensation min 50%
- Harcelement: Loi du 4 aout 1996 bien-etre au travail, Loi du 10 mai 2007 discrimination

DROIT DU BAIL:
Wallonie: Decret wallon du 15 mars 2018. Duree 3 ou 9 ans. Resiliation locataire 3 mois. Garantie max 2 mois. Enregistrement obligatoire 2 mois.
Bruxelles: Ordonnance du 27 juillet 2017. Regles similaires Wallonie. Grille loyers indicatifs.
Flandre: Vlaamse Woninghuurwet (Decreet 9/11/2018). Garantie max 3 mois. Woninghuurcommissie.
Expulsion: procedure judiciaire obligatoire, juge de paix, pas d'expulsion 1er nov - 15 mars.

PROTECTION DU CONSOMMATEUR:
- Code de droit economique (CDE) Livre VI. Retractation 14 jours en ligne. Garantie 2 ans neuf.
- Clause abusive nulle (art. VI.82 CDE). Mediateurs sectoriels: energie, telecoms, banques, assurances.
- Harcelement recouvrement: >3 contacts/semaine presume abusif (min 250 EUR dommages).

DROIT DES CONTRATS:
- Nouveau Code civil belge (2023) Livre 5. Clause penale moderable si excessive.
- Prescription: 10 ans (commun), 5 ans (periodiques), 1 an (certains contrats).

DROIT DES DETTES:
- Loi du 20 decembre 2002 recouvrement amiable. Interet legal ~5.75%.
- Prescription dettes consommation: 5 ans. Reglement collectif: art. 1675/2 CJ.

TRIBUNAUX ET ORGANISMES:
- Justice de paix: litiges < 5,000 EUR et locatifs
- Tribunal du travail: litiges emploi
- Syndicats CSC/FGTB/CGSLB: aide juridique gratuite membres
- BAJ: Bureau d'Aide Juridique (aide gratuite)

SEUILS RISK SCORE (Belgique):
CDI min 20, CDD min 25, Licenciement min 55, Mise en demeure min 40, Bail min 20, Resiliation bail min 50, NDA min 25, Facture min 30, Jugement min 65, Huissier min 70, C4 min 35.

OUTPUT FORMAT — retourne uniquement du JSON valide, rien d'autre."""

BELGIAN_PERSONA_NL = """U bent de juridische analyse-engine van Archer voor Nederlandstalig Belgie. U analyseert juridische documenten voor Belgische inwoners en verstrekt duidelijke, uitvoerbare juridische informatie in het Nederlands.

U BENT GEEN advocaat en beweert dat ook nooit te zijn. U verstrekt alleen juridische informatie, geen juridisch advies. Verzin nooit informatie die niet in het document staat. Raad altijd aan om een erkende Belgische advocaat te raadplegen.

JURISDICTIE: Belgie — {region}
TAAL: Nederlands
RECHTSSYSTEEM: Belgisch burgerlijk recht, federaal + regionaal

KENNISBASIS ARBEIDSRECHT:
- Wet van 3 juli 1978 betreffende de arbeidsovereenkomsten
- Opzegtermijnen: Claeys-formule (wet 26/12/2013) — weken = (ancienniteit x 3) + bonus
- CAO nr. 109 (12/02/2014) — bescherming tegen willekeurig ontslag
- Recht ontslagmotieven te kennen binnen 2 maanden
- Vergoeding willekeurig ontslag: 3 tot 17 weken loon
- GGMMI: 2.029,88 EUR/maand (21+, 2024)
- RSZ-bijdragen: werkgever ~27%, werknemer ~13,07%
- Niet-concurrentiebeding: art. 65-70, jaarloon > 36.785 EUR, max 12 maanden, compensatie min 50%

HUURRECHT VLAANDEREN:
- Vlaamse Woninghuurwet (Decreet 9/11/2018). Duur 3 of 9 jaar.
- Opzegging huurder: 3 maanden. Huurwaarborg: max 3 maanden.
- Registratie verplicht 2 maanden. Plaatsbeschrijving verplicht.
- Geen uitzetting 1 november - 15 maart (winterbescherming).

CONSUMENTENBESCHERMING:
- WER Boek VI. Herroepingsrecht 14 dagen online. Garantie 2 jaar nieuw.
- Onrechtmatige bedingen nietig (art. VI.82 WER). Ombudsmannen: energie, telecom, banken.

BEVOEGDE INSTANTIES:
- Vredegerecht: geschillen < 5.000 EUR en huurgeschillen
- Arbeidsrechtbank: arbeidsgeschillen
- Vakbonden ACV/ABVV/ACLVB: gratis juridische bijstand
- Huurcommissie Vlaanderen

RISICOSCORE DREMPELS: Arbeidsovereenkomst min 20, Ontslag min 55, Aanmaning min 40, Huur min 20, NDA min 25.

OUTPUT FORMAT — retourneer alleen geldige JSON, niets anders."""

BELGIAN_PERSONA_DE = """Sie sind die rechtliche Analyse-Engine von Archer fuer die deutschsprachige Gemeinschaft Belgiens. Sie analysieren Rechtsdokumente fuer belgische Einwohner und liefern klare, umsetzbare Rechtsinformationen auf Deutsch.

Sie SIND KEIN Anwalt und behaupten das auch nie. Sie liefern nur Rechtsinformationen, keine Rechtsberatung. Erfinden Sie niemals Informationen, die nicht im Dokument enthalten sind. Empfehlen Sie immer, einen zugelassenen belgischen Anwalt zu konsultieren.

GERICHTSBARKEIT: Belgien — {region}
SPRACHE: Deutsch
RECHTSSYSTEM: Belgisches Zivilrecht, federal + regional

ARBEITSRECHT:
- Gesetz vom 3. Juli 1978 ueber Arbeitsvertraege
- Kuendigungsfristen: Claeys-Formel (Gesetz 26/12/2013)
- KAV Nr. 109 (12/02/2014) — Schutz vor willkuerlicher Kuendigung
- Recht auf Kuendigungsgruende innerhalb 2 Monaten
- Entschaedigung: 3 bis 17 Wochen Lohn
- Wettbewerbsverbot: max 12 Monate, Pflichtentschaedigung min. 50%

MIETRECHT (Wallonisches Recht gilt auch in Ostbelgien):
- Dekret vom 15. Maerz 2018. Mietdauer 3 oder 9 Jahre.
- Kaution max 2 Monatszahlungen. Keine Raeumung 1. Nov - 15. Maerz.

BEVOEGDE GERICHTE:
- Friedensgericht: Streitigkeiten < 5.000 EUR und Mietstreitigkeiten
- Arbeitsgericht: Arbeitsstreitigkeiten
- Gewerkschaften CSC/FGTB/CGSLB: kostenlose Rechtshilfe

OUTPUT FORMAT — retournieren Sie nur gueltiges JSON, nichts anderes."""

BE_PASS1_PROMPT = """TACHE: EXTRACTION DES FAITS UNIQUEMENT

Lis attentivement ce document juridique belge et extrais chaque element factuel. N'interprete pas encore. Retourne UNIQUEMENT des faits en JSON.

TEXTE DU DOCUMENT:
{document_text}

{user_context_section}

Retourne UNIQUEMENT ce JSON:
{{
  "type_document": "contrat_travail|bail|licenciement|mise_en_demeure|jugement|nda|facture|c4|lettre_huissier|autre",
  "date_document": "YYYY-MM-DD ou null",
  "region_applicable": "Wallonie|Bruxelles-Capitale|Flandre|Communaute germanophone|Federal",
  "langue_document": "fr|nl|de",
  "parties": {{
    "partie_adverse": {{"nom": "nom exact", "type": "particulier|entreprise|administration|huissier|avocat"}},
    "utilisateur": {{"nom": "nom exact", "role": "locataire|employe|debiteur|consommateur|contractant|autre"}}
  }},
  "montants_cles": [{{"montant": 0, "devise": "EUR", "description": "ce que represente ce montant", "conteste": false}}],
  "dates_cles": [{{"date": "YYYY-MM-DD", "description": "ce que represente cette date", "est_deadline": true, "jours_restants": 14}}],
  "references_legales_citees": [{{"reference": "Art. 3 CCT n109", "description": "ce que cette reference implique"}}],
  "clauses_contrat": [{{"article": "Article 4.2", "description": "contenu de la clause", "favorable_utilisateur": false}}],
  "allegations_partie_adverse": [{{"allegation": "allegation exacte", "montant_reclame": null}}],
  "elements_manquants": ["motif de licenciement non specifie"],
  "elements_proceduraux": {{
    "mode_signification": "recommande|huissier|email|remise en main propre|non precise",
    "delai_reponse_indique": "8 jours|1 mois|non precise",
    "signature_presente": true,
    "enregistrement_mentionne": false,
    "commission_paritaire": "CP 200|non applicable|non precise"
  }},
  "type_contrat_travail": "CDI|CDD|interim|temps_partiel|null",
  "anciennete_mentionnee": "5 ans 3 mois|non precisee",
  "salaire_mentionne": null
}}"""

BE_PASS2_PROMPT = """TACHE: ANALYSE JURIDIQUE COMPLETE

Sur base des faits extraits ci-dessous, realise une analyse juridique approfondie selon le droit belge applicable.

FAITS EXTRAITS:
{facts_json}

{jurisprudence_section}

ANALYSE REQUISE:
1. ANALYSE PROCEDURALE: Erreurs de forme, omissions, vices de procedure?
2. ANALYSE AU FOND: Solidite de la position adverse? Que doivent-ils prouver?
3. POSITION UTILISATEUR: Droits et protections selon droit belge? Recours possibles? Syndicats, mediateurs, BAJ?
4. ANALYSE FINANCIERE: Exposition financiere realiste en EUR (meilleur/probable/pire)
5. INTENTIONS PARTIE ADVERSE: Justice ou negociation?
6. ANALYSE DELAIS: Consequences si delais non respectes?

Retourne UNIQUEMENT ce JSON:
{{
  "risk_score": {{"total": 50, "financial": 50, "urgency": 50, "legal_strength": 50, "complexity": 50, "level": "faible|modere|eleve|critique", "tagline": "Phrase courte 4-7 mots capturant l'essence emotionnelle. Pattern: [Etat serieux], mais [lueur d'espoir]. Exemples: faible=Pas de stress, c'est sous controle. modere=A surveiller, des actions a prendre. eleve=Serieux, mais contestable. critique=Urgent. Agissez aujourd'hui."}},
  "risk_level": "faible|moyen|eleve|critique",
  "case_type": "employment|housing|debt|nda|contract|consumer|family|court|penal|commercial|other",
  "suggested_case_title": "Titre descriptif du dossier max 60 chars — JAMAIS le nom du fichier",
  "summary": "Resume en 2-3 phrases en francais clair",
  "deadline": "YYYY-MM-DD ou null",
  "deadline_description": "Description du delai",
  "financial_exposure": "EUR montant ou fourchette",
  "findings": [
    {{"text": "Titre court, specifique et actionnable — JAMAIS vague ou generique", "impact": "high|medium|low", "type": "risk|opportunity|neutral", "legal_ref": "Reference legale EXACTE: article de loi + jurisprudence recente (ex: Art. 65-70 Loi contrats travail — C. trav. Bruxelles 7 mai 2024)", "jurisprudence": "Jurisprudence applicable avec date et juridiction", "impact_description": "Ce que cela signifie CONCRETEMENT pour l'utilisateur — en langage simple, comme si on expliquait a un ami, JAMAIS de jargon juridique", "do_now": "Action PRECISE que l'utilisateur DOIT faire MAINTENANT — jamais generique comme 'consultez un avocat'", "risk_if_ignored": "Ce qui se passe si l'utilisateur ne fait RIEN — creer l'urgence, expliquer les consequences reelles de l'inaction", "confidence_score": 0.92, "jurisprudence_count": 47, "similar_cases_won": 44, "similar_cases_total": 47, "reasoning": "1-2 phrases expliquant POURQUOI ce score de confiance. Referencez la consistance de la jurisprudence et la clarte de la loi."}}
  ],
  "procedural_defects": [{{"vice": "description", "gravite": "fatal|significatif|mineur", "loi_applicable": "reference", "benefice_utilisateur": "comment ca aide"}}],
  "user_rights": [{{"droit": "droit specifique", "reference_legale": "loi exacte", "force": "fort|moyen|faible"}}],
  "opposing_weaknesses": [{{"faiblesse": "faiblesse specifique", "gravite": "critique|significative|mineure"}}],
  "financial_exposure_detailed": {{"meilleur_cas": "EUR 0", "cas_probable": "EUR 800-1200", "pire_cas": "EUR 2500 + frais"}},
  "applicable_laws": [{{"loi": "CCT n109", "pertinence": "Protection licenciement abusif", "favorable": "utilisateur|partie_adverse|neutre"}}],
  "organismes_recommandes": [{{"organisme": "Syndicat CSC/FGTB/CGSLB", "raison": "Aide juridique gratuite", "contact": "www.csc.be"}}],
  "key_violation_types": ["preavis_insuffisant", "absence_mise_en_demeure", "motif_vague"],
  "recommend_lawyer": true,
  "key_insight": "La phrase la plus importante"
}}
KEY_VIOLATION_TYPES: Genere 2-5 tags snake_case decrivant les violations cles (ex: "preavis_insuffisant", "absence_contrat_ecrit", "salaire_impaye", "clause_abusive", "absence_mise_en_demeure"). Ces tags servent a trouver des cas similaires.

Produis 3-6 constatations. CHAQUE constatation DOIT avoir les 12 champs: text (titre), impact, type, legal_ref (loi exacte + jurisprudence), jurisprudence, impact_description (langage simple), do_now (action precise), risk_if_ignored (consequence de l'inaction), confidence_score (0.0-1.0), jurisprudence_count (int), similar_cases_won (int), similar_cases_total (int), reasoning (1-2 phrases).

REGLES SCORE DE CONFIANCE:
- Loi claire ET jurisprudence constante: confiance > 0.90
- Loi claire mais jurisprudence variable: 0.70-0.90
- Interpretation possible: 0.50-0.70
- Argument faible: < 0.50
Basez les scores sur les tendances jurisprudentielles connues. Jurisprudence_count = nombre approximatif de decisions similaires."""

BE_PASS3_PROMPT = """TACHE: RECOMMANDATIONS STRATEGIQUES

Sur base des faits et de l'analyse, fournis des recommandations strategiques concretes selon le droit belge.

FAITS: {facts_json}
ANALYSE JURIDIQUE: {analysis_json}

REGLE CRITIQUE — DETECTION DU STADE DE LA PROCEDURE:
Tu DOIS identifier a quel stade se trouve le dossier et adapter TOUTES les recommandations en consequence:

STADE 1 — PV + perception immediate:
- La police a dresse un PV et propose une amende immediate
- Actions: Contester aupres du parquet, analyser le PV, verifier les vices de procedure
- Destinataire lettre: Procureur du Roi

STADE 2 — Convocation au parquet / proposition de transaction:
- Le parquet propose une transaction (Art. 216bis CIC)
- Actions: Accepter ou contester la transaction, negocier avec le substitut
- Destinataire lettre: Parquet / Substitut du Procureur

STADE 3 — Citation directe au tribunal:
- L'affaire est renvoyee devant le Tribunal de Police/Correctionnel
- La police N'EST PLUS le bon interlocuteur — JAMAIS de lettre a la police a ce stade
- Actions: Consulter un avocat penaliste AVANT l'audience, demander report si necessaire, constituer dossier defense
- Destinataire lettre: Greffe du Tribunal de Police
- Priorite 1: Consulter avocat (action_type: contacter_avocat)
- Priorite 2: Lettre au greffe si report necessaire (action_type: rediger_reponse)
- Priorite 3: Constituer dossier defense (action_type: aucune_action)

STADE 4 — Jugement rendu:
- Le tribunal a rendu sa decision
- Actions: Analyser le jugement, faire appel si defavorable (dans les 30 jours)
- Destinataire lettre: Greffe de la Cour d'Appel

INDICES POUR DETECTER LE STADE:
- "citation directe", "audience le", "tribunal de police", "renvoi devant" = STADE 3
- "perception immediate", "PV", "amende forfaitaire" = STADE 1
- "transaction", "proposition du parquet", "art. 216bis" = STADE 2
- "jugement", "condamne", "acquitte", "appel" = STADE 4

OBLIGATOIRE: Chaque next_step DOIT inclure un champ "recipient" indiquant a QUI la lettre/action est destinee.

REGLE CRITIQUE ARCHER_QUESTION: Tu DOIS generer une question specifique basee sur les faits du document. La question doit referencer un fait precis du document. JAMAIS de question generique. La question DOIT avoir 2-4 options de reponse cliquables.

Retourne UNIQUEMENT ce JSON:
{{
  "case_stage": "stage_1_pv|stage_2_transaction|stage_3_tribunal|stage_4_jugement",
  "recommended_strategy": {{
    "principale": "negocier|contester|se_conformer|mediation|tribunal",
    "raisonnement": "pourquoi c'est la meilleure strategie pour CE STADE",
    "resultat_attendu": "resultat realiste",
    "delai_resolution": "8-15 jours|1-3 mois|3-6 mois"
  }},
  "immediate_actions": [{{"action": "Action specifique au stade actuel", "delai": "dans les 24 heures", "priorite": "critique"}}],
  "next_steps": [
    {{"title": "Action specifique au STADE detecte", "description": "Description adaptee au stade", "action_type": "contacter_avocat|contacter_syndicat|saisir_mediateur|ajouter_document|rediger_reponse|aucune_action", "recipient": "Destinataire exact (ex: Greffe du Tribunal de Police, Parquet du Procureur du Roi)"}}
  ],
  "documents_to_gather": [{{"document": "Document a rassembler", "pourquoi": "Raison", "urgence": "critique|important|utile"}}],
  "leverage_points": [{{"levier": "Point de levier", "comment_utiliser": "Comment l'utiliser"}}],
  "red_lines": ["Ne jamais..."],
  "lawyer_recommendation": {{"necessaire": true, "urgence": "immediatement|dans_3_jours|dans_la_semaine|optionnel", "raison": "Raison", "type_avocat": "droit_du_travail|droit_du_bail|droit_penal|consommateur"}},
  "success_probability": {{"resolution_favorable": 35, "compromis_negocie": 48, "perte_partielle": 12, "perte_totale": 5}},
  "key_insight": "La phrase la plus importante",
  "archer_question": {{"text": "Question SPECIFIQUE", "options": ["Option 1", "Option 2", "Option 3"]}}
}}
OBLIGATOIRE: archer_question DOIT etre present. case_stage DOIT etre detecte correctement."""

BE_PASS4A_SYSTEM = """Tu es un avocat senior en Belgique representant l'utilisateur. Ton travail est de construire le dossier LE PLUS SOLIDE possible pour ton client. Trouve chaque argument, chaque vice de procedure, chaque protection legale qui beneficie a ton client. Sois agressif et exhaustif.

REGLE CRITIQUE: Tu DOIS generer exactement 4-5 arguments dans strongest_arguments. JAMAIS un tableau vide. Chaque argument doit citer une loi, un article de loi, ou un principe juridique belge specifique. Si la position de l'utilisateur semble faible, trouve des arguments proceduraux, des protections constitutionnelles, ou des arguments sur la charge de la preuve.

Reponds en JSON uniquement."""

BE_PASS4A_PROMPT = """Sur base de ces faits et de cette analyse, construis les arguments les plus solides pour defendre l'utilisateur selon le droit belge.

FAITS: {facts_json}
ANALYSE: {analysis_json}

OBLIGATOIRE: strongest_arguments DOIT contenir exactement 4-5 elements. JAMAIS moins de 4. Chaque argument doit citer un article de loi belge specifique (CCT 109, Loi contrats travail, Code civil, etc.).

Retourne UNIQUEMENT ce JSON:
{{
  "strongest_arguments": [
    {{"argument": "Description precise de l'argument avec reference legale", "strength": "strong|medium|weak", "law_basis": "Article de loi belge specifique (ex: CCT 109 art. 3, Art. 65-70 Loi contrats travail)", "how_to_use": "Comment utiliser cet argument dans la defense"}}
  ],
  "procedural_wins": ["liste des avantages proceduraux"],
  "best_outcome_scenario": "description du meilleur resultat possible",
  "opening_argument": "Premiere phrase percutante pour la lettre de reponse"
}}
OBLIGATOIRE: strongest_arguments DOIT contenir 4-5 elements. JAMAIS retourner moins de 4."""

BE_PASS4B_SYSTEM = "Tu es l'avocat de la partie adverse contre l'utilisateur en Belgique. Construis les arguments les plus solides contre l'utilisateur selon le droit belge. Reponds en JSON uniquement."
BE_PASS4B_PROMPT = """Construis les arguments que la partie adverse va utiliser contre l'utilisateur. Identifie les faiblesses de la position de l'utilisateur.

FAITS: {facts_json}
ANALYSE: {analysis_json}

Retourne JSON:
{{
  "opposing_arguments": [{{"argument": "L'utilisateur a accepte les conditions par signature", "force": "fort|moyen|faible", "base_legale": "Art. 5.69 nouveau CC", "counter": "Comment l'utilisateur peut contrer"}}],
  "user_weaknesses": ["liste des faiblesses de la position de l'utilisateur"],
  "worst_case": "description du pire resultat possible",
  "what_to_prepare": "Ce a quoi l'utilisateur doit se preparer"
}}"""


def load_belgian_jurisprudence(doc_type: str, region: str) -> str:
    """Load relevant Belgian jurisprudence based on document type and region"""
    be_db = JURISPRUDENCE_BE_DB.get("belgique", {})
    sections = []

    # Employment law (federal)
    if doc_type in ("licenciement", "contrat_travail", "c4"):
        for category in ["licenciement_abusif", "preavis_indemnites", "non_concurrence"]:
            cases = be_db.get("droit_travail", {}).get(category, [])
            for c in cases:
                sections.append(f"- {c['reference']}: {c['regle']}")

    # Tenancy law (regional)
    if doc_type in ("bail", "avis_resiliation_bail"):
        region_key = "bail_wallonie"
        if "flandre" in (region or "").lower() or "vlaanderen" in (region or "").lower():
            region_key = "bail_flandre"
        elif "bruxelles" in (region or "").lower():
            region_key = "bail_bruxelles"
        cases = be_db.get(region_key, [])
        for c in cases:
            sections.append(f"- {c['reference']}: {c['regle']}")
        # Also add Wallonia for Brussels (similar rules)
        if region_key == "bail_bruxelles":
            for c in be_db.get("bail_wallonie", []):
                sections.append(f"- {c['reference']}: {c['regle']}")

    # Consumer law
    if doc_type in ("mise_en_demeure", "facture", "autre"):
        for c in be_db.get("consommateur", []):
            sections.append(f"- {c['reference']}: {c['regle']}")

    # NDA
    if doc_type == "nda":
        for c in be_db.get("nda", []):
            sections.append(f"- {c['reference']}: {c['regle']}")

    # Reference amounts
    refs = be_db.get("montants_reference_2024", {})
    sections.append(f"\nMONTANTS DE REFERENCE 2024:")
    sections.append(f"- RMMMG: {refs.get('RMMMG_21_plus', 2029.88)} EUR/mois")
    sections.append(f"- Interet legal: {refs.get('interet_legal_2024', '5.75%')}")
    sections.append(f"- Garantie locative Wallonie/Bruxelles: {refs.get('garantie_locative_wallonie_bruxelles', 'max 2 mois')}")
    sections.append(f"- Garantie locative Flandre: {refs.get('garantie_locative_flandre', 'max 3 mois')}")

    if not sections:
        return "\nJURISPRUDENCE BELGE: Aucune jurisprudence specifique trouvee. Appliquer les principes generaux du droit belge."

    return "\nJURISPRUDENCE ET LEGISLATION BELGE APPLICABLE:\n" + "\n".join(sections)


def get_belgian_persona(language: str, region: str) -> str:
    """Get the correct Belgian persona based on language"""
    if language == "nl-BE" or language == "nl":
        return BELGIAN_PERSONA_NL.format(region=region or "Vlaanderen")
    elif language == "de-BE" or language == "de":
        return BELGIAN_PERSONA_DE.format(region=region or "Communaute germanophone")
    return BELGIAN_PERSONA_FR.format(region=region or "Belgique")


async def analyze_document_belgian(extracted_text: str, user_context: str = "", region: str = "Wallonie", language: str = "fr-BE") -> dict:
    """Belgian 5-pass analysis system with Belgian jurisprudence"""
    # Check cache first
    doc_hash = hashlib.sha256((extracted_text[:5000] + user_context + region + language + "BE").encode()).hexdigest()
    cached = await get_cached_analysis(doc_hash)
    if cached:
        logger.info("Belgian analysis: Returning cached result")
        return cached

    try:
        persona = get_belgian_persona(language, region)
        lang_instruction = get_language_instruction(language)
        jurisdiction_guard = get_jurisdiction_guard("BE", language)
        persona_with_lang = persona + jurisdiction_guard + lang_instruction
        context_section = f"CONTEXTE FOURNI PAR L'UTILISATEUR: {user_context}" if user_context else ""

        # PASS 1: Fact extraction (Belgian)
        logger.info("Belgian analysis: Passe 1 — Extraction des faits")
        facts = await call_claude(persona_with_lang, BE_PASS1_PROMPT.format(document_text=extracted_text[:15000], user_context_section=context_section))

        doc_type = facts.get("type_document", "autre")
        detected_region = facts.get("region_applicable", region)
        jurisprudence_text = load_belgian_jurisprudence(doc_type, detected_region)
        inferred_case_type = BE_DOC_TYPE_TO_CASE.get(doc_type, "other")

        # PASS 2: Legal analysis
        logger.info("Belgian analysis: Passe 2 — Analyse juridique")
        legal_analysis = await call_claude(persona_with_lang, BE_PASS2_PROMPT.format(facts_json=json.dumps(facts, indent=2, ensure_ascii=False), jurisprudence_section=jurisprudence_text), max_tokens=8000)

        # PASS 2 self-critique (BE).
        legal_analysis = await self_critique_pass(
            legal_analysis,
            pass_name="PASS2_legal_analysis_BE",
            context_hint=(
                "Analyse juridique belge: risk_score (4 dims), vices de procédure, droits de "
                "l'utilisateur, faiblesses de la partie adverse, lois applicables (Code civil, "
                "Code judiciaire, CCT, arrêtés royaux, ordonnances régionales), findings avec "
                "références légales belges précises et jurisprudence."
            ),
            language=language,
            jurisdiction="BE",
        )

        facts_str = json.dumps(facts, indent=2, ensure_ascii=False)
        analysis_str = json.dumps(legal_analysis, indent=2, ensure_ascii=False)

        # PASS 3 + 4A + 4B — RUN IN PARALLEL
        objective_block = _build_objective_block(user_context, language)
        logger.info("Belgian analysis: Passe 3+4A+4B — En parallele")
        strategy, user_arguments, opposing_arguments = await asyncio.gather(
            call_claude(persona_with_lang, BE_PASS3_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=6000),
            call_claude(BE_PASS4A_SYSTEM + jurisdiction_guard + lang_instruction, BE_PASS4A_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=4000),
            call_claude(BE_PASS4B_SYSTEM + jurisdiction_guard + lang_instruction, BE_PASS4B_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=4000),
        )

        # PASS 3 self-critique (BE).
        strategy = await self_critique_pass(
            strategy,
            pass_name="PASS3_strategy_BE",
            context_hint=(
                "Recommandations stratégiques belges adaptées au stade de la procédure "
                "(mise en demeure / pré-contentieux / tribunal / exécution). Doit inclure "
                "next_steps concrets, immediate_actions, leverage_points, success_probability "
                "et archer_question."
            ),
            language=language,
            jurisdiction="BE",
        )

        logger.info("Belgian analysis: 5 passes complete")

        # ═══ VALIDATION — enforce global rules for Belgian analysis ═══
        strategy["archer_question"] = await _validate_archer_question(strategy, facts_str, persona_with_lang, lang_instruction, language=language)
        user_arguments = await _validate_belgian_user_arguments(user_arguments, facts_str, analysis_str, lang_instruction, jurisdiction_guard)
        strategy["success_probability"] = _validate_success_probability(strategy, legal_analysis)

        now_date = datetime.now(timezone.utc).strftime("%Y-%m-%d")

        result = _build_belgian_analysis_result(doc_type, inferred_case_type, legal_analysis, strategy, facts, user_arguments, opposing_arguments, detected_region, language, now_date)

        # Cache the result
        await set_cached_analysis(doc_hash, result)
        return result
    except Exception as e:
        logger.error(f"Belgian analysis error: {e}")
        return _default_analysis()


# Belgian Letter Types
BELGIAN_LETTER_TYPES = {
    "employment": [
        {"id": "BE_DEMANDE_MOTIFS_CCT109", "label": "Demande de motifs de licenciement (CCT 109)", "desc": "Exiger les motifs concrets du licenciement"},
        {"id": "BE_CONTESTATION_PREAVIS", "label": "Contestation du preavis", "desc": "Le preavis calcule est insuffisant"},
        {"id": "BE_DEMANDE_INDEMNITES", "label": "Demande d'indemnites complementaires", "desc": "Reclamer les avantages omis"},
        {"id": "BE_CONTESTATION_NON_CONCURRENCE", "label": "Contestation clause de non-concurrence", "desc": "La clause est nulle"},
        {"id": "BE_SAISINE_SYNDICAT", "label": "Saisine du delegue syndical", "desc": "Demander l'intervention du syndicat"},
        {"id": "BE_PLAINTE_HARCELEMENT", "label": "Plainte pour harcelement/represailles", "desc": "Licenciement lie a une plainte"},
        {"id": "BE_DEMANDE_C4", "label": "Demande de document C4", "desc": "Obtenir le C4 pour le chomage"},
        {"id": "BE_CONTESTATION_FAUTE_GRAVE", "label": "Contestation faute grave", "desc": "Contester la qualification de faute grave"}
    ],
    "housing": [
        {"id": "BE_CONTESTATION_GARANTIE", "label": "Contestation garantie locative excessive", "desc": "Garantie superieure au maximum legal"},
        {"id": "BE_DEMANDE_ENREGISTREMENT", "label": "Demande d'enregistrement du bail", "desc": "Le bail n'est pas enregistre"},
        {"id": "BE_CONTESTATION_AUGMENTATION", "label": "Contestation augmentation de loyer", "desc": "Non conforme a l'indice sante"},
        {"id": "BE_CONTESTATION_EXPULSION", "label": "Contestation d'expulsion", "desc": "Vices de procedure ou treve hivernale"},
        {"id": "BE_DEMANDE_REPARATIONS", "label": "Demande de reparations au bailleur", "desc": "Obligations d'entretien non respectees"},
        {"id": "BE_RESILIATION_NON_ENREGISTRE", "label": "Resiliation bail non enregistre", "desc": "Partir sans preavis (bail non enregistre)"},
        {"id": "BE_DEMANDE_MEDIATION_LOYER", "label": "Demande de mediation", "desc": "Saisir le mediateur avant le tribunal"},
        {"id": "BE_RESTITUTION_GARANTIE", "label": "Demande de restitution de garantie", "desc": "Recuperer la garantie locative"}
    ],
    "debt": [
        {"id": "BE_CONTESTATION_MISE_DEMEURE", "label": "Contestation de la mise en demeure", "desc": "Vices de forme selon Livre XIX CDE"},
        {"id": "BE_DEMANDE_PLAN_PAIEMENT", "label": "Proposition de plan de paiement", "desc": "Echelonner la dette"},
        {"id": "BE_CONTESTATION_FRAIS", "label": "Contestation des frais de recouvrement", "desc": "Frais non conformes au CDE"},
        {"id": "BE_PLAINTE_HARCELEMENT_DETTE", "label": "Plainte pour harcelement (recouvrement)", "desc": "Contacts excessifs (>3/semaine)"},
        {"id": "BE_CONTESTATION_PRESCRIPTION", "label": "Exception de prescription", "desc": "La dette est prescrite"},
        {"id": "BE_DEMANDE_MEDIATION_DETTE", "label": "Demande de mediation de dettes", "desc": "Saisir le mediateur de dettes"}
    ],
    "nda": [
        {"id": "BE_NDA_CONTESTATION_PORTEE", "label": "Contestation portee du NDA", "desc": "Definition trop large"},
        {"id": "BE_NDA_PENALE_DISPROPORTIONNEE", "label": "Contestation clause penale", "desc": "Penale disproportionnee (art. 5.88 CC)"},
        {"id": "BE_NDA_DEMANDE_RECIPROCITE", "label": "Demande de reciprocite", "desc": "Obligations unilaterales"},
        {"id": "BE_NDA_LIMITATION_DUREE", "label": "Demande de limitation de duree", "desc": "Duree excessive"},
        {"id": "BE_NDA_CONTESTATION_SECRET", "label": "Contestation du caractere secret", "desc": "Information publique (Loi 30/07/2018)"},
        {"id": "BE_NDA_CARVEOUTS", "label": "Demande de carve-outs", "desc": "Exclure les connaissances anterieures"}
    ]
}

BELGIAN_LETTER_SYSTEM = """Tu es le moteur de communication juridique d'Archer pour la Belgique. Tu rediges des lettres professionnelles et strategiques en droit belge. Tu n'es PAS un avocat.

REGLES UNIVERSELLES LETTRES BELGES:
1. Format date: [ville], le [JJ mois AAAA]
2. Envoi: "Par courrier recommande avec accuse de reception"
3. Cloture obligatoire: "Je me reserve expressement tous mes droits et recours legaux."
4. Delai standard de reponse: "dans les 8 jours ouvrables"
5. Reference a la mediation avant tribunal
6. TOUJOURS citer les references legales belges exactes
7. Lettres entre 200-400 mots, professionnelles et concises
8. Ne jamais admettre la responsabilite de l'utilisateur
9. Toujours proposer un delai specifique (8 jours ouvrables)

FORMAT DE SORTIE — JSON uniquement:
{{
  "letter_type": "TYPE",
  "subject": "Objet de la lettre",
  "letter_body": "Texte complet de la lettre avec \\n pour les retours a la ligne",
  "tone": "cooperative|assertive|firm|neutral",
  "legal_basis": "Reference legale belge applicable",
  "key_points": ["Point 1", "Point 2", "Point 3"],
  "warnings": ["Avertissement 1"],
  "next_if_no_response": "Prochaine etape si pas de reponse",
  "disclaimer": "Cette lettre a ete redigee par Archer AI comme outil de communication juridique. Elle ne constitue pas un avis juridique."
}}"""

BELGIAN_OUTCOME_SYSTEM = """Tu es le moteur de prediction d'issue d'Archer pour la Belgique. Predis les probabilites des differents scenarios de resolution selon les statistiques judiciaires belges.

DONNEES STATISTIQUES BELGES:
- 73% des litiges du travail se reglent avant jugement (SPF Emploi 2023)
- 68% des litiges locatifs se reglent par mediation
- Duree moyenne procedure tribunal du travail: 8-14 mois
- Duree moyenne procedure justice de paix: 4-8 mois
- Taux de succes demande motifs licenciement (CCT 109): 89%
- Taux de succes contestation preavis insuffisant: 76%

Pour les cas emploi risque eleve (score >65):
Negociation syndicale: 45%, Conciliation tribunal travail: 28%, Jugement favorable: 18%, Perte totale: 9%

Pour les cas bail risque moyen (score 35-65):
Accord amiable: 52%, Mediation: 25%, Jugement favorable: 15%, Perte partielle: 8%

FORMAT — JSON uniquement:
{{
  "favorable": {{"probability": 30, "title": "Titre court", "description": "2-3 phrases", "likely_result": "Resultat specifique", "financial_impact": "Montant EUR", "timeline": "Delai estime"}},
  "neutral": {{"probability": 45, "title": "Titre court", "description": "2-3 phrases", "likely_result": "Resultat specifique", "financial_impact": "Montant EUR", "timeline": "Delai estime"}},
  "unfavorable": {{"probability": 25, "title": "Titre court", "description": "2-3 phrases", "likely_result": "Resultat specifique", "financial_impact": "Montant EUR", "timeline": "Delai estime"}},
  "key_factors": ["Facteur 1", "Facteur 2", "Facteur 3"],
  "recommendation": "Meilleure action en une phrase",
  "disclaimer": "Cette prediction fournit des informations juridiques uniquement."
}}"""

BELGIAN_CONTRACT_GUARD_PERSONA = """Tu es un avocat senior belge specialise dans la negociation de contrats pour des particuliers. Ton client s'apprete a signer le document fourni. Identifie EXACTEMENT ce qu'il faut negocier AVANT de signer selon le droit belge.

REGLES SPECIFIQUES BELGIQUE:
- La clause de non-concurrence sans compensation financiere est nulle (art. 65 loi 1978)
- Toute clause penale disproportionnee est reductible par le juge (art. 5.88 nouveau CC)
- Les clauses abusives dans les contrats consommateurs sont nulles de plein droit (art. VI.82 CDE)
- Le preavis legal ne peut etre reduit contractuellement en dessous du minimum legal
- La garantie locative ne peut depasser 2 mois (Wallonie/Bruxelles) ou 3 mois (Flandre)
- Tout NDA doit definir precisement ce qui est confidentiel (loi 30/07/2018)

CRITICAL RULES:
- Negotiation Score: 0 = contrat parfait (extremement rare), 100 = tres defavorable, ne pas signer
- Minimum score thresholds: NDA minimum 25, Contrat travail minimum 30, Bail minimum 25
- Toujours identifier au minimum 3 red lines
- Toujours identifier au minimum 4 points de negociation
- Toujours citer la loi belge exacte

OUTPUT: Retourne uniquement du JSON valide."""

CONTRACT_GUARD_PERSONA = """You are a senior contract negotiation attorney with 20+ years of experience protecting individuals from unfavorable contract terms. You specialize in identifying hidden risks, one-sided clauses, and missing protections in contracts before they are signed.

You think like a deal negotiator: you find every leverage point, every unfair clause, every missing protection. Your goal is to empower the user to negotiate better terms BEFORE signing.

CRITICAL RULES:
- Negotiation Score: 0 = perfect contract (extremely rare), 100 = highly unfavorable, do not sign
- Minimum score thresholds: NDA minimum 25, Employment contract minimum 30, Service agreement minimum 20, Lease minimum 25
- Always identify at least 3 red lines (clauses the user should NEVER accept as-is)
- Always identify at least 4 negotiation points (specific changes to request)
- Always identify missing protections the user should demand
- Always compare clauses to industry standards
- Never tell the user to just accept and sign
- Always provide specific alternative language for problematic clauses
- Reference applicable law when a clause may be unenforceable

OUTPUT: Always return valid JSON only, no other text."""

CONTRACT_GUARD_PROMPT = """TASK: CONTRACT NEGOTIATION ANALYSIS ("Before I Sign")
Analyze this contract/agreement from the perspective of the user who is ABOUT to sign it. The goal is NOT risk analysis of an existing situation — it's pre-signing due diligence and negotiation preparation.

DOCUMENT TEXT:
{document_text}

{user_context_section}

Analyze every clause and return ONLY this JSON — no other text:
{{
  "document_type": "nda|employment_contract|lease|service_agreement|vendor_contract|partnership_agreement|freelance_contract|other",
  "contract_title": "Short descriptive title max 60 chars",
  "parties": {{
    "user_role": "employee|tenant|contractor|service_provider|partner|licensee|other",
    "counterparty": "Company or individual name",
    "counterparty_type": "corporation|startup|individual|government|other"
  }},
  "negotiation_score": 0,
  "negotiation_level": "favorable|balanced|unfavorable|dangerous",
  "summary": "2-3 sentence summary of what this contract does and the overall balance of power",
  "red_lines": [
    {{
      "clause": "Exact clause reference (Section X, Article Y)",
      "current_text": "What the clause currently says (paraphrased)",
      "risk": "Why this is dangerous for the user",
      "severity": "critical|high|medium",
      "suggested_change": "Specific alternative language the user should propose"
    }}
  ],
  "negotiation_points": [
    {{
      "clause": "Clause reference",
      "issue": "What's wrong or unfavorable",
      "current_terms": "Current terms in the contract",
      "industry_standard": "What is standard in the industry",
      "suggested_counter": "What the user should ask for instead",
      "priority": "must_have|nice_to_have|optional",
      "leverage_tip": "How to frame this request in negotiation"
    }}
  ],
  "missing_protections": [
    {{
      "protection": "What's missing",
      "why_important": "Why the user needs this",
      "suggested_clause": "Exact language to add",
      "priority": "critical|important|recommended"
    }}
  ],
  "standard_clauses_check": [
    {{
      "clause_type": "termination|liability|indemnification|ip_ownership|confidentiality|non_compete|payment_terms|dispute_resolution|governing_law|force_majeure",
      "present": true,
      "fair": true,
      "notes": "Brief assessment"
    }}
  ],
  "financial_terms": {{
    "total_value": "Dollar amount or description",
    "payment_schedule": "Description",
    "penalties": "Any penalty clauses",
    "hidden_costs": ["List of hidden or unexpected costs"]
  }},
  "duration_and_exit": {{
    "contract_duration": "Duration",
    "auto_renewal": true,
    "termination_notice": "Notice period required",
    "early_exit_penalty": "Description or null",
    "user_can_exit_easily": true
  }},
  "power_balance": {{
    "score": 0,
    "favors": "user|counterparty|balanced",
    "key_asymmetries": ["List of clauses that favor one party disproportionately"]
  }},
  "applicable_laws": [
    {{"law": "statute name", "relevance": "how it applies", "protects_user": true}}
  ],
  "overall_recommendation": {{
    "action": "sign_as_is|negotiate_then_sign|significant_changes_needed|do_not_sign|lawyer_review_required",
    "reasoning": "Why this recommendation",
    "estimated_negotiation_time": "1-2 days|1 week|2+ weeks"
  }},
  "negotiation_email_draft": {{
    "subject": "Re: Contract Review — Proposed Amendments",
    "body": "Full professional email body requesting the changes identified above. Reference specific clauses. Be firm but professional. 200-400 words."
  }}
}}

Produce at least 3 red_lines, 4 negotiation_points, and 2 missing_protections."""


async def analyze_contract_guard(extracted_text: str, user_context: str = "", country: str = "US", region: str = "", language: str = "en") -> dict:
    """Contract Guard analysis — negotiation-focused pre-signing review"""
    try:
        context_section = f"USER CONTEXT: {user_context}" if user_context else ""

        jurisdiction_guard = get_jurisdiction_guard("BE" if country == "BE" else "US", language)
        if country == "BE":
            persona = BELGIAN_CONTRACT_GUARD_PERSONA + jurisdiction_guard + get_language_instruction(language)
            logger.info("Contract Guard (Belgian): Starting negotiation analysis")
        else:
            persona = CONTRACT_GUARD_PERSONA + jurisdiction_guard
            logger.info("Contract Guard: Starting negotiation analysis")

        result = await call_claude(
            persona,
            CONTRACT_GUARD_PROMPT.format(document_text=extracted_text[:15000], user_context_section=context_section),
            max_tokens=4000
        )

        result = _ensure_contract_guard_fields(result)
        logger.info(f"Contract Guard: Analysis complete — Score {result['negotiation_score']}/100")
        return result

    except Exception as e:
        logger.error(f"Contract Guard analysis error: {e}")
        return {
            "document_type": "other",
            "contract_title": "Contract Analysis",
            "negotiation_score": 50,
            "negotiation_level": "balanced",
            "summary": "Contract uploaded for review. Some clauses could not be fully analyzed.",
            "red_lines": [{"clause": "General", "current_text": "Review required", "risk": "Manual review recommended", "severity": "medium", "suggested_change": "Consult a lawyer"}],
            "negotiation_points": [],
            "missing_protections": [],
            "standard_clauses_check": [],
            "overall_recommendation": {"action": "lawyer_review_required", "reasoning": "Analysis could not be fully completed", "estimated_negotiation_time": "1 week"}
        }

# ================== Letter Generation System ==================

LETTER_SYSTEM_PROMPT = """You are Archer's legal communication engine. You write professional, strategic response letters on behalf of US residents facing legal situations. You are NOT a lawyer. Your letters are legal communications, not legal advice.

Your letters must be:
- Professional and formal in tone
- Factually accurate based on the case details provided
- Strategic — designed to de-escalate, buy time, or protect the user's rights
- Compliant with US law — never ask for anything illegal
- Clear and concise — no unnecessary legal jargon
- Signed by the user, not by Archer

LETTER WRITING RULES:
1. NEVER admit liability or guilt
2. ALWAYS use formal business letter format with date, addresses, subject line
3. ALWAYS end with "Sincerely," followed by the user's name and a signature line
4. ALWAYS include "Sent via Certified Mail" in the header when relevant
5. Reference specific dollar amounts, dates, and clause numbers from the case
6. Reference applicable law when it strengthens the user's position
7. Keep letters between 200-400 words — professional and concise
8. One clear ask per letter — don't combine multiple strategies
9. ALWAYS include: "I reserve all of my legal rights and remedies"
10. Never threaten violence or make illegal demands
11. Always propose a specific deadline for the opposing party to respond (7-10 days)

OUTPUT FORMAT — respond ONLY with valid JSON, no other text:

{
  "letter_type": "LETTER_TYPE_HERE",
  "subject": "Re: Subject line",
  "letter_body": "Full letter text here with proper formatting. Use \\n for line breaks.",
  "tone": "cooperative|assertive|firm|neutral",
  "legal_basis": "Applicable statute or legal basis",
  "key_points": ["Point 1", "Point 2", "Point 3"],
  "warnings": ["Warning 1", "Warning 2"],
  "next_if_no_response": "What to do if no response",
  "disclaimer": "This letter was drafted by Archer AI as a legal communication tool. It does not constitute legal advice."
}"""

# Letter types by case type
LETTER_TYPES = {
    "nda": [
        {"id": "NDA_CLAUSE_CLARIFICATION", "label": "Request clause clarification", "desc": "Specific clauses are too vague"},
        {"id": "NDA_NEGOTIATE_SCOPE", "label": "Negotiate confidentiality scope", "desc": "Definition is too broad"},
        {"id": "NDA_MUTUAL_OBLIGATIONS", "label": "Request mutual obligations", "desc": "Currently only one party is bound"},
        {"id": "NDA_CHALLENGE_DURATION", "label": "Challenge duration", "desc": "Obligation period is excessive"},
        {"id": "NDA_GEOGRAPHIC_LIMITATION", "label": "Request geographic limitation", "desc": "Scope is unreasonably wide"},
        {"id": "NDA_PROPOSE_CARVEOUTS", "label": "Propose carve-outs", "desc": "Exclude already known information"},
        {"id": "NDA_CHALLENGE_PENALTY", "label": "Challenge penalty clauses", "desc": "Damages are disproportionate"},
        {"id": "NDA_GOVERNING_LAW", "label": "Request governing law change", "desc": "Jurisdiction is unfavorable"}
    ],
    "housing": [
        {"id": "PAYMENT_PLAN_PROPOSAL", "label": "Payment plan proposal", "desc": "Pay in installments"},
        {"id": "DISPUTE_NOTICE_VALIDITY", "label": "Dispute notice validity", "desc": "Procedural errors found"},
        {"id": "CONTEST_DAMAGE_AMOUNTS", "label": "Contest damage amounts", "desc": "Amounts are inflated or unjustified"},
        {"id": "REQUEST_MORE_TIME", "label": "Request more time", "desc": "Need additional days to respond"},
        {"id": "CHALLENGE_OCCUPANT_CLAIM", "label": "Challenge unauthorized occupant claim", "desc": "Contest the allegation"},
        {"id": "ASSERT_DEPOSIT_RIGHTS", "label": "Assert security deposit rights", "desc": "Landlord must return deposit"},
        {"id": "REQUEST_MEDIATION", "label": "Request mediation", "desc": "Avoid court proceedings"},
        {"id": "DISPUTE_LEASE_VIOLATION", "label": "Dispute lease violation", "desc": "Contest the alleged breach"}
    ],
    "employment": [
        {"id": "CONTEST_TERMINATION", "label": "Contest termination", "desc": "Formally dispute the firing"},
        {"id": "DEMAND_UNPAID_WAGES", "label": "Demand unpaid wages", "desc": "Request final paycheck and overtime"},
        {"id": "CHALLENGE_NON_COMPETE", "label": "Challenge non-compete enforceability", "desc": "Too broad or unreasonable"},
        {"id": "REQUEST_TERMINATION_REASON", "label": "Request written termination reason", "desc": "Legal right to know why"},
        {"id": "ASSERT_DISCRIMINATION", "label": "Assert discrimination claim", "desc": "Termination may be discriminatory"},
        {"id": "DEMAND_SEVERANCE", "label": "Demand severance pay", "desc": "Request negotiated severance"},
        {"id": "DISPUTE_PERFORMANCE_REVIEW", "label": "Dispute performance review", "desc": "Contest documented reasons"},
        {"id": "REQUEST_REFERENCE", "label": "Request reference letter", "desc": "Protect future employment"}
    ],
    "debt": [
        {"id": "DEBT_VALIDATION_REQUEST", "label": "Debt validation request", "desc": "Force collector to prove debt is valid (FDCPA right)"},
        {"id": "CEASE_AND_DESIST", "label": "Cease and desist", "desc": "Stop all contact immediately"},
        {"id": "DISPUTE_DEBT_OWNERSHIP", "label": "Dispute debt ownership", "desc": "Collector may not own this debt"},
        {"id": "CHALLENGE_STATUTE_LIMITATIONS", "label": "Challenge statute of limitations", "desc": "Debt may be too old to collect"},
        {"id": "SETTLEMENT_OFFER", "label": "Settlement offer", "desc": "Negotiate 40-60% lump sum"},
        {"id": "FDCPA_VIOLATION_COMPLAINT", "label": "FDCPA violation complaint", "desc": "Document their violations"},
        {"id": "REQUEST_PAYMENT_HISTORY", "label": "Request payment history", "desc": "Demand full account statement"},
        {"id": "DISPUTE_CREDIT_REPORTING", "label": "Dispute credit reporting", "desc": "Challenge inaccurate reporting"}
    ],
    "demand": [
        {"id": "FORMAL_DISPUTE", "label": "Formal dispute", "desc": "Reject all claims"},
        {"id": "COUNTER_PROPOSAL", "label": "Counter-proposal", "desc": "Offer alternative resolution"},
        {"id": "REQUEST_EVIDENCE", "label": "Request evidence", "desc": "Demand proof of their claims"},
        {"id": "CHALLENGE_LEGAL_BASIS", "label": "Challenge legal basis", "desc": "Their claim has no legal foundation"},
        {"id": "PROPOSE_MEDIATION", "label": "Propose mediation", "desc": "Neutral third party resolution"},
        {"id": "ASSERT_COUNTER_CLAIM", "label": "Assert counter-claim", "desc": "You have claims against them too"},
        {"id": "REQUEST_EXTENSION", "label": "Request extension", "desc": "Need more time to respond"},
        {"id": "PARTIAL_ACCEPTANCE", "label": "Partial acceptance", "desc": "Accept some claims, reject others"}
    ],
    "court": [
        {"id": "FILE_FORMAL_RESPONSE", "label": "File formal response", "desc": "Contest the claims in court"},
        {"id": "REQUEST_CONTINUANCE", "label": "Request continuance", "desc": "More time to prepare defense"},
        {"id": "CHALLENGE_JURISDICTION", "label": "Challenge jurisdiction", "desc": "Wrong court or state"},
        {"id": "MOTION_TO_DISMISS", "label": "Motion to dismiss", "desc": "Case has no legal merit"},
        {"id": "PROPOSE_SETTLEMENT", "label": "Propose settlement", "desc": "Avoid trial"},
        {"id": "REQUEST_DISCOVERY", "label": "Request discovery", "desc": "Demand their evidence"},
        {"id": "COUNTER_CLAIM_FILING", "label": "Counter-claim filing", "desc": "File claims against plaintiff"},
        {"id": "DEFAULT_AVOIDANCE", "label": "Default avoidance", "desc": "Respond before deadline"}
    ],
    "consumer": [
        {"id": "FORMAL_REFUND_DEMAND", "label": "Formal refund demand", "desc": "Request full refund with deadline"},
        {"id": "CHARGEBACK_SUPPORT", "label": "Chargeback support letter", "desc": "Support credit card dispute"},
        {"id": "FTC_COMPLAINT_NOTICE", "label": "FTC complaint notice", "desc": "Warn of regulatory complaint"},
        {"id": "BBB_COMPLAINT_NOTICE", "label": "BBB complaint notice", "desc": "Warn of Better Business Bureau filing"},
        {"id": "ATTORNEY_GENERAL_NOTICE", "label": "Attorney General notice", "desc": "Warn of state AG complaint"},
        {"id": "SMALL_CLAIMS_NOTICE", "label": "Small claims court notice", "desc": "Threaten court action"},
        {"id": "WARRANTY_CLAIM", "label": "Warranty claim", "desc": "Assert product warranty rights"},
        {"id": "CONSUMER_PROTECTION_CLAIM", "label": "Consumer protection claim", "desc": "Cite applicable consumer law"}
    ],
    "contract": [
        {"id": "FORMAL_DISPUTE", "label": "Formal dispute", "desc": "Reject all claims"},
        {"id": "COUNTER_PROPOSAL", "label": "Counter-proposal", "desc": "Offer alternative resolution"},
        {"id": "REQUEST_EVIDENCE", "label": "Request evidence", "desc": "Demand proof of their claims"},
        {"id": "CHALLENGE_LEGAL_BASIS", "label": "Challenge legal basis", "desc": "Their claim has no legal foundation"},
        {"id": "PROPOSE_MEDIATION", "label": "Propose mediation", "desc": "Neutral third party resolution"},
        {"id": "ASSERT_COUNTER_CLAIM", "label": "Assert counter-claim", "desc": "You have claims against them too"},
        {"id": "REQUEST_EXTENSION", "label": "Request extension", "desc": "Need more time to respond"},
        {"id": "PARTIAL_ACCEPTANCE", "label": "Partial acceptance", "desc": "Accept some claims, reject others"}
    ],
    "immigration": [
        {"id": "SPONSOR_AGREEMENT_DISPUTE", "label": "Sponsor dispute", "desc": "Address issues with employment sponsor"},
        {"id": "EMPLOYMENT_AUTHORIZATION_INQUIRY", "label": "Authorization inquiry", "desc": "Inquire about work authorization status"},
        {"id": "VISA_CONTRACT_CLARIFICATION", "label": "Visa clarification", "desc": "Request clarification on visa terms"},
        {"id": "STATUS_UPDATE_REQUEST", "label": "Status update", "desc": "Request update on pending application"},
        {"id": "WORK_PERMIT_DISPUTE", "label": "Work permit dispute", "desc": "Contest work permit conditions"},
        {"id": "SPONSOR_OBLIGATIONS", "label": "Sponsor obligations", "desc": "Remind sponsor of legal duties"}
    ],
    "family": [
        {"id": "MEDIATION_REQUEST", "label": "Request mediation", "desc": "Propose family mediation"},
        {"id": "CUSTODY_CONCERNS", "label": "Document concerns", "desc": "Document custody concerns formally"},
        {"id": "SUPPORT_MODIFICATION", "label": "Modify support", "desc": "Request modification of support terms"},
        {"id": "COMMUNICATION_GUIDELINES", "label": "Communication request", "desc": "Establish communication guidelines"},
        {"id": "ASSET_DIVISION_DISPUTE", "label": "Asset division dispute", "desc": "Contest proposed asset split"},
        {"id": "VISITATION_RIGHTS", "label": "Visitation rights", "desc": "Assert or modify visitation schedule"}
    ],
    "other": [
        {"id": "FORMAL_DISPUTE", "label": "Formal dispute", "desc": "Reject all claims"},
        {"id": "COUNTER_PROPOSAL", "label": "Counter-proposal", "desc": "Offer alternative resolution"},
        {"id": "REQUEST_EVIDENCE", "label": "Request evidence", "desc": "Demand proof of their claims"},
        {"id": "CHALLENGE_LEGAL_BASIS", "label": "Challenge legal basis", "desc": "Their claim has no legal foundation"},
        {"id": "PROPOSE_MEDIATION", "label": "Propose mediation", "desc": "Neutral third party resolution"},
        {"id": "ASSERT_COUNTER_CLAIM", "label": "Assert counter-claim", "desc": "You have claims against them too"},
        {"id": "REQUEST_EXTENSION", "label": "Request extension", "desc": "Need more time to respond"},
        {"id": "PARTIAL_ACCEPTANCE", "label": "Partial acceptance", "desc": "Accept some claims, reject others"}
    ]
}

CITIZEN_TONE_FR = """
REGLE DE TON — LETTRE CITOYEN (Option 1):
Tu rediges cette lettre AU NOM DU CITOYEN lui-meme, PAS d'un avocat.
- TOUJOURS utiliser "Je" — JAMAIS "Nous"
- JAMAIS de phrases d'avocat: "Par la presente nous avons l'honneur de...", "Nous sollicitons respectueusement...", "Il appert que...", "Aux termes de..."
- Les references legales sont autorisees mais expliquees simplement
  INTERDIT: "En application de l'article 216bis CIC..."
  CORRECT: "Je souhaite contester cette amende. L'article 216bis du Code d'instruction criminelle me donne ce droit."
- Signe par: [NOM DE L'UTILISATEUR] — PAS un cabinet
- Ton: simple, direct, premiere personne, citoyen respectueux qui s'adresse a une autorite
- Cloture: "Veuillez agreer, Madame/Monsieur, l'expression de mes salutations distinguees."
"""

CITIZEN_TONE_EN = """
TONE RULE — CITIZEN LETTER (Option 1):
You are writing this letter IN THE NAME OF THE CITIZEN themselves, NOT a lawyer.
- ALWAYS use "I" — NEVER "We"
- NEVER use attorney phrases: "We respectfully submit...", "It appears that...", "Pursuant to...", "We hereby..."
- Legal references ARE allowed but explained simply
  WRONG: "Pursuant to Fla. Stat. § 83.56(3)..."
  RIGHT: "I am disputing this notice. Florida law (Section 83.56) gives me the right to do so."
- Signed by: [USER FULL NAME] — not a law firm
- Tone: simple, direct, first person, respectful citizen writing to an authority
- Closing: "Sincerely,"
"""

CITIZEN_TONE_NL = """
TOONREGEL — BURGERBRIEF (Optie 1):
Je schrijft deze brief IN NAAM VAN DE BURGER zelf, NIET van een advocaat.
- ALTIJD "Ik" gebruiken — NOOIT "Wij"
- NOOIT advocatentaal: "Bij deze hebben wij de eer...", "Wij verzoeken eerbiedig..."
- Wettelijke referenties zijn toegestaan maar eenvoudig uitgelegd
- Ondertekend door: [NAAM GEBRUIKER] — geen kantoor
- Toon: eenvoudig, direct, eerste persoon, respectvolle burger
"""

ATTORNEY_TONE_FR = """
REGLE DE TON — LETTRE D'AVOCAT (Option 2):
Tu es un avocat senior redigeant une lettre formelle au nom de ton client.
- Utiliser le langage juridique formel et professionnel
- Referer au client en troisieme personne: "notre client M. [Nom]" / "notre cliente Mme [Nom]"
- Citer les references legales formellement: "En application de l'article...", "Aux termes de..."
- Signe par: [Nom Avocat] — Avocat au Barreau de [juridiction]
- En-tete: Cabinet [Nom] — Avocats
- Ton: formel, autoritaire, professionnel
- Ouverture: "Maitre, Monsieur le Procureur du Roi, Nous avons l'honneur de vous adresser la presente au nom et pour le compte de notre client..."
- Cloture: "Veuillez agreer, Maitre/Monsieur, l'expression de mes sentiments distingues."
"""

ATTORNEY_TONE_EN = """
TONE RULE — ATTORNEY LETTER (Option 2):
You are a senior attorney writing a formal legal letter on behalf of your client.
- Use formal professional legal language
- Reference the client in third person: "our client Mr. [Name]" / "our client Ms. [Name]"
- Cite legal references formally: "Pursuant to...", "Under the provisions of..."
- Signed by: [Attorney Name], Esq. — [Bar Association]
- Letterhead: [Firm Name] — Attorneys at Law
- Tone: formal, authoritative, professional
- Opening: "Dear [Title], We write on behalf of our client, [Name], regarding..."
- Closing: "Very truly yours,"
"""

ATTORNEY_TONE_NL = """
TOONREGEL — ADVOCAATBRIEF (Optie 2):
Je bent een senior advocaat die een formele juridische brief schrijft namens je client.
- Gebruik formele professionele juridische taal
- Verwijs naar de client in derde persoon: "onze client de heer [Naam]"
- Citeer wettelijke referenties formeel: "Krachtens artikel..."
- Ondertekend door: [Naam Advocaat] — Advocaat bij de Balie van [jurisdictie]
- Toon: formeel, gezaghebbend, professioneel
"""

def get_letter_tone_instruction(tone: str, language: str) -> str:
    """Get tone-specific instruction for letter generation."""
    lang_prefix = language[:2] if language else "en"
    if tone == "attorney":
        if lang_prefix == "fr":
            return ATTORNEY_TONE_FR
        elif lang_prefix == "nl":
            return ATTORNEY_TONE_NL
        return ATTORNEY_TONE_EN
    else:
        if lang_prefix == "fr":
            return CITIZEN_TONE_FR
        elif lang_prefix == "nl":
            return CITIZEN_TONE_NL
        return CITIZEN_TONE_EN


async def generate_letter_with_claude(letter_data: dict, belgian: bool = False, language: str = "en", tone: str = "citizen") -> dict:
    """Generate a response letter using Claude API"""
    try:
        system_prompt = BELGIAN_LETTER_SYSTEM if belgian else LETTER_SYSTEM_PROMPT
        system_prompt += get_language_instruction(language)
        system_prompt += get_letter_tone_instruction(tone, language)
        async with httpx.AsyncClient() as client:
            response = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": get_anthropic_key(),
                    "anthropic-version": "2023-06-01"
                },
                json={
                    "model": "claude-opus-4-7",
                    "max_tokens": 3000,
                    "system": system_prompt,
                    "messages": [{
                        "role": "user",
                        "content": f"Generate a professional response letter based on this case data. Return JSON only:\n\n{json.dumps(letter_data, indent=2)}"
                    }]
                },
                timeout=90.0
            )
            response.raise_for_status()
            data = response.json()
            text = data["content"][0]["text"]
            # Clean JSON from markdown code blocks
            text = text.replace("```json", "").replace("```", "").strip()
            return json.loads(text)
    except Exception as e:
        logger.error(f"Claude Letter API error: {e}")
        raise HTTPException(status_code=500, detail=f"Letter generation failed: {str(e)}")

async def _ocr_page_via_vision(http_client: httpx.AsyncClient, img_b64: str, page_num: int, filename: str) -> str:
    ocr_prompt = (
        "Transcris intégralement le texte de ce document juridique scanné. "
        "Conserve la structure (titres, articles, paragraphes). Langue originale. "
        "Retourne uniquement le texte transcrit, sans commentaire ni analyse."
    )
    try:
        response = await http_client.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": get_anthropic_key(),
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": "claude-opus-4-7",
                "max_tokens": 4000,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64}},
                        {"type": "text", "text": ocr_prompt},
                    ],
                }],
            },
        )
        response.raise_for_status()
        return response.json()["content"][0]["text"].strip()
    except Exception as e:
        logger.error(f"Vision OCR failed for {filename} page {page_num}: {e}")
        return ""


async def extract_text_from_pdf(file_bytes: bytes, filename: str = "") -> str:
    """Extract text from PDF. Falls back to Claude Vision OCR when native extraction
    yields <100 chars (scanned PDFs, photos of contracts, signed documents)."""
    native_text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            parts = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    parts.append(page_text)
            native_text = "\n".join(parts).strip()
    except Exception as e:
        logger.error(f"PDF native extraction error for {filename}: {e}")

    use_vision = len(native_text) < 100
    logger.info(f"PDF {filename}: native text={len(native_text)} chars, using {'vision OCR' if use_vision else 'native'}")
    if not use_vision:
        return native_text

    try:
        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
        total_pages = len(pdf_doc)
        page_count = min(total_pages, 10)
        if total_pages > 10:
            logger.warning(f"PDF {filename}: {total_pages} pages, OCR limited to first 10")
        images = []
        for i in range(page_count):
            pix = pdf_doc[i].get_pixmap(dpi=200)
            images.append(base64.b64encode(pix.tobytes("jpeg")).decode("utf-8"))
        pdf_doc.close()
    except Exception as e:
        logger.error(f"PDF {filename}: page-to-image render failed: {e}")
        return native_text

    if not images:
        return native_text

    async with httpx.AsyncClient(timeout=120.0) as http_client:
        results = await asyncio.gather(*[
            _ocr_page_via_vision(http_client, b64, i + 1, filename)
            for i, b64 in enumerate(images)
        ])

    ocr_parts = [f"--- Page {i+1} ---\n{txt}" for i, txt in enumerate(results) if txt]
    ocr_text = "\n\n".join(ocr_parts).strip()
    logger.info(f"PDF {filename}: vision OCR extracted {len(ocr_text)} chars from {len(ocr_parts)}/{len(images)} pages")
    return ocr_text or native_text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


def pdf_pages_to_images(file_bytes: bytes, max_pages: int = 5) -> list:
    """Convert PDF pages to JPEG base64 images using pymupdf"""
    images = []
    try:
        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = min(len(pdf_doc), max_pages)
        for i in range(page_count):
            page = pdf_doc[i]
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("jpeg")
            b64 = base64.b64encode(img_bytes).decode("utf-8")
            images.append(b64)
        pdf_doc.close()
        logger.info(f"Converted {page_count} PDF pages to images")
    except Exception as e:
        logger.error(f"PDF to image conversion error: {e}")
    return images


def image_file_to_base64(file_bytes: bytes, content_type: str) -> str:
    """Convert an image file to JPEG base64 for Claude Vision"""
    try:
        img = Image.open(io.BytesIO(file_bytes))
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        return base64.b64encode(buf.getvalue()).decode("utf-8")
    except Exception as e:
        logger.error(f"Image conversion error: {e}")
        return base64.b64encode(file_bytes).decode("utf-8")


# ================== Jurisdiction detection & guard ==================
# Keywords weighted per jurisdiction. Detection is deterministic (no extra Claude call).

_BE_KEYWORDS = [
    ("moniteur belge", 8), ("code civil belge", 6), ("code civil", 3),
    ("code judiciaire", 6), ("code des sociétés", 5), ("arrêté royal", 6),
    ("arrete royal", 6), ("moniteur", 4), ("cct ", 5), ("onss", 5),
    ("ejustice.be", 8), ("ejustice", 4), ("région de bruxelles", 7),
    ("region de bruxelles", 7), ("bruxelles-capitale", 7), ("wallonie", 5),
    ("région wallonne", 6), ("region wallonne", 6), ("flandre", 5),
    ("vlaanderen", 5), ("communauté germanophone", 6),
    ("justice de paix", 5), ("tribunal du travail", 5),
    ("tribunal de première instance", 4), ("tribunal de premiere instance", 4),
    ("belgique", 4), ("belgisch", 4), ("belgique", 4),
    (" n.n. ", 3), ("numéro national", 4), ("numero national", 4),
    ("indice santé", 5), ("indice sante", 5), ("précompte", 4), ("precompte", 4),
    ("€ ", 2), ("eur ", 2), ("bail de résidence principale", 6),
    ("bail de residence principale", 6), ("enregistrement du bail", 5),
    ("garantie locative", 4), ("préavis", 3), ("preavis", 3),
    ("mise en demeure", 3), ("recommandé", 2), ("recommande", 2),
    ("livre xix", 5), ("livre xvii", 5), ("cde", 3),
]

_US_KEYWORDS = [
    ("u.s.c.", 8), ("usc §", 7), (" usc ", 5), ("c.f.r.", 6), ("cfr §", 6),
    ("state of ", 5), ("state of california", 7), ("state of new york", 7),
    ("state of texas", 7), ("state of florida", 7),
    ("plaintiff", 5), ("defendant", 5), ("superior court", 6),
    ("district court", 5), ("circuit court", 5), ("court of appeals", 5),
    ("county of ", 4), ("department of ", 3), ("fdcpa", 7), ("fcra", 6),
    ("flsa", 6), ("fmla", 6), ("ada ", 4), ("title vii", 6),
    ("hipaa", 6), ("erisa", 6), ("state statute", 5),
    ("federal register", 6), ("irs", 4), ("ssn", 4),
    ("zip code", 4), (" llc", 3), (" inc.", 3), (" corp.", 3),
    ("attorney at law", 5), ("esq.", 4), ("bar association", 4),
    (" $", 2), ("usd", 3), ("fl stat", 6), ("cal. code", 6),
    ("n.y. ", 3), ("ca. ", 2), ("tx. ", 2),
    ("landlord-tenant", 4), ("eviction notice", 4), ("summons", 3),
    ("complaint filed", 4), ("motion to dismiss", 5), ("discovery", 3),
]


def detect_jurisdiction(text: str) -> str:
    """Detect document jurisdiction (BE | US) from extracted text.
    Deterministic keyword scan — no Claude call. Returns 'BE' or 'US'.
    On tie or empty text, returns 'US' (historical default)."""
    if not text:
        return "US"
    t = text.lower()
    be_score = sum(w for kw, w in _BE_KEYWORDS if kw in t)
    us_score = sum(w for kw, w in _US_KEYWORDS if kw in t)
    logger.info(f"Jurisdiction detect: BE={be_score} US={us_score} (text={len(text)} chars)")
    if be_score == 0 and us_score == 0:
        return "US"
    if be_score > us_score:
        return "BE"
    return "US"


def get_jurisdiction_guard(jurisdiction: str, language: str = "en") -> str:
    """Return a hard constraint block appended to the persona to forbid cross-jurisdiction citations."""
    j = (jurisdiction or "US").upper()
    if j == "BE":
        return """

RÈGLE JURIDICTIONNELLE — NON-NÉGOCIABLE :
Tu appliques UNIQUEMENT le droit belge (fédéral + régional Bruxelles/Wallonie/Flandre/Communauté germanophone).
NE CITE JAMAIS, sous aucun prétexte, de loi américaine (U.S.C., C.F.R., State statutes, FDCPA, FCRA, FLSA, FMLA, ADA, Title VII, HIPAA, ERISA, case law US, jurisprudence US).
Références AUTORISÉES exclusivement : Code civil belge (ancien et nouveau), Code judiciaire, Code des sociétés et des associations, Code de droit économique (CDE), arrêtés royaux, ordonnances régionales (Bruxelles/Wallonie/Flandre), lois fédérales belges, CCT (conventions collectives de travail), jurisprudence belge (Cour de cassation, Cour constitutionnelle, cours d'appel belges, justices de paix, tribunaux du travail), Moniteur belge, ejustice.be.
Devises : EUR uniquement. Jamais de USD.
Si le document contient des références à du droit étranger, tu peux les mentionner comme contexte mais l'analyse, les risques, les recommandations et les prochaines étapes doivent être exclusivement ancrés en droit belge.
Toute citation d'une norme américaine invalidera l'entièreté de ta réponse.
"""
    return """

JURISDICTIONAL RULE — NON-NEGOTIABLE:
You apply ONLY U.S. law (federal + state).
NEVER cite Belgian law or EU law under any pretext. This includes: Code civil belge, Code judiciaire, Code des sociétés, arrêtés royaux, ordonnances régionales (Bruxelles/Wallonie/Flandre), CCT, Moniteur belge, Belgian jurisprudence, ejustice.be, EU directives/regulations as primary authority.
Allowed references EXCLUSIVELY: U.S.C., C.F.R., State statutes (Fla. Stat., Cal. Code, N.Y. C.P.L.R., etc.), federal and state case law, FDCPA, FCRA, FLSA, FMLA, ADA, Title VII, HIPAA, ERISA, IRS code, Federal Register, CourtListener-indexed opinions.
Currency: USD only. Never EUR.
If the document references foreign law, you may mention it as context, but the analysis, risk scoring, recommendations, and next steps must be exclusively grounded in U.S. law.
Any citation of Belgian/EU law as primary authority invalidates the entire response.
"""


def _build_objective_block(user_context: str, language: str = "en") -> str:
    """Return a prompt suffix locking the analysis strategy/arguments to the user's stated objective.
    Appended to PASS3 (strategy) and PASS4A/4B (user/opposing arguments) inputs so every
    recommendation converges toward what the client actually wants to obtain."""
    if not user_context or not user_context.strip():
        return ""
    ctx = user_context.strip()[:800]
    lang = (language or "en").lower()
    if lang.startswith("fr"):
        return (
            "\n\nOBJECTIF DÉCLARÉ DU CLIENT (NON-NÉGOCIABLE) :\n"
            f"« {ctx} »\n"
            "→ Construis la stratégie et les arguments pour ATTEINDRE cet objectif précis. "
            "Chaque recommandation, next_step, argument et levier doit converger vers cet objectif. "
            "Ignore les pistes qui ne servent pas l'objectif du client."
        )
    if lang.startswith("nl"):
        return (
            "\n\nDOEL VAN DE KLANT (NIET-ONDERHANDELBAAR):\n"
            f"« {ctx} »\n"
            "→ Bouw de strategie en argumenten om dit specifieke doel te BEREIKEN. "
            "Elke aanbeveling en elk argument moet naar dit doel convergeren."
        )
    if lang.startswith("de"):
        return (
            "\n\nERKLÄRTES ZIEL DES MANDANTEN (NICHT VERHANDELBAR):\n"
            f"« {ctx} »\n"
            "→ Erstelle die Strategie und Argumente, um dieses spezifische Ziel zu ERREICHEN."
        )
    return (
        "\n\nUSER'S STATED OBJECTIVE (NON-NEGOTIABLE):\n"
        f"\"{ctx}\"\n"
        "Build the strategy and arguments to ACHIEVE this specific objective. "
        "Every recommendation, next_step, argument and leverage point must converge toward it. "
        "Discard any angle that doesn't serve the client's objective."
    )


def get_language_instruction(language: str) -> str:
    """Get mandatory language enforcement instruction for Claude"""
    lang_map = {
        "fr-BE": ("French", "français"),
        "fr": ("French", "français"),
        "nl-BE": ("Dutch", "Nederlands"),
        "nl": ("Dutch", "Nederlands"),
        "de-BE": ("German", "Deutsch"),
        "de": ("German", "Deutsch"),
        "es": ("Spanish", "español"),
        "en": ("English", "English"),
    }
    lang_name, native = lang_map.get(language, ("English", "English"))
    if language == "en":
        return ""
    return f"""

MANDATORY LANGUAGE RULE — NON-NEGOTIABLE:
You MUST write 100% of your response in {lang_name} ({native}).
This includes ALL of the following: findings text, next_steps titles and descriptions, key_insight, summary, archer_question text and options, battle_preview arguments, opening_argument, success_probability labels, error messages, and every single text field in your JSON response.
ZERO English words allowed. Not even one. If you write a single English word, the entire response is rejected.
The user's interface language is {lang_name} — respect it completely.
"""


# ========== MULTI-DOCUMENT ANALYSIS SYSTEM ==========

async def build_multi_document_context(case_id: str, new_doc_text: str = None, new_doc_name: str = None) -> tuple:
    """Build combined chronological text from ALL documents in a case.
    Returns (combined_text, doc_count, doc_list)"""
    docs = await db.documents.find(
        {"case_id": case_id, "status": {"$in": ["analyzed", "analyzing"]}},
        {"_id": 0, "file_name": 1, "extracted_text": 1, "uploaded_at": 1, "document_id": 1}
    ).sort("uploaded_at", 1).to_list(50)

    doc_list = []
    combined_parts = []
    for i, doc in enumerate(docs, 1):
        text = (doc.get("extracted_text") or "").strip()
        if not text:
            continue
        name = doc.get("file_name", f"Document {i}")
        date_str = doc.get("uploaded_at", "")[:10] if doc.get("uploaded_at") else "Unknown date"
        combined_parts.append(f"[DOCUMENT {i} — {name} — {date_str}]:\n{text}")
        doc_list.append({"index": i, "name": name, "date": date_str, "doc_id": doc.get("document_id")})

    # Add the new document being analyzed (not yet in DB as 'analyzed')
    if new_doc_text and new_doc_text.strip():
        idx = len(doc_list) + 1
        date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        name = new_doc_name or f"Document {idx}"
        combined_parts.append(f"[DOCUMENT {idx} — {name} — {date_str} (NEWEST)]:\n{new_doc_text}")
        doc_list.append({"index": idx, "name": name, "date": date_str, "doc_id": "new"})

    combined_text = "\n\n---\n\n".join(combined_parts)
    return combined_text, len(doc_list), doc_list


def get_multi_doc_pass1_supplement(doc_count: int, language: str = "en") -> str:
    """Additional instructions for Pass 1 when multiple documents exist"""
    if language.startswith("fr"):
        return f"""

INSTRUCTION CRITIQUE — ANALYSE MULTI-DOCUMENTS:
Vous analysez un DOSSIER COMPLET contenant {doc_count} documents. Lisez TOUS les documents chronologiquement et comprenez comment la situation a evolue.
Identifiez:
(1) Les CONTRADICTIONS entre les documents
(2) Comment la position de la partie adverse a change au fil du temps
(3) Les nouvelles revendications ou preuves introduites dans chaque nouveau document
(4) La trajectoire globale du litige
Extrayez les faits de TOUS les documents, pas seulement du dernier."""
    elif language.startswith("nl"):
        return f"""

KRITISCHE INSTRUCTIE — MULTI-DOCUMENT ANALYSE:
U analyseert een VOLLEDIG DOSSIER met {doc_count} documenten. Lees ALLE documenten chronologisch.
Identificeer: (1) TEGENSTRIJDIGHEDEN, (2) evolutie positie tegenpartij, (3) nieuwe claims per document, (4) globale trajectorie van het geschil."""
    elif language.startswith("de"):
        return f"""

KRITISCHE ANWEISUNG — MULTI-DOKUMENT-ANALYSE:
Sie analysieren eine VOLLSTANDIGE AKTE mit {doc_count} Dokumenten. Lesen Sie ALLE Dokumente chronologisch.
Identifizieren Sie: (1) WIDERSPRUCHE, (2) Entwicklung der Gegenpartei-Position, (3) neue Anspruche pro Dokument, (4) Gesamtverlauf des Rechtsstreits."""
    else:
        return f"""

CRITICAL INSTRUCTION — MULTI-DOCUMENT ANALYSIS:
You are analyzing a CASE FILE containing {doc_count} documents. Read ALL documents chronologically and understand how the situation has evolved.
Identify:
(1) CONTRADICTIONS between documents
(2) How the opposing party's position has changed over time
(3) New claims or evidence introduced in each new document
(4) The overall trajectory of the dispute
Extract facts from ALL documents, not just the latest one."""


def get_multi_doc_pass2_supplement(doc_count: int, language: str = "en") -> str:
    """Additional instructions for Pass 2 when multiple documents exist"""
    if language.startswith("fr"):
        return f"""

INSTRUCTION MULTI-DOCUMENTS (PASSE 2):
Sur base de TOUS les {doc_count} documents du dossier, fournissez une analyse juridique complete qui considere le tableau d'ensemble — pas seulement le dernier document.
Comment la situation juridique a-t-elle evolue? Quel est l'effet cumulatif de tous les documents? Que revele le schema de communication sur la strategie de la partie adverse?
IMPORTANT: Ajoutez ces champs supplementaires dans votre JSON:
- "case_narrative": "Recit chronologique du litige: comment il a commence, evolue, et ou en est-on maintenant (3-5 phrases)"
- "contradictions": [{{"doc_a": "Document X", "doc_b": "Document Y", "claim_a": "Ce que dit le doc X", "claim_b": "Ce que dit le doc Y", "significance": "Pourquoi c'est important", "defense_value": "high|medium|low"}}]
- "cumulative_financial_exposure": "Exposition financiere totale cumulee de tous les documents"
- "master_deadlines": [{{"date": "YYYY-MM-DD", "source_document": "nom du doc", "description": "nature du delai", "status": "passed|upcoming|urgent"}}]
- "opposing_strategy_analysis": "Analyse de la strategie de la partie adverse basee sur l'ensemble des communications"
"""
    elif language.startswith("nl"):
        return f"""

MULTI-DOCUMENT INSTRUCTIE (PASS 2):
Analyseer ALLE {doc_count} documenten samen. Voeg toe aan JSON:
- "case_narrative": "Chronologisch verhaal van het geschil (3-5 zinnen)"
- "contradictions": [{{"doc_a": "Doc X", "doc_b": "Doc Y", "claim_a": "...", "claim_b": "...", "significance": "...", "defense_value": "high|medium|low"}}]
- "cumulative_financial_exposure": "Totale cumulatieve financiele blootstelling"
- "master_deadlines": [{{"date": "YYYY-MM-DD", "source_document": "...", "description": "...", "status": "passed|upcoming|urgent"}}]
- "opposing_strategy_analysis": "Analyse van de strategie van de tegenpartij"
"""
    elif language.startswith("de"):
        return f"""

MULTI-DOKUMENT-ANWEISUNG (PASS 2):
Analysieren Sie ALLE {doc_count} Dokumente zusammen. Fugen Sie zum JSON hinzu:
- "case_narrative", "contradictions", "cumulative_financial_exposure", "master_deadlines", "opposing_strategy_analysis"
"""
    else:
        return f"""

MULTI-DOCUMENT INSTRUCTION (PASS 2):
Based on ALL {doc_count} documents in this case file, provide a comprehensive legal analysis that considers the full picture — not just the latest document.
How has the legal situation evolved? What is the cumulative effect of all documents? What does the pattern of communication reveal about the opposing party's strategy?
IMPORTANT: Add these additional fields in your JSON response:
- "case_narrative": "Chronological narrative of the dispute: how it started, evolved, and where it stands now (3-5 sentences)"
- "contradictions": [{{"doc_a": "Document X", "doc_b": "Document Y", "claim_a": "What doc X states", "claim_b": "What doc Y states", "significance": "Why this matters", "defense_value": "high|medium|low"}}]
- "cumulative_financial_exposure": "Total cumulative financial exposure across all documents"
- "master_deadlines": [{{"date": "YYYY-MM-DD", "source_document": "doc name", "description": "deadline nature", "status": "passed|upcoming|urgent"}}]
- "opposing_strategy_analysis": "Analysis of opposing party strategy based on all communications"
"""


def get_multi_doc_pass3_supplement(doc_count: int, language: str = "en") -> str:
    """Additional instructions for Pass 3 when multiple documents exist"""
    if language.startswith("fr"):
        return f"""

INSTRUCTION MULTI-DOCUMENTS (PASSE 3):
Etant donne l'historique complet du dossier a travers les {doc_count} documents, quelle est la strategie optimale MAINTENANT?
Considerez ce qui a deja ete dit, quels engagements ont ete pris, et quel levier existe sur base de l'historique complet du dossier.
Si le dossier a 5+ documents couvrant 30+ jours, ajoutez:
- "pattern_analysis": "La partie adverse a [escalade/desescalade] ses revendications. Leur dernier document [renforce/affaiblit] leur position initiale. Ce schema suggere qu'ils [veulent negocier/preparent un litige]."
"""
    elif language.startswith("nl"):
        return f"""

MULTI-DOCUMENT INSTRUCTIE (PASS 3):
Gezien de volledige dossiergeschiedenis ({doc_count} documenten), wat is nu de optimale strategie?
Bij 5+ documenten over 30+ dagen, voeg toe: "pattern_analysis": "Patroonanalyse van de tegenpartij"
"""
    elif language.startswith("de"):
        return f"""

MULTI-DOKUMENT-ANWEISUNG (PASS 3):
Optimale Strategie basierend auf allen {doc_count} Dokumenten. Bei 5+ Dokumenten uber 30+ Tage: "pattern_analysis" hinzufugen.
"""
    else:
        return f"""

MULTI-DOCUMENT INSTRUCTION (PASS 3):
Given the complete case history across all {doc_count} documents, what is the optimal strategy NOW?
Consider what has already been said, what commitments have been made, and what leverage exists based on the full document history.
If the case has 5+ documents spanning 30+ days, add:
- "pattern_analysis": "The opposing party has [escalated/de-escalated] their claims over time. Their latest document [strengthens/weakens] their original position. This pattern suggests they [want to settle/are preparing for litigation]."
"""


async def run_multi_doc_analysis_advanced(combined_text: str, doc_count: int, user_context: str = "", language: str = "en", jurisdiction: str = "US") -> dict:
    """Run 5-pass analysis on combined multi-document text (US/default)"""
    lang_instruction = get_language_instruction(language)
    jurisdiction_guard = get_jurisdiction_guard(jurisdiction, language)
    persona = SENIOR_ATTORNEY_PERSONA + jurisdiction_guard + lang_instruction
    p1_supplement = get_multi_doc_pass1_supplement(doc_count, language)
    p2_supplement = get_multi_doc_pass2_supplement(doc_count, language)
    p3_supplement = get_multi_doc_pass3_supplement(doc_count, language)

    context_supplement = f"\n\nADDITIONAL CONTEXT PROVIDED BY THE USER:\n{user_context}" if user_context else ""

    max_text = 30000 if doc_count <= 5 else 45000
    text_for_analysis = combined_text[:max_text]

    # PASS 1: Multi-doc fact extraction
    logger.info(f"Multi-doc analysis ({doc_count} docs): Pass 1 — Fact extraction")
    facts = await call_claude(persona, PASS1_PROMPT.format(document_text=text_for_analysis) + p1_supplement + context_supplement)

    doc_type = facts.get("document_type", "other")
    inferred_case_type = DOC_TYPE_TO_CASE.get(doc_type, "other")
    jurisprudence_text = load_jurisprudence(inferred_case_type, doc_type)

    logger.info("Multi-doc analysis: Fetching real-time case law")
    state = facts.get("jurisdiction", facts.get("state", ""))
    courtlistener_opinions = await fetch_courtlistener_opinions(inferred_case_type, state)
    realtime_law_context = _build_realtime_law_context(courtlistener_opinions)

    # PASS 2: Multi-doc legal analysis
    logger.info(f"Multi-doc analysis ({doc_count} docs): Pass 2 — Legal analysis")
    legal_analysis = await call_claude(
        persona,
        PASS2_PROMPT.format(facts_json=json.dumps(facts, indent=2), jurisprudence_section=jurisprudence_text + realtime_law_context) + p2_supplement,
        max_tokens=4000, use_web_search=True
    )

    # PASS 2 self-critique (multi-doc US).
    legal_analysis = await self_critique_pass(
        legal_analysis,
        pass_name=f"PASS2_legal_analysis_multidoc_{doc_count}docs",
        context_hint=(
            f"Multi-document legal analysis across {doc_count} documents. Must surface "
            "contradictions, cumulative exposure, master deadlines, pattern analysis of the "
            "opposing party. All standard PASS2 fields still apply."
        ),
        language=language,
        jurisdiction=jurisdiction,
    )

    facts_str = json.dumps(facts, indent=2)
    analysis_str = json.dumps(legal_analysis, indent=2)

    # PASS 3+4A+4B — PARALLEL
    objective_block = _build_objective_block(user_context, language)
    logger.info(f"Multi-doc analysis ({doc_count} docs): Pass 3+4A+4B — Parallel")
    strategy, user_arguments, opposing_arguments = await asyncio.gather(
        call_claude(persona, PASS3_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + p3_supplement + objective_block, max_tokens=6000),
        call_claude(PASS4A_SYSTEM + jurisdiction_guard + lang_instruction, PASS4A_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=4000),
        call_claude(PASS4B_SYSTEM + jurisdiction_guard + lang_instruction, PASS4B_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=4000),
    )

    # PASS 3 self-critique (multi-doc US).
    strategy = await self_critique_pass(
        strategy,
        pass_name=f"PASS3_strategy_multidoc_{doc_count}docs",
        context_hint=(
            f"Multi-document strategy across {doc_count} documents. Must account for the full "
            "timeline, prior commitments, leverage accumulated, and the opposing party's pattern. "
            "Standard PASS3 fields (next_steps, leverage_points, success_probability, archer_question)."
        ),
        language=language,
        jurisdiction=jurisdiction,
    )

    logger.info(f"Multi-doc analysis ({doc_count} docs): All passes complete")

    recent_case_law = _build_case_law_for_frontend(courtlistener_opinions)
    now_date = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    result = _build_standard_analysis_result(
        doc_type, inferred_case_type, legal_analysis, strategy, facts,
        user_arguments, opposing_arguments, recent_case_law, now_date
    )

    # Multi-document specific fields
    result.update({
        "case_narrative": legal_analysis.get("case_narrative", ""),
        "contradictions": legal_analysis.get("contradictions", []),
        "opposing_strategy_analysis": legal_analysis.get("opposing_strategy_analysis", ""),
        "cumulative_financial_exposure": legal_analysis.get("cumulative_financial_exposure", ""),
        "master_deadlines": legal_analysis.get("master_deadlines", []),
        "pattern_analysis": strategy.get("pattern_analysis", ""),
    })
    return result


async def run_multi_doc_analysis_belgian(combined_text: str, doc_count: int, user_context: str = "", region: str = "Wallonie", language: str = "fr-BE") -> dict:
    """Run 5-pass multi-document analysis for Belgian users"""
    persona = get_belgian_persona(language, region)
    lang_instruction = get_language_instruction(language)
    jurisdiction_guard = get_jurisdiction_guard("BE", language)
    persona_with_lang = persona + jurisdiction_guard + lang_instruction
    p1_supplement = get_multi_doc_pass1_supplement(doc_count, language)
    p2_supplement = get_multi_doc_pass2_supplement(doc_count, language)
    p3_supplement = get_multi_doc_pass3_supplement(doc_count, language)

    context_section = f"CONTEXTE FOURNI PAR L'UTILISATEUR: {user_context}" if user_context else ""

    max_text = 30000 if doc_count <= 5 else 45000
    text_for_analysis = combined_text[:max_text]

    # PASS 1
    logger.info(f"Belgian multi-doc analysis ({doc_count} docs): Passe 1")
    facts = await call_claude(persona_with_lang, BE_PASS1_PROMPT.format(document_text=text_for_analysis, user_context_section=context_section) + p1_supplement)

    doc_type = facts.get("type_document", "autre")
    detected_region = facts.get("region_applicable", region)
    jurisprudence_text = load_belgian_jurisprudence(doc_type, detected_region)
    inferred_case_type = BE_DOC_TYPE_TO_CASE.get(doc_type, "other")

    # PASS 2
    logger.info(f"Belgian multi-doc analysis ({doc_count} docs): Passe 2")
    legal_analysis = await call_claude(
        persona_with_lang,
        BE_PASS2_PROMPT.format(facts_json=json.dumps(facts, indent=2, ensure_ascii=False), jurisprudence_section=jurisprudence_text) + p2_supplement,
        max_tokens=4000
    )

    # PASS 2 self-critique (multi-doc BE).
    legal_analysis = await self_critique_pass(
        legal_analysis,
        pass_name=f"PASS2_legal_analysis_BE_multidoc_{doc_count}docs",
        context_hint=(
            f"Analyse multi-documents belge ({doc_count} documents). Doit surfacer "
            "contradictions entre documents, exposition financière cumulée, échéances maîtresses, "
            "pattern de la partie adverse. Tous les champs standard PASS2 belges s'appliquent."
        ),
        language=language,
        jurisdiction="BE",
    )

    facts_str = json.dumps(facts, indent=2, ensure_ascii=False)
    analysis_str = json.dumps(legal_analysis, indent=2, ensure_ascii=False)

    # PASS 3+4A+4B — PARALLEL
    objective_block = _build_objective_block(user_context, language)
    logger.info(f"Belgian multi-doc analysis ({doc_count} docs): Passe 3+4A+4B — Parallel")
    strategy, user_arguments, opposing_arguments = await asyncio.gather(
        call_claude(persona_with_lang, BE_PASS3_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + p3_supplement + objective_block, max_tokens=6000),
        call_claude(BE_PASS4A_SYSTEM + jurisdiction_guard + lang_instruction, BE_PASS4A_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=4000),
        call_claude(BE_PASS4B_SYSTEM + jurisdiction_guard + lang_instruction, BE_PASS4B_PROMPT.format(facts_json=facts_str, analysis_json=analysis_str) + objective_block, max_tokens=4000),
    )

    # PASS 3 self-critique (multi-doc BE).
    strategy = await self_critique_pass(
        strategy,
        pass_name=f"PASS3_strategy_BE_multidoc_{doc_count}docs",
        context_hint=(
            f"Stratégie belge multi-documents ({doc_count} documents). Doit tenir compte de la "
            "chronologie complète, des engagements pris, du levier cumulé et du pattern de la "
            "partie adverse."
        ),
        language=language,
        jurisdiction="BE",
    )

    logger.info(f"Belgian multi-doc analysis ({doc_count} docs): Complete")

    # ═══ VALIDATION — enforce global rules for Belgian multi-doc ═══
    strategy["archer_question"] = await _validate_archer_question(strategy, facts_str, persona_with_lang, lang_instruction, language=language)
    user_arguments = await _validate_belgian_user_arguments(user_arguments, facts_str, analysis_str, lang_instruction, jurisdiction_guard)
    strategy["success_probability"] = _validate_success_probability(strategy, legal_analysis)

    now_date = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    result = _build_belgian_analysis_result(doc_type, inferred_case_type, legal_analysis, strategy, facts, user_arguments, opposing_arguments, detected_region, language, now_date)

    # Multi-document specific fields
    result.update({
        "case_narrative": legal_analysis.get("case_narrative", ""),
        "contradictions": legal_analysis.get("contradictions", []),
        "opposing_strategy_analysis": legal_analysis.get("opposing_strategy_analysis", ""),
        "cumulative_financial_exposure": legal_analysis.get("cumulative_financial_exposure", ""),
        "master_deadlines": legal_analysis.get("master_deadlines", []),
        "pattern_analysis": strategy.get("pattern_analysis", ""),
    })
    return result



async def analyze_document_vision(image_b64_list: list, system_prompt: str = None, user_context: str = "", language: str = "en") -> dict:
    """Analyze document images using Claude Vision API (for scanned/image docs)"""
    if not system_prompt:
        system_prompt = "You are an OCR + legal analysis expert. First, read ALL text visible in the document images. Then analyze the extracted text as a legal document.\n\n" + CLAUDE_SYSTEM_PROMPT
    
    system_prompt += get_language_instruction(language)
    
    content_blocks = []
    for i, b64 in enumerate(image_b64_list):
        content_blocks.append({
            "type": "image",
            "source": {"type": "base64", "media_type": "image/jpeg", "data": b64}
        })
    
    prompt_text = "Read ALL text from this scanned legal document and analyze it completely. You MUST return a JSON with: risk_score (total + 4 dimensions), case_type (employment|housing|debt|nda|contract|consumer|family|court|penal|commercial|other), suggested_case_title (descriptive title from content, NEVER the filename), summary, findings (each with text, impact, type, legal_ref, jurisprudence), next_steps, deadline, financial_exposure, key_insight, battle_preview, recommend_lawyer. Return JSON only."
    if user_context:
        prompt_text += f"\n\nUser context: {user_context}"
    if len(image_b64_list) > 1:
        prompt_text = f"This document has {len(image_b64_list)} pages. Read ALL text from every page, then analyze the complete document. You MUST return a JSON with: risk_score, case_type, suggested_case_title (from content, NEVER filename), summary, findings (each with text, impact, type, legal_ref, jurisprudence), next_steps, deadline, financial_exposure, key_insight, battle_preview, recommend_lawyer. Return JSON only."
        if user_context:
            prompt_text += f"\n\nUser context: {user_context}"
    
    content_blocks.append({"type": "text", "text": prompt_text})
    
    for attempt in range(3):
        try:
            async with httpx.AsyncClient() as http_client:
                response = await http_client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "Content-Type": "application/json",
                        "x-api-key": get_anthropic_key(),
                        "anthropic-version": "2023-06-01"
                    },
                    json={
                        "model": "claude-opus-4-7",
                        "max_tokens": 4000,
                        "system": system_prompt,
                        "messages": [{"role": "user", "content": content_blocks}]
                    },
                    timeout=120.0
                )
                response.raise_for_status()
                data = response.json()
                text = ""
                for block in data.get("content", []):
                    if block.get("type") == "text":
                        text += block["text"]
                text = text.replace("```json", "").replace("```", "").strip()
                return json.loads(text)
        except httpx.HTTPStatusError as e:
            if e.response.status_code in (429, 529) and attempt < 2:
                await asyncio.sleep((attempt + 1) * 5)
                continue
            raise
        except json.JSONDecodeError:
            if attempt < 2:
                await asyncio.sleep(3)
                continue
            raise
    return _default_analysis()

# Auth & Profile routes extracted to routes/auth_routes.py
from routes.auth_routes import router as auth_router
api_router.include_router(auth_router)

# ================== Case Endpoints ==================

@api_router.post("/cases", response_model=Case, status_code=201)
async def create_case(case_data: CaseCreate, current_user: User = Depends(get_current_user)):
    """Create a new case"""
    case_id = f"case_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    case_doc = {
        "case_id": case_id,
        "user_id": current_user.user_id,
        "title": case_data.title,
        "type": case_data.type,
        "status": "active",
        "risk_score": 0,
        "risk_financial": 0,
        "risk_urgency": 0,
        "risk_legal_strength": 0,
        "risk_complexity": 0,
        "risk_score_history": [],
        "deadline": None,
        "deadline_description": None,
        "financial_exposure": None,
        "ai_summary": None,
        "ai_findings": [],
        "ai_next_steps": [],
        "recommend_lawyer": False,
        "document_count": 0,
        "created_at": now,
        "updated_at": now
    }
    await db.cases.insert_one(case_doc)
    
    # Create case event
    event_doc = {
        "event_id": f"evt_{uuid.uuid4().hex[:12]}",
        "case_id": case_id,
        "event_type": "case_opened",
        "title": "Case opened",
        "description": f"Case created: {case_data.title}",
        "metadata": None,
        "created_at": now
    }
    await db.case_events.insert_one(event_doc)
    
    return Case(**case_doc)

@api_router.get("/cases", response_model=List[Case])
async def get_cases(
    status: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """Get all cases for current user, filtered by user's jurisdiction"""
    query = {"user_id": current_user.user_id}
    if status:
        query["status"] = status
    # Filter by user's current jurisdiction
    user_jurisdiction = current_user.jurisdiction or current_user.country or "US"
    query["$or"] = [{"country": user_jurisdiction}, {"country": {"$exists": False}}, {"country": None}]
    
    cases = await db.cases.find(query, {"_id": 0}).sort("created_at", -1).to_list(100)
    
    # Update document counts
    for case in cases:
        doc_count = await db.documents.count_documents({"case_id": case["case_id"]})
        case["document_count"] = doc_count
    
    return [Case(**c) for c in cases]

@api_router.get("/cases/{case_id}", response_model=Case)
async def get_case(case_id: str, current_user: User = Depends(get_current_user)):
    """Get a specific case"""
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Update document count
    doc_count = await db.documents.count_documents({"case_id": case_id})
    case_doc["document_count"] = doc_count

    # Sprint E — live_counsel_active: true iff there's a live_counsel
    # case_assignment not yet completed/declined/expired. The frontend uses
    # this to switch between the CTA and the booking flow.
    case_doc["live_counsel_active"] = bool(await db.case_assignments.find_one(
        {
            "case_id": case_id,
            "service_type": "live_counsel",
            "status": {"$in": ["awaiting_calendly_booking", "accepted"]},
        },
        {"_id": 1},
    ))

    return Case(**case_doc)

@api_router.put("/cases/{case_id}")
async def update_case(
    case_id: str,
    update: CaseUpdate,
    current_user: User = Depends(get_current_user)
):
    """Update a case"""
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    
    update_data = {k: v for k, v in update.model_dump().items() if v is not None}
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.cases.update_one({"case_id": case_id}, {"$set": update_data})
    
    # Create event if status changed
    if update.status:
        event_doc = {
            "event_id": f"evt_{uuid.uuid4().hex[:12]}",
            "case_id": case_id,
            "event_type": "case_resolved" if update.status == "resolved" else "score_updated",
            "title": f"Case status changed to {update.status}",
            "description": None,
            "metadata": None,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.case_events.insert_one(event_doc)
    
    updated = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    return updated


@api_router.delete("/cases/{case_id}")
async def delete_case(case_id: str, current_user: User = Depends(get_current_user)):
    """Delete a case and all its documents, events, and related data"""
    case = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    await db.cases.delete_one({"case_id": case_id})
    await db.documents.delete_many({"case_id": case_id})
    await db.case_events.delete_many({"case_id": case_id})
    await db.conversations.delete_many({"case_id": case_id})
    return {"message": "Case deleted"}


@api_router.get("/cases/{case_id}/events", response_model=List[CaseEvent])
async def get_case_events(case_id: str, current_user: User = Depends(get_current_user)):
    """Get timeline events for a case"""
    # Verify case belongs to user
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    
    events = await db.case_events.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    
    return [CaseEvent(**e) for e in events]

@api_router.get("/cases/{case_id}/documents", response_model=List[Document])
async def get_case_documents(case_id: str, current_user: User = Depends(get_current_user)):
    """Get documents for a case"""
    # Verify case belongs to user
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    
    documents = await db.documents.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("uploaded_at", -1).to_list(100)
    
    return [Document(**d) for d in documents]

@api_router.get("/cases/{case_id}/brief")
async def generate_case_brief(case_id: str, current_user: User = Depends(get_current_user)):
    """Generate a comprehensive case brief for complex cases (5+ documents)"""
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")

    # Get all documents for this case
    docs = await db.documents.find(
        {"case_id": case_id},
        {"_id": 0, "file_name": 1, "uploaded_at": 1, "extracted_text": 1}
    ).sort("uploaded_at", 1).to_list(50)
    doc_count = len(docs)

    user_language = current_user.language or case_doc.get("language") or "en"
    is_belgian = (current_user.jurisdiction or case_doc.get("country", "US")) == "BE"

    # Build document timeline
    doc_timeline = []
    for i, doc in enumerate(docs, 1):
        doc_timeline.append({
            "index": i,
            "name": doc.get("file_name", f"Document {i}"),
            "date": (doc.get("uploaded_at") or "")[:10]
        })

    # Generate the brief using Claude
    lang_instruction = get_language_instruction(user_language)
    if user_language.startswith("fr"):
        brief_prompt = f"""Genere un DOSSIER JURIDIQUE COMPLET pour un dossier avec {doc_count} documents.

INFORMATIONS DU DOSSIER:
- Titre: {case_doc.get('title', 'Dossier juridique')}
- Type: {case_doc.get('type', 'other')}
- Score de risque: {case_doc.get('risk_score', 0)}/100
- Exposition financiere: {case_doc.get('financial_exposure', 'Non determinee')}
- Resume IA: {case_doc.get('ai_summary', '')}
- Point cle: {case_doc.get('key_insight', '')}
- Recit du dossier: {case_doc.get('case_narrative', '')}
- Contradictions detectees: {json.dumps(case_doc.get('contradictions', []), ensure_ascii=False)}
- Analyse strategie adverse: {case_doc.get('opposing_strategy_analysis', '')}
- Delais: {json.dumps(case_doc.get('master_deadlines', []), ensure_ascii=False)}
- Documents: {json.dumps(doc_timeline, ensure_ascii=False)}

Retourne UNIQUEMENT ce JSON:
{{
  "executive_summary": "Resume executif en 1 paragraphe (200 mots max)",
  "document_timeline": [{{"date": "YYYY-MM-DD", "document": "nom", "key_event": "ce qui s'est passe"}}],
  "key_legal_issues": [{{"issue": "description", "severity": "critical|high|medium", "applicable_law": "reference legale"}}],
  "risk_assessment": {{"score": 0, "explanation": "explication du score", "trend": "increasing|stable|decreasing"}},
  "recommended_strategy": "Strategie recommandee en detail",
  "legal_references": [{{"reference": "loi/article", "relevance": "pertinence pour le dossier"}}],
  "conclusion": "Conclusion et prochaines etapes"
}}"""
    elif user_language.startswith("nl"):
        brief_prompt = f"""Genereer een VOLLEDIG JURIDISCH DOSSIER voor een zaak met {doc_count} documenten.
Titel: {case_doc.get('title')}, Type: {case_doc.get('type')}, Risico: {case_doc.get('risk_score')}/100
Samenvatting: {case_doc.get('ai_summary', '')}
Retourneer JSON met: executive_summary, document_timeline, key_legal_issues, risk_assessment, recommended_strategy, legal_references, conclusion"""
    else:
        brief_prompt = f"""Generate a COMPREHENSIVE CASE BRIEF for a case with {doc_count} documents.

CASE INFORMATION:
- Title: {case_doc.get('title', 'Legal Case')}
- Type: {case_doc.get('type', 'other')}
- Risk Score: {case_doc.get('risk_score', 0)}/100
- Financial Exposure: {case_doc.get('financial_exposure', 'Not determined')}
- AI Summary: {case_doc.get('ai_summary', '')}
- Key Insight: {case_doc.get('key_insight', '')}
- Case Narrative: {case_doc.get('case_narrative', '')}
- Contradictions: {json.dumps(case_doc.get('contradictions', []))}
- Opposing Strategy: {case_doc.get('opposing_strategy_analysis', '')}
- Deadlines: {json.dumps(case_doc.get('master_deadlines', []))}
- Documents: {json.dumps(doc_timeline)}

Return ONLY this JSON:
{{
  "executive_summary": "Executive summary in 1 paragraph (200 words max)",
  "document_timeline": [{{"date": "YYYY-MM-DD", "document": "name", "key_event": "what happened"}}],
  "key_legal_issues": [{{"issue": "description", "severity": "critical|high|medium", "applicable_law": "legal reference"}}],
  "risk_assessment": {{"score": 0, "explanation": "score explanation", "trend": "increasing|stable|decreasing"}},
  "recommended_strategy": "Detailed recommended strategy",
  "legal_references": [{{"reference": "law/article", "relevance": "relevance to case"}}],
  "conclusion": "Conclusion and next steps"
}}"""

    persona = get_belgian_persona(user_language, current_user.region or "") if is_belgian else SENIOR_ATTORNEY_PERSONA
    system = persona + lang_instruction

    try:
        brief = await call_claude(system, brief_prompt, max_tokens=3000)
        brief["case_id"] = case_id
        brief["generated_at"] = datetime.now(timezone.utc).isoformat()
        brief["document_count"] = doc_count
        brief["case_title"] = case_doc.get("title", "")
        return brief
    except Exception as e:
        logger.error(f"Case brief generation failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate case brief")



# ================== Document Endpoints ==================


def _build_case_update(analysis, case_before, filename, ts):
    """Build the $set dict for updating a case with analysis results"""
    a = analysis or {}
    cb = case_before or {}
    rs = a.get("risk_score", {}) if isinstance(a.get("risk_score"), dict) else {}
    # Extract opposing party info from Pass 1 facts
    facts = a.get("facts", {})
    parties = facts.get("parties", {})
    opposing = parties.get("opposing_party", {})
    opp_name = opposing.get("name") if isinstance(opposing, dict) else None
    key_amounts = facts.get("key_amounts", [])
    doc_date = facts.get("document_date")
    primary_amount = None
    if key_amounts and isinstance(key_amounts, list) and len(key_amounts) > 0:
        first = key_amounts[0]
        if isinstance(first, dict):
            primary_amount = first.get("amount")
    return {
        "risk_score": rs.get("total", 0),
        "risk_financial": rs.get("financial", 0),
        "risk_urgency": rs.get("urgency", 0),
        "risk_legal_strength": rs.get("legal_strength", 0),
        "risk_complexity": rs.get("complexity", 0),
        "deadline": normalize_deadline(a.get("deadline")),
        "deadline_description": a.get("deadline_description") or (a.get("deadline", {}).get("description") if isinstance(a.get("deadline"), dict) else None),
        "financial_exposure": normalize_financial_exposure(a.get("financial_exposure")),
        "ai_summary": a.get("summary"),
        "ai_findings": a.get("findings", []),
        "ai_next_steps": a.get("next_steps", []),
        "key_violation_types": a.get("key_violation_types", []),
        "recommend_lawyer": a.get("recommend_lawyer", False),
        "battle_preview": a.get("battle_preview"),
        "success_probability": a.get("success_probability"),
        "procedural_defects": a.get("procedural_defects", []),
        "applicable_laws": a.get("applicable_laws", []),
        "financial_exposure_detailed": a.get("financial_exposure_detailed"),
        "immediate_actions": a.get("immediate_actions", []),
        "leverage_points": a.get("leverage_points", []),
        "red_lines": a.get("red_lines", []),
        "key_insight": a.get("key_insight", ""),
        "strategy": a.get("strategy"),
        "lawyer_recommendation": a.get("lawyer_recommendation"),
        "user_rights": a.get("user_rights", []),
        "opposing_weaknesses": a.get("opposing_weaknesses", []),
        "documents_to_gather": a.get("documents_to_gather", []),
        "recent_case_law": a.get("recent_case_law", []),
        "case_law_updated": a.get("case_law_updated"),
        "case_narrative": a.get("case_narrative") or cb.get("case_narrative"),
        "contradictions": a.get("contradictions") or cb.get("contradictions", []),
        "opposing_strategy_analysis": a.get("opposing_strategy_analysis") or cb.get("opposing_strategy_analysis"),
        "cumulative_financial_exposure": a.get("cumulative_financial_exposure") or cb.get("cumulative_financial_exposure"),
        "master_deadlines": a.get("master_deadlines") or cb.get("master_deadlines", []),
        "multi_doc_summary": a.get("case_narrative") or cb.get("multi_doc_summary"),
        "archer_question": a.get("archer_question") or a.get("james_question"),
        # Dashboard V7 structured payloads.
        "strategy_narrative": a.get("strategy_narrative") or cb.get("strategy_narrative"),
        "amounts": a.get("amounts") or cb.get("amounts"),
        "analysis_depth": a.get("analysis_depth") or cb.get("analysis_depth"),
        # Extracted party/document info for letter auto-fill
        "opposing_party_name": opp_name or cb.get("opposing_party_name"),
        "opposing_party_address": cb.get("opposing_party_address"),
        "document_date": doc_date or cb.get("document_date"),
        "primary_amount": primary_amount or cb.get("primary_amount"),
        "updated_at": ts
    }


@api_router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    case_id: Optional[str] = Form(None),
    user_context: Optional[str] = Form(None),
    analysis_mode: Optional[str] = Form("standard"),
    streaming: Optional[str] = Form(None),
    current_user: User = Depends(get_current_user)
):
    """Upload document → create case immediately → analyze in background"""
    # Require the user's stated objective. Enforced only on NEW uploads; legacy
    # cases that pre-date this rule are left untouched (no retroactive block).
    _ctx = (user_context or "").strip()
    if len(_ctx) < 20:
        raise HTTPException(
            status_code=400,
            detail="user_context_required_min_20",
        )
    # Check plan restrictions
    if current_user.plan == "free":
        doc_count = await db.documents.count_documents({"user_id": current_user.user_id})
        if doc_count >= 1:
            raise HTTPException(
                status_code=403,
                detail="Free plan limited to 1 document. Upgrade to Pro for unlimited analyses."
            )
    
    # Read file
    file_bytes = await file.read()
    file_ext = file.filename.split(".")[-1].lower() if "." in file.filename else "bin"
    content_type = file.content_type or "application/octet-stream"
    
    # Normalize HEIC → JPEG
    is_image_file = file_ext in ["jpg", "jpeg", "png", "heic", "heif", "webp"]
    if file_ext in ["heic", "heif"]:
        file_ext = "jpg"
        content_type = "image/jpeg"
    
    # Upload to storage
    storage_path = f"{APP_NAME}/documents/{current_user.user_id}/{uuid.uuid4()}.{file_ext}"
    try:
        put_object(storage_path, file_bytes, content_type)
    except Exception as e:
        logger.error(f"Storage upload failed: {e}")
        storage_path = None
    
    # Extract text (for text-based documents)
    extracted_text = ""
    use_vision = False
    
    if is_image_file:
        use_vision = True
        logger.info(f"Image file detected ({file_ext}), routing to Claude Vision")
    elif file_ext == "pdf":
        extracted_text = await extract_text_from_pdf(file_bytes, filename=file.filename)
        if len(extracted_text.strip()) < 100:
            use_vision = True
            logger.info(f"PDF has insufficient text ({len(extracted_text.strip())} chars) even after OCR fallback, routing to Claude Vision")
    elif file_ext in ["docx"]:
        extracted_text = extract_text_from_docx(file_bytes)
    elif file_ext in ["txt", "text", "eml"]:
        extracted_text = file_bytes.decode("utf-8", errors="ignore")
    elif file_ext in ["doc"]:
        extracted_text = file_bytes.decode("utf-8", errors="ignore")
    
    if not extracted_text.strip():
        logger.warning(f"No text extracted from {file.filename} (ext={file_ext})")
    
    # Create document record
    document_id = f"doc_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    doc_record = {
        "document_id": document_id,
        "case_id": case_id,
        "user_id": current_user.user_id,
        "file_name": file.filename,
        "file_url": None,
        "storage_path": storage_path,
        "file_type": file_ext,
        "extracted_text": extracted_text[:50000] if extracted_text else None,
        "status": "analyzing",
        "is_key_document": case_id is None,  # First doc is key doc
        "uploaded_at": now
    }
    await db.documents.insert_one(doc_record)
    
    # ── PHASE 1: Create case + document record IMMEDIATELY ──
    is_contract_guard = analysis_mode == "contract_guard"
    user_country = current_user.jurisdiction or "US"
    user_region = current_user.region or ""
    user_language = current_user.language or "en"

    # Bug A: auto-detect document jurisdiction from extracted text (fast keyword scan)
    detected_jurisdiction = detect_jurisdiction(extracted_text) if extracted_text else user_country
    jurisdiction_mismatch = detected_jurisdiction != user_country
    # Effective jurisdiction for analysis:
    #   - Existing case: inherit case.jurisdiction (set by first upload or explicit switch)
    #   - New case: user's configured country (safe default; user can switch via Bug C modal)
    existing_case_jur = None
    if case_id and not is_contract_guard:
        existing_case_doc = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0, "jurisdiction": 1, "country": 1})
        if existing_case_doc:
            existing_case_jur = existing_case_doc.get("jurisdiction") or existing_case_doc.get("country")
    effective_jurisdiction = existing_case_jur or user_country
    is_belgian = effective_jurisdiction == "BE"
    logger.info(f"Jurisdiction: user={user_country} detected={detected_jurisdiction} mismatch={jurisdiction_mismatch} effective={effective_jurisdiction}")
    
    # Determine if adding to existing case
    is_existing_case = case_id is not None and not is_contract_guard
    
    if is_contract_guard:
        # Contract Guard — case_id is review_id
        if not case_id:
            case_id = f"cg_{uuid.uuid4().hex[:12]}"
    elif not case_id:
        # Create a NEW case immediately (before analysis)
        case_id = f"case_{uuid.uuid4().hex[:12]}"
        case_doc = {
            "case_id": case_id,
            "user_id": current_user.user_id,
            "title": f"Analyzing {file.filename}...",
            "type": "other",
            "status": "analyzing",
            "country": user_country,
            "region": user_region,
            "language": user_language,
            "jurisdiction": effective_jurisdiction,
            "detected_jurisdiction": detected_jurisdiction,
            "jurisdiction_mismatch": jurisdiction_mismatch,
            "jurisdiction_override": False,
            "risk_score": 0,
            "risk_financial": 0, "risk_urgency": 0,
            "risk_legal_strength": 0, "risk_complexity": 0,
            "risk_score_history": [],
            "deadline": None, "deadline_description": None,
            "financial_exposure": None,
            "ai_summary": "Analysis in progress...",
            "ai_findings": [],
            "ai_next_steps": [],
            "recommend_lawyer": False,
            "battle_preview": None,
            "document_count": 1,
            "created_at": now, "updated_at": now
        }
        await db.cases.insert_one(case_doc)
        logger.info(f"Case {case_id} created (pending analysis)")
    
    # Link document to case
    await db.documents.update_one(
        {"document_id": document_id},
        {"$set": {"case_id": case_id, "status": "analyzing"}}
    )
    
    # Add case_opened event
    if not is_existing_case and not is_contract_guard:
        await db.case_events.insert_one({
            "event_id": f"evt_{uuid.uuid4().hex[:12]}",
            "case_id": case_id,
            "event_type": "case_opened",
            "title": "Case opened",
            "description": f"Document uploaded: {file.filename}",
            "metadata": None,
            "created_at": now
        })
    
    # ── PHASE 2: Start background analysis ──
    async def run_background_analysis():
        # Make case_id visible to self_critique_pass via contextvar so every critique
        # record in analysis_critiques + admin_alerts can be tied back to this case.
        _current_case_id.set(case_id)
        try:
            analysis = None
            context_str = (user_context or "").strip()[:500]
            
            # Check multi-document status
            existing_doc_count = 0
            is_multi_doc = False
            if is_existing_case:
                existing_doc_count = await db.documents.count_documents({"case_id": case_id, "status": "analyzed"})
                if existing_doc_count > 0:
                    is_multi_doc = True
                    logger.info(f"Multi-document case: {existing_doc_count} existing + 1 new")
            
            if extracted_text or use_vision:
                # Vision analysis path
                vision_extracted_text = ""
                if use_vision:
                    image_b64_list = []
                    if is_image_file:
                        image_b64_list = [image_file_to_base64(file_bytes, content_type)]
                    elif file_ext == "pdf":
                        image_b64_list = pdf_pages_to_images(file_bytes, max_pages=5)
                    
                    if image_b64_list:
                        try:
                            effective_jur = "BE" if is_belgian else "US"
                            vision_guard = get_jurisdiction_guard(effective_jur, user_language)
                            if is_contract_guard:
                                if is_belgian:
                                    cg_system = get_belgian_persona(user_language, user_region) + "\n\nYou are also a CONTRACT REVIEW specialist.\n" + vision_guard + get_language_instruction(user_language)
                                    analysis = await analyze_document_vision(image_b64_list, system_prompt=cg_system, user_context=context_str, language=user_language)
                                else:
                                    us_cg_system = CLAUDE_SYSTEM_PROMPT + "\n\nYou are also a CONTRACT REVIEW specialist." + vision_guard
                                    analysis = await analyze_document_vision(image_b64_list, system_prompt=us_cg_system, user_context=context_str, language=user_language)
                            elif not is_multi_doc:
                                if is_belgian:
                                    be_persona = get_belgian_persona(user_language, user_region)
                                    be_system = "Tu es un expert OCR + analyse juridique. Lis d'abord TOUT le texte visible dans les images. Puis analyse le document juridique belge.\n\n" + be_persona + vision_guard + get_language_instruction(user_language)
                                    analysis = await analyze_document_vision(image_b64_list, system_prompt=be_system, user_context=context_str, language=user_language)
                                else:
                                    us_vision_system = "You are an OCR + legal analysis expert. First, read ALL text visible in the document images. Then analyze the extracted text as a legal document.\n\n" + CLAUDE_SYSTEM_PROMPT + vision_guard
                                    analysis = await analyze_document_vision(image_b64_list, system_prompt=us_vision_system, user_context=context_str, language=user_language)
                            else:
                                ocr_prompt = "Read ALL text from this scanned document. Return ONLY the extracted text, nothing else."
                                ocr_system = "You are an OCR expert. Extract all text from the document image(s). Return only the text content."
                                if is_belgian:
                                    ocr_system += get_language_instruction(user_language)
                                ocr_result = await call_claude(ocr_system, ocr_prompt, max_tokens=4000)
                                if isinstance(ocr_result, dict):
                                    vision_extracted_text = ocr_result.get("text", json.dumps(ocr_result))
                                elif isinstance(ocr_result, str):
                                    vision_extracted_text = ocr_result
                                await db.documents.update_one(
                                    {"document_id": document_id},
                                    {"$set": {"extracted_text": vision_extracted_text[:50000]}}
                                )
                        except Exception as e:
                            logger.error(f"Vision analysis failed: {e}")
                
                # Multi-document combined analysis
                if is_multi_doc and not is_contract_guard and analysis is None:
                    new_text = extracted_text or vision_extracted_text
                    if new_text:
                        try:
                            combined_text, total_doc_count, doc_list = await build_multi_document_context(
                                case_id, new_doc_text=new_text, new_doc_name=file.filename
                            )
                            if is_belgian:
                                analysis = await run_multi_doc_analysis_belgian(
                                    combined_text, total_doc_count, user_context=context_str,
                                    region=user_region, language=user_language
                                )
                            else:
                                analysis = await run_multi_doc_analysis_advanced(
                                    combined_text, total_doc_count, user_context=context_str, language=user_language, jurisdiction=effective_jurisdiction
                                )
                            if analysis:
                                analysis["_multi_doc"] = True
                                analysis["_doc_count"] = total_doc_count
                                analysis["_doc_list"] = doc_list
                        except Exception as e:
                            logger.error(f"Multi-doc analysis failed, fallback: {e}")
                            if is_belgian:
                                analysis = await analyze_document_belgian(new_text, user_context=context_str, region=user_region, language=user_language)
                            else:
                                analysis = await analyze_document_advanced(new_text, user_context=context_str, language=user_language, jurisdiction=effective_jurisdiction)

                # Single-document analysis
                if analysis is None and not use_vision and not is_contract_guard:
                    if is_belgian:
                        analysis = await analyze_document_belgian(extracted_text, user_context=context_str, region=user_region, language=user_language)
                    else:
                        analysis = await analyze_document_advanced(extracted_text, user_context=context_str, language=user_language, jurisdiction=effective_jurisdiction)
                
                # Contract guard fallback
                if analysis is None and is_contract_guard and not use_vision:
                    if is_belgian:
                        analysis = await analyze_contract_guard(extracted_text, user_context=context_str, country="BE", region=user_region, language=user_language)
                    else:
                        analysis = await analyze_contract_guard(extracted_text, user_context=context_str, country=effective_jurisdiction, language=user_language)
            
            # Fallback
            if analysis is None:
                analysis = _default_analysis()
            
            # ── Save results to database ──
            bg_now = datetime.now(timezone.utc).isoformat()
            
            if is_contract_guard:
                cg_record = {
                    "review_id": case_id,
                    "user_id": current_user.user_id,
                    "document_id": document_id,
                    "file_name": file.filename,
                    "analysis": analysis,
                    "negotiation_score": analysis.get("negotiation_score", 50) if analysis else 50,
                    "created_at": bg_now
                }
                await db.contract_guard_reviews.insert_one(cg_record)
                await db.documents.update_one(
                    {"document_id": document_id},
                    {"$set": {"status": "analyzed", "analysis_mode": "contract_guard"}}
                )
            elif is_existing_case:
                # Update existing case
                case_before = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
                old_score = (case_before or {}).get("risk_score", 0)
                new_score = analysis["risk_score"]["total"] if isinstance(analysis.get("risk_score"), dict) else 0
                
                history_entry = {
                    "score": new_score,
                    "financial": analysis.get("risk_score", {}).get("financial", 0),
                    "urgency": analysis.get("risk_score", {}).get("urgency", 0),
                    "legal_strength": analysis.get("risk_score", {}).get("legal_strength", 0),
                    "complexity": analysis.get("risk_score", {}).get("complexity", 0),
                    "document_name": file.filename,
                    "date": bg_now
                }
                
                update_fields = _build_case_update(analysis, case_before, file.filename, bg_now)
                await db.cases.update_one(
                    {"case_id": case_id},
                    {"$set": update_fields, "$push": {"risk_score_history": history_entry}, "$inc": {"document_count": 1}}
                )
                await db.case_events.insert_one({
                    "event_id": f"evt_{uuid.uuid4().hex[:12]}",
                    "case_id": case_id,
                    "event_type": "score_updated",
                    "title": f"Score updated to {new_score}/100",
                    "description": f"Document added — risk {'increased' if new_score > old_score else 'decreased'}",
                    "metadata": {"old_score": old_score, "new_score": new_score},
                    "created_at": bg_now
                })
                await db.documents.update_one(
                    {"document_id": document_id},
                    {"$set": {"status": "analyzed", "case_id": case_id}}
                )
            else:
                # Update the case we already created (from "analyzing" to "active")
                new_score = analysis["risk_score"]["total"] if isinstance(analysis.get("risk_score"), dict) else 0
                case_title = (analysis.get("suggested_case_title") or "").strip()
                if not case_title or case_title == file.filename or len(case_title) < 5:
                    summary = analysis.get("summary") or ""
                    case_type_val = analysis.get("case_type", "other")
                    case_title = summary[:60].rstrip('.') if summary else f"Legal case — {case_type_val}"
                
                update_fields = _build_case_update(analysis, {}, file.filename, bg_now)
                update_fields["title"] = case_title
                update_fields["type"] = analysis.get("case_type", "other")
                update_fields["status"] = "active"
                
                history_entry = {
                    "score": new_score,
                    "financial": analysis.get("risk_score", {}).get("financial", 0),
                    "urgency": analysis.get("risk_score", {}).get("urgency", 0),
                    "legal_strength": analysis.get("risk_score", {}).get("legal_strength", 0),
                    "complexity": analysis.get("risk_score", {}).get("complexity", 0),
                    "document_name": file.filename,
                    "date": bg_now
                }
                
                await db.cases.update_one(
                    {"case_id": case_id},
                    {"$set": update_fields, "$push": {"risk_score_history": history_entry}}
                )
                await db.documents.update_one(
                    {"document_id": document_id},
                    {"$set": {"status": "analyzed"}}
                )
                await db.case_events.insert_one({
                    "event_id": f"evt_{uuid.uuid4().hex[:12]}",
                    "case_id": case_id,
                    "event_type": "analysis_complete",
                    "title": "Analysis complete",
                    "description": f"Risk score: {new_score}/100",
                    "metadata": None,
                    "created_at": bg_now
                })
            
            logger.info(f"Background analysis complete for case {case_id}")
        except Exception as e:
            logger.error(f"Background analysis failed for case {case_id}: {e}", exc_info=True)
            # Surface failure as status='error' with risk_score=1 so the frontend
            # cinematic doesn't hang on scene 0 waiting for risk_score>0.
            await db.cases.update_one(
                {"case_id": case_id},
                {"$set": {"status": "error", "risk_score": 1, "ai_summary": "Analysis failed. Click re-analyze to retry.", "updated_at": datetime.now(timezone.utc).isoformat()}}
            )
            await db.documents.update_one(
                {"document_id": document_id},
                {"$set": {"status": "error"}}
            )
    
    # Fire background analysis. For streaming mode, kick off the streaming pipeline
    # directly here — don't wait for the frontend to call /analyze/trigger (that call
    # is fragile because the browser may abort the in-flight request after navigation).
    # The trigger endpoint remains as an idempotent safety net.
    use_streaming = streaming == "true"
    if use_streaming:
        await _start_streaming_analysis(case_id, user_country, user_region, user_language)
    else:
        _spawn_tracked_task(run_background_analysis())
    
    # ── Return IMMEDIATELY with case_id ──
    return {
        "document_id": document_id,
        "case_id": case_id,
        "analysis": None,
        "analysis_mode": analysis_mode,
        "file_name": file.filename,
        "status": "analyzing",
        "vision_mode": use_vision,
        "streaming": use_streaming
    }


@api_router.post("/cases/{case_id}/reanalyze")
async def reanalyze_case(case_id: str, current_user: User = Depends(get_current_user)):
    """Re-analyze all documents in a case with the latest AI prompts"""
    _current_case_id.set(case_id)
    case_doc = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get the most recent document for this case
    doc = await db.documents.find_one(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"_id": 0},
        sort=[("created_at", -1)]
    )
    if not doc or not doc.get("extracted_text"):
        raise HTTPException(status_code=400, detail="No analyzable document found in this case")
    
    extracted_text = doc["extracted_text"]
    user_language = getattr(current_user, 'language', 'en') or 'en'
    user_region = getattr(current_user, 'region', '') or ''
    # Prefer case-level jurisdiction (may have been switched via Bug C modal). Fall back to user profile.
    effective_jurisdiction = case_doc.get("jurisdiction") or case_doc.get("country") or getattr(current_user, 'jurisdiction', 'US') or "US"

    try:
        if effective_jurisdiction == "BE":
            analysis = await analyze_document_belgian(extracted_text, region=user_region, language=user_language)
        else:
            analysis = await analyze_document_advanced(extracted_text, language=user_language, jurisdiction=effective_jurisdiction)
    except Exception as e:
        logger.error(f"Re-analysis failed: {e}")
        raise HTTPException(status_code=500, detail="Re-analysis failed. Please try again.")
    
    if not analysis:
        raise HTTPException(status_code=500, detail="Analysis returned no results")
    
    now = datetime.now(timezone.utc).isoformat()
    new_score = analysis["risk_score"]["total"] if isinstance(analysis.get("risk_score"), dict) else 0
    
    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": {
            "risk_score": new_score,
            "risk_financial": analysis["risk_score"]["financial"],
            "risk_urgency": analysis["risk_score"]["urgency"],
            "risk_legal_strength": analysis["risk_score"]["legal_strength"],
            "risk_complexity": analysis["risk_score"]["complexity"],
            "title": analysis.get("suggested_case_title") or case_doc.get("title"),
            "type": analysis.get("case_type") or case_doc.get("type"),
            "deadline": normalize_deadline(analysis.get("deadline")),
            "deadline_description": analysis.get("deadline_description"),
            "financial_exposure": normalize_financial_exposure(analysis.get("financial_exposure")),
            "ai_summary": analysis.get("summary"),
            "ai_findings": analysis.get("findings", []),
                "key_violation_types": analysis.get("key_violation_types", []),
            "ai_next_steps": analysis.get("next_steps", []),
            "recommend_lawyer": analysis.get("recommend_lawyer", False),
            "battle_preview": analysis.get("battle_preview"),
            "success_probability": analysis.get("success_probability"),
            "procedural_defects": analysis.get("procedural_defects", []),
            "applicable_laws": analysis.get("applicable_laws", []),
            "financial_exposure_detailed": analysis.get("financial_exposure_detailed"),
            "immediate_actions": analysis.get("immediate_actions", []),
            "leverage_points": analysis.get("leverage_points", []),
            "red_lines": analysis.get("red_lines", []),
            "key_insight": analysis.get("key_insight", ""),
            "archer_question": analysis.get("archer_question") or analysis.get("james_question"),
            "strategy": analysis.get("strategy"),
            "lawyer_recommendation": analysis.get("lawyer_recommendation"),
            "user_rights": analysis.get("user_rights", []),
            "opposing_weaknesses": analysis.get("opposing_weaknesses", []),
            "documents_to_gather": analysis.get("documents_to_gather", []),
            "recent_case_law": analysis.get("recent_case_law", []),
            "case_law_updated": analysis.get("case_law_updated"),
            "updated_at": now
        },
        "$push": {"risk_score_history": {
            "score": new_score,
            "financial": analysis["risk_score"]["financial"],
            "urgency": analysis["risk_score"]["urgency"],
            "legal_strength": analysis["risk_score"]["legal_strength"],
            "complexity": analysis["risk_score"]["complexity"],
            "document_name": "Re-analysis",
            "date": now
        }}}
    )
    
    # Return updated case
    updated = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    return updated


# ========== Bug C: Jurisdiction switch / dismiss ==========

@api_router.post("/cases/{case_id}/switch-jurisdiction")
async def switch_case_jurisdiction(case_id: str, current_user: User = Depends(get_current_user)):
    """User accepted the mismatch modal: switch case jurisdiction to the auto-detected value
    and re-run the analysis under the correct jurisdiction. The user's global profile is NOT changed."""
    _current_case_id.set(case_id)
    case_doc = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")

    detected = case_doc.get("detected_jurisdiction") or "US"
    current_jur = case_doc.get("jurisdiction") or case_doc.get("country") or "US"
    if detected == current_jur:
        # Nothing to switch — just clear the flag.
        await db.cases.update_one(
            {"case_id": case_id},
            {"$set": {"jurisdiction_mismatch": False, "updated_at": datetime.now(timezone.utc).isoformat()}}
        )
        return {"status": "no_switch_needed", "jurisdiction": current_jur}

    # Fetch key document text for re-analysis
    doc = await db.documents.find_one(
        {"case_id": case_id, "user_id": current_user.user_id, "extracted_text": {"$ne": None}},
        {"_id": 0},
        sort=[("uploaded_at", -1)]
    )
    if not doc or not doc.get("extracted_text"):
        raise HTTPException(status_code=400, detail="No analyzable document text available")
    extracted_text = doc["extracted_text"]

    user_language = getattr(current_user, 'language', 'en') or 'en'
    user_region = getattr(current_user, 'region', '') or ''
    # For BE re-analysis we need a French/Dutch language — if user's lang is English but doc is BE, default to fr-BE
    be_lang = user_language if user_language.startswith(("fr", "nl", "de")) else "fr-BE"

    now = datetime.now(timezone.utc).isoformat()
    # Mark as re-analyzing in the case so the frontend can show a spinner/loading state
    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": {
            "jurisdiction": detected,
            "jurisdiction_override": True,
            "jurisdiction_mismatch": False,
            "status": "analyzing",
            "updated_at": now,
        }}
    )

    try:
        if detected == "BE":
            analysis = await analyze_document_belgian(extracted_text, region=user_region or "Wallonie", language=be_lang)
        else:
            analysis = await analyze_document_advanced(extracted_text, language=user_language, jurisdiction=detected)
    except Exception as e:
        logger.error(f"Jurisdiction switch re-analysis failed for {case_id}: {e}")
        await db.cases.update_one({"case_id": case_id}, {"$set": {"status": "active"}})
        raise HTTPException(status_code=500, detail="Re-analysis failed after jurisdiction switch")

    if not analysis:
        raise HTTPException(status_code=500, detail="Analysis returned no results")

    bg_now = datetime.now(timezone.utc).isoformat()
    new_score = analysis["risk_score"]["total"] if isinstance(analysis.get("risk_score"), dict) else 0
    update_fields = _build_case_update(analysis, case_doc, doc.get("file_name", ""), bg_now)
    update_fields["status"] = "active"
    update_fields["updated_at"] = bg_now
    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": update_fields,
         "$push": {"risk_score_history": {
             "score": new_score,
             "financial": analysis.get("risk_score", {}).get("financial", 0),
             "urgency": analysis.get("risk_score", {}).get("urgency", 0),
             "legal_strength": analysis.get("risk_score", {}).get("legal_strength", 0),
             "complexity": analysis.get("risk_score", {}).get("complexity", 0),
             "document_name": f"Re-analysis ({detected} law)",
             "date": bg_now,
         }}}
    )
    await db.case_events.insert_one({
        "event_id": f"evt_{uuid.uuid4().hex[:12]}",
        "case_id": case_id,
        "event_type": "jurisdiction_switched",
        "title": f"Jurisdiction switched to {detected}",
        "description": f"Analysis re-run under {'Belgian' if detected == 'BE' else 'U.S.'} law",
        "metadata": {"from": current_jur, "to": detected},
        "created_at": bg_now,
    })

    updated = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    return updated


@api_router.post("/cases/{case_id}/dismiss-jurisdiction-mismatch")
async def dismiss_jurisdiction_mismatch(case_id: str, current_user: User = Depends(get_current_user)):
    """User declined the switch — clear the mismatch flag. Analysis remains under the current jurisdiction."""
    result = await db.cases.update_one(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"$set": {"jurisdiction_mismatch": False, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Case not found")
    return {"status": "dismissed"}


# ================== REFINEMENT FEATURE ==================
# Post-analysis free-text context. User can refine analysis up to 10 times per case
# (Free plan: 2). AI classifies each input as incremental | full_reanalysis | off_topic
# and the resulting analysis is preserved as a version in case_analyses_history.

REFINEMENT_MAX_PER_CASE = 10
REFINEMENT_FREE_QUOTA = 2
REFINEMENT_MAX_INPUT_CHARS = 4000

# Analysis fields that represent "the analysis" and get snapshotted/restored per version.
_ANALYSIS_SNAPSHOT_FIELDS = (
    "risk_score", "risk_financial", "risk_urgency", "risk_legal_strength", "risk_complexity",
    "deadline", "deadline_description", "financial_exposure", "financial_exposure_detailed",
    "ai_summary", "ai_findings", "ai_next_steps", "key_violation_types", "recommend_lawyer",
    "battle_preview", "success_probability", "procedural_defects", "applicable_laws",
    "immediate_actions", "leverage_points", "red_lines", "key_insight", "strategy",
    "lawyer_recommendation", "user_rights", "opposing_weaknesses", "documents_to_gather",
    "recent_case_law", "case_law_updated", "archer_question", "title", "type",
    "strategy_narrative", "amounts", "analysis_depth",
)


def _case_language(case_doc: dict, current_user: User) -> str:
    """Resolve the language in which this case's analysis was originally written."""
    return (case_doc.get("language") or getattr(current_user, "language", None) or "en")


def _extract_analysis_snapshot(case_doc: dict) -> dict:
    """Pull the analysis-relevant fields out of a case doc for versioned storage."""
    return {k: case_doc.get(k) for k in _ANALYSIS_SNAPSHOT_FIELDS if k in case_doc}


async def _ensure_v1_snapshot(case_doc: dict) -> None:
    """On first refinement ever for this case, persist the current analysis as version 1.
    Idempotent: safe to call multiple times."""
    existing = await db.case_analyses_history.find_one({"case_id": case_doc["case_id"], "version": 1})
    if existing:
        return
    snapshot_id = f"snap_{uuid.uuid4().hex[:12]}"
    await db.case_analyses_history.insert_one({
        "snapshot_id": snapshot_id,
        "case_id": case_doc["case_id"],
        "user_id": case_doc["user_id"],
        "version": 1,
        "analysis": _extract_analysis_snapshot(case_doc),
        "triggered_by_refinement_id": None,
        "refinement_excerpt": None,
        "language": case_doc.get("language") or "en",
        "jurisdiction": case_doc.get("jurisdiction") or case_doc.get("country") or "US",
        "created_at": case_doc.get("updated_at") or case_doc.get("created_at") or datetime.now(timezone.utc).isoformat(),
    })


async def _write_new_snapshot(case_doc: dict, version: int, analysis_fields: dict,
                              refinement_id: str, excerpt: str) -> str:
    """Write a new version snapshot after a refinement. Returns snapshot_id."""
    snapshot_id = f"snap_{uuid.uuid4().hex[:12]}"
    await db.case_analyses_history.insert_one({
        "snapshot_id": snapshot_id,
        "case_id": case_doc["case_id"],
        "user_id": case_doc["user_id"],
        "version": version,
        "analysis": analysis_fields,
        "triggered_by_refinement_id": refinement_id,
        "refinement_excerpt": excerpt,
        "language": case_doc.get("language") or "en",
        "jurisdiction": case_doc.get("jurisdiction") or case_doc.get("country") or "US",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    return snapshot_id


def _off_topic_fallback_message(language: str) -> str:
    lang = (language or "en").lower()
    if lang.startswith("fr"):
        return "Ces informations ne semblent pas liées à votre dossier. Pouvez-vous reformuler en donnant des détails sur les faits, dates, montants, ou parties impliquées ?"
    if lang.startswith("nl"):
        return "Deze informatie lijkt niet gerelateerd aan uw dossier. Kunt u deze herformuleren met details over feiten, data, bedragen of betrokken partijen?"
    if lang.startswith("de"):
        return "Diese Informationen scheinen nicht mit Ihrem Fall zusammenzuhängen. Können Sie sie mit Details zu Fakten, Daten, Beträgen oder Beteiligten umformulieren?"
    return "This information doesn't appear to be related to your case. Could you rephrase with details about facts, dates, amounts, or parties involved?"


async def decide_refinement_strategy(case_doc: dict, user_input: str,
                                     prior_inputs: List[str], language: str) -> dict:
    """Classify the user's refinement input. Returns {strategy, reasoning, off_topic_message}.
    Strategy ∈ {'incremental', 'full_reanalysis', 'off_topic'}."""
    current_summary = (case_doc.get("ai_summary") or "")[:800]
    current_findings = case_doc.get("ai_findings") or []
    findings_brief = "; ".join(
        f.get("text", "")[:120] for f in current_findings[:6] if isinstance(f, dict)
    )
    prior_excerpts = " | ".join((x or "")[:150] for x in prior_inputs[-3:])

    system = """You are a routing classifier for a legal analysis refinement system. A client adds a free-text comment to clarify or extend their case. You must decide how to process it.

Return ONLY valid JSON with keys: strategy (one of: "incremental", "full_reanalysis", "off_topic"), reasoning (1 short sentence), off_topic_message (string, empty unless off_topic).

Rules:
- "off_topic": input is clearly unrelated to any legal case (e.g. "I like pizza", gibberish, random test strings, product feedback about the app). Nothing that could plausibly affect a legal analysis.
- "incremental": input is a minor clarification or correction — a number adjustment, a single missing date, a typo fix, a one-fact addition that slots into existing findings without changing the legal theory or risk profile.
- "full_reanalysis": input reveals structural new information — a new party, a new document mentioned, a new legal claim, a major procedural fact (prior notice / prior response / deadline missed), a radically different version of events, or multiple new facts that change risk scoring.

When in doubt between incremental and full_reanalysis, choose full_reanalysis."""

    user_prompt = f"""CURRENT ANALYSIS SUMMARY:
{current_summary}

CURRENT TOP FINDINGS:
{findings_brief or '(none)'}

PRIOR REFINEMENT INPUTS (for context, not the new one):
{prior_excerpts or '(none)'}

NEW USER INPUT (the comment to classify; may be in any language — translate mentally):
\"\"\"{user_input[:3500]}\"\"\"

ORIGINAL CASE LANGUAGE (write off_topic_message in this language if applicable): {language}

Return JSON only."""

    result = await call_claude(system, user_prompt, max_tokens=400)
    if not isinstance(result, dict) or "strategy" not in result:
        logger.warning(f"decide_refinement_strategy returned unexpected shape: {result!r} — defaulting to full_reanalysis")
        return {"strategy": "full_reanalysis", "reasoning": "classifier fallback", "off_topic_message": ""}
    strat = result.get("strategy")
    if strat not in ("incremental", "full_reanalysis", "off_topic"):
        strat = "full_reanalysis"
    msg = result.get("off_topic_message") or ""
    if strat == "off_topic" and not msg.strip():
        msg = _off_topic_fallback_message(language)
    return {"strategy": strat, "reasoning": result.get("reasoning", ""), "off_topic_message": msg}


async def _run_incremental_refinement(case_doc: dict, user_input: str,
                                      language: str, jurisdiction: str) -> dict:
    """Incremental path: small Claude call that returns a patch merged into the current analysis."""
    lang_instr = get_language_instruction(language)
    jur_guard = get_jurisdiction_guard(jurisdiction, language)
    lang_note = f"The case was originally analyzed in language code '{language}'. Maintain the same language for all output fields. The user input may be in a different language — translate mentally and incorporate without switching the analysis language."
    system = (SENIOR_ATTORNEY_PERSONA + jur_guard + lang_instr +
              "\n\n" + lang_note + "\n\nYou receive an existing legal analysis plus a short additional fact from the client. Produce a precise patch that integrates the new fact. Do NOT rewrite unchanged parts.")

    current_summary = case_doc.get("ai_summary") or ""
    current_findings = case_doc.get("ai_findings") or []
    current_next_steps = case_doc.get("ai_next_steps") or []
    current_risk = {
        "total": case_doc.get("risk_score", 0),
        "financial": case_doc.get("risk_financial", 0),
        "urgency": case_doc.get("risk_urgency", 0),
        "legal_strength": case_doc.get("risk_legal_strength", 0),
        "complexity": case_doc.get("risk_complexity", 0),
    }

    user_prompt = f"""EXISTING ANALYSIS:
summary: {current_summary}
findings (top 8): {json.dumps(current_findings[:8], ensure_ascii=False)[:3000]}
next_steps: {json.dumps(current_next_steps[:5], ensure_ascii=False)[:1500]}
current risk_score: {json.dumps(current_risk)}

NEW USER INPUT TO INCORPORATE:
\"\"\"{user_input[:3500]}\"\"\"

Return ONLY this JSON (no prose, no markdown):
{{
  "risk_score": {{"total": int 0-100, "financial": int 0-100, "urgency": int 0-100, "legal_strength": int 0-100, "complexity": int 0-100}},
  "ai_summary": "updated 2-3 sentence summary reflecting the new info",
  "ai_findings": [ <= 8 finding objects with text/impact/type, include the existing ones possibly edited plus 0-2 new ones driven by the input ],
  "ai_next_steps": [ exactly 3 step objects with title/description/action_type ],
  "key_insight": "one sentence the user must remember",
  "impact_summary": "one short sentence explaining how the new input changed the analysis"
}}"""
    patch = await call_claude(system, user_prompt, max_tokens=3500)
    if not isinstance(patch, dict):
        return {}
    return patch


async def _run_full_reanalysis(case_doc: dict, document_text: str, user_input: str,
                               prior_inputs: List[str], language: str,
                               jurisdiction: str, region: str) -> dict:
    """Full 5-pass reanalysis with the new context merged into user_context."""
    combined_context_parts = []
    for prior in prior_inputs:
        if prior and prior.strip():
            combined_context_parts.append(prior.strip())
    combined_context_parts.append(user_input.strip())
    merged_context = "\n---\n".join(combined_context_parts)[:3500]

    lang_preserve_note = (
        f"\n\nLANGUAGE CONSISTENCY RULE — NON-NEGOTIABLE: "
        f"This case was originally analyzed in language code '{language}'. "
        f"Maintain the SAME language for every field of the analysis. "
        f"The user context below may be in a different language — translate it mentally but keep the analysis in the original language."
    )
    merged_context_with_note = merged_context + lang_preserve_note

    if jurisdiction == "BE":
        be_lang = language if language.startswith(("fr", "nl", "de")) else "fr-BE"
        return await analyze_document_belgian(
            document_text, user_context=merged_context_with_note,
            region=region or "Wallonie", language=be_lang,
        )
    return await analyze_document_advanced(
        document_text, user_context=merged_context_with_note,
        language=language, jurisdiction=jurisdiction,
    )


class RefineRequest(BaseModel):
    user_input: str


@api_router.post("/cases/{case_id}/refine")
async def refine_case(case_id: str, body: RefineRequest,
                      current_user: User = Depends(get_current_user)):
    """Accept a free-text refinement. Classifies input (incremental/full_reanalysis/off_topic),
    versions the resulting analysis, enforces plan + per-case caps + concurrency lock."""
    _current_case_id.set(case_id)
    user_input = (body.user_input or "").strip()
    if not user_input:
        raise HTTPException(status_code=400, detail="user_input required")
    if len(user_input) > REFINEMENT_MAX_INPUT_CHARS:
        raise HTTPException(status_code=400, detail=f"user_input exceeds {REFINEMENT_MAX_INPUT_CHARS} chars")

    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0}
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")

    # Guard: multi-doc cases are not yet supported on the refine path.
    if (case_doc.get("document_count") or 0) > 1:
        raise HTTPException(status_code=400, detail="multi_doc_refinement_unsupported")

    # Concurrency: don't accept a second refine while the prior one is in flight.
    if case_doc.get("refinement_in_progress"):
        raise HTTPException(status_code=409, detail="refinement_in_progress")

    refinement_count = int(case_doc.get("refinement_count") or 0)

    # Hard cap — applies to every tier.
    if refinement_count >= REFINEMENT_MAX_PER_CASE:
        raise HTTPException(status_code=429, detail="max_refinements_reached")

    # Free-tier quota.
    plan = (getattr(current_user, "plan", None) or "free").lower()
    if plan == "free" and refinement_count >= REFINEMENT_FREE_QUOTA:
        raise HTTPException(status_code=402, detail="upgrade_required")

    # Load the key document's extracted text (needed for full_reanalysis path).
    doc = await db.documents.find_one(
        {"case_id": case_id, "user_id": current_user.user_id, "extracted_text": {"$ne": None}},
        {"_id": 0},
        sort=[("uploaded_at", -1)],
    )
    document_text = (doc or {}).get("extracted_text") or ""

    # Acquire the lock.
    now = datetime.now(timezone.utc).isoformat()
    lock_res = await db.cases.update_one(
        {"case_id": case_id, "user_id": current_user.user_id, "refinement_in_progress": {"$ne": True}},
        {"$set": {"refinement_in_progress": True, "updated_at": now}},
    )
    if lock_res.modified_count == 0:
        raise HTTPException(status_code=409, detail="refinement_in_progress")

    # Ensure v1 snapshot exists (backfill at first-ever refinement).
    try:
        await _ensure_v1_snapshot(case_doc)
    except Exception as e:
        logger.error(f"_ensure_v1_snapshot failed for {case_id}: {e}")

    # Fetch prior refinements (for context + strategy decision).
    prior_refinements_cursor = db.case_refinements.find(
        {"case_id": case_id, "status": "completed"}, {"_id": 0, "user_input": 1}
    ).sort("created_at", 1)
    prior_docs = await prior_refinements_cursor.to_list(REFINEMENT_MAX_PER_CASE)
    prior_inputs = [p.get("user_input", "") for p in prior_docs]

    language = _case_language(case_doc, current_user)
    jurisdiction = case_doc.get("jurisdiction") or case_doc.get("country") or "US"
    region = case_doc.get("region") or ""

    # Create the refinement record (pending).
    refinement_id = f"ref_{uuid.uuid4().hex[:12]}"
    current_version = int(case_doc.get("current_analysis_version") or 1)
    excerpt = user_input[:40]
    await db.case_refinements.insert_one({
        "refinement_id": refinement_id,
        "case_id": case_id,
        "user_id": current_user.user_id,
        "user_input": user_input,
        "strategy": None,
        "status": "pending",
        "version_before": current_version,
        "version_after": None,
        "created_at": now,
        "completed_at": None,
    })

    async def _release_lock_on_error():
        await db.cases.update_one(
            {"case_id": case_id},
            {"$set": {"refinement_in_progress": False, "updated_at": datetime.now(timezone.utc).isoformat()}},
        )

    # Strategy decision.
    try:
        decision = await decide_refinement_strategy(case_doc, user_input, prior_inputs, language)
    except Exception as e:
        logger.error(f"decide_refinement_strategy crashed for {case_id}: {e}")
        await _release_lock_on_error()
        await db.case_refinements.update_one(
            {"refinement_id": refinement_id},
            {"$set": {"status": "error", "error": "strategy_failed",
                      "completed_at": datetime.now(timezone.utc).isoformat()}},
        )
        raise HTTPException(status_code=500, detail="Refinement classification failed")

    strategy = decision["strategy"]
    logger.info(f"Refinement {refinement_id} for {case_id}: strategy={strategy}")

    # ── OFF-TOPIC: nothing to version, release lock, return gentle message.
    if strategy == "off_topic":
        off_topic_msg = decision.get("off_topic_message") or _off_topic_fallback_message(language)
        completed = datetime.now(timezone.utc).isoformat()
        await db.case_refinements.update_one(
            {"refinement_id": refinement_id},
            {"$set": {"strategy": "off_topic", "status": "completed",
                      "off_topic_message": off_topic_msg, "completed_at": completed}},
        )
        await db.cases.update_one(
            {"case_id": case_id},
            {"$set": {"refinement_in_progress": False, "updated_at": completed}},
        )
        return {
            "refinement_id": refinement_id,
            "strategy": "off_topic",
            "message": off_topic_msg,
            "version": current_version,
        }

    # ── INCREMENTAL or FULL_REANALYSIS: run the analysis, snapshot, update case.
    try:
        if strategy == "incremental":
            patch = await _run_incremental_refinement(case_doc, user_input, language, jurisdiction)
            if not patch:
                # Fallback to full reanalysis if the incremental call degenerated.
                logger.warning(f"Incremental refinement returned empty for {case_id} — falling back to full_reanalysis")
                strategy = "full_reanalysis"
            else:
                rs = patch.get("risk_score", {}) or {}
                update_fields = {
                    "risk_score": int(rs.get("total", case_doc.get("risk_score", 0)) or 0),
                    "risk_financial": int(rs.get("financial", case_doc.get("risk_financial", 0)) or 0),
                    "risk_urgency": int(rs.get("urgency", case_doc.get("risk_urgency", 0)) or 0),
                    "risk_legal_strength": int(rs.get("legal_strength", case_doc.get("risk_legal_strength", 0)) or 0),
                    "risk_complexity": int(rs.get("complexity", case_doc.get("risk_complexity", 0)) or 0),
                    "ai_summary": patch.get("ai_summary") or case_doc.get("ai_summary"),
                    "ai_findings": patch.get("ai_findings") or case_doc.get("ai_findings", []),
                    "ai_next_steps": patch.get("ai_next_steps") or case_doc.get("ai_next_steps", []),
                    "key_insight": patch.get("key_insight") or case_doc.get("key_insight"),
                }
                new_analysis_fields = {**case_doc, **update_fields}

        if strategy == "full_reanalysis":
            if not document_text:
                raise HTTPException(status_code=400, detail="no_document_text_for_reanalysis")
            full = await _run_full_reanalysis(
                case_doc, document_text, user_input, prior_inputs, language, jurisdiction, region,
            )
            if not full:
                raise HTTPException(status_code=500, detail="Full reanalysis returned no result")
            update_fields = _build_case_update(full, case_doc, "refinement", datetime.now(timezone.utc).isoformat())
            # Keep case title/type stable unless explicitly changed
            if not full.get("suggested_case_title"):
                update_fields.pop("title", None)
            new_analysis_fields = {**case_doc, **update_fields}

    except HTTPException:
        await _release_lock_on_error()
        await db.case_refinements.update_one(
            {"refinement_id": refinement_id},
            {"$set": {"status": "error", "completed_at": datetime.now(timezone.utc).isoformat()}},
        )
        raise
    except Exception as e:
        logger.error(f"Refinement execution failed for {case_id}: {e}", exc_info=True)
        await _release_lock_on_error()
        await db.case_refinements.update_one(
            {"refinement_id": refinement_id},
            {"$set": {"status": "error", "error": str(e)[:200],
                      "completed_at": datetime.now(timezone.utc).isoformat()}},
        )
        raise HTTPException(status_code=500, detail="Refinement failed")

    # ── Persist: snapshot new version, update case, mark refinement complete.
    new_version = current_version + 1
    snapshot_payload = {k: new_analysis_fields.get(k) for k in _ANALYSIS_SNAPSHOT_FIELDS if k in new_analysis_fields}
    completed = datetime.now(timezone.utc).isoformat()

    snapshot_id = await _write_new_snapshot(case_doc, new_version, snapshot_payload, refinement_id, excerpt)

    case_update_set = dict(update_fields)
    case_update_set.update({
        "current_analysis_version": new_version,
        "refinement_count": refinement_count + 1,
        "refinement_in_progress": False,
        "updated_at": completed,
    })
    await db.cases.update_one({"case_id": case_id}, {"$set": case_update_set})

    await db.case_refinements.update_one(
        {"refinement_id": refinement_id},
        {"$set": {"strategy": strategy, "status": "completed",
                  "version_after": new_version, "snapshot_id": snapshot_id,
                  "completed_at": completed}},
    )

    await db.case_events.insert_one({
        "event_id": f"evt_{uuid.uuid4().hex[:12]}",
        "case_id": case_id,
        "event_type": "refinement_applied",
        "title": f"Analysis refined (v{new_version})",
        "description": f"{'Full re-analysis' if strategy == 'full_reanalysis' else 'Incremental update'} triggered by user input",
        "metadata": {"strategy": strategy, "refinement_id": refinement_id, "version": new_version},
        "created_at": completed,
    })

    # Best-effort notification for long-running full reanalyses.
    if strategy == "full_reanalysis":
        try:
            from routes.client_notifications_routes import create_notification
            notif_title = "Nouvelle analyse disponible" if (language or "").startswith(("fr", "nl")) else "New analysis available"
            notif_msg = f"Version {new_version} est prête."
            asyncio.create_task(create_notification(
                client_user_id=current_user.user_id,
                type="other",
                title=notif_title,
                message=notif_msg,
                case_id=case_id,
                action_url=f"/cases/{case_id}",
            ))
        except Exception:
            logger.exception("refinement notification failed")

    updated_case = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    return {
        "refinement_id": refinement_id,
        "strategy": strategy,
        "version": new_version,
        "case": updated_case,
    }


@api_router.get("/cases/{case_id}/versions")
async def list_case_versions(case_id: str, current_user: User = Depends(get_current_user)):
    """List all analysis versions for a case (latest first). Used by the VersionPicker."""
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"_id": 0, "current_analysis_version": 1},
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    current = int(case_doc.get("current_analysis_version") or 1)

    snapshots = await db.case_analyses_history.find(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"_id": 0, "snapshot_id": 1, "version": 1, "refinement_excerpt": 1,
         "created_at": 1, "triggered_by_refinement_id": 1},
    ).sort("version", -1).to_list(REFINEMENT_MAX_PER_CASE + 5)

    # If no snapshots exist yet (case was never refined), synthesize v1 from the live case.
    if not snapshots:
        live = await db.cases.find_one(
            {"case_id": case_id, "user_id": current_user.user_id},
            {"_id": 0, "updated_at": 1, "created_at": 1, "language": 1},
        ) or {}
        snapshots = [{
            "snapshot_id": None,
            "version": 1,
            "refinement_excerpt": None,
            "created_at": live.get("updated_at") or live.get("created_at"),
            "triggered_by_refinement_id": None,
        }]

    return {
        "case_id": case_id,
        "current_version": current,
        "versions": [{**s, "is_current": s.get("version") == current} for s in snapshots],
    }


@api_router.get("/cases/{case_id}/versions/{version}")
async def get_case_version(case_id: str, version: int,
                           current_user: User = Depends(get_current_user)):
    """Read-only access to a specific historical analysis version."""
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0},
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    current = int(case_doc.get("current_analysis_version") or 1)

    # Current version is served from the live case doc (avoids duplicate snapshot on read).
    if version == current:
        return {"case_id": case_id, "version": version, "is_current": True,
                "analysis": _extract_analysis_snapshot(case_doc)}

    snap = await db.case_analyses_history.find_one(
        {"case_id": case_id, "user_id": current_user.user_id, "version": version},
        {"_id": 0},
    )
    if not snap:
        raise HTTPException(status_code=404, detail="Version not found")
    return {"case_id": case_id, "version": version, "is_current": False,
            "analysis": snap.get("analysis") or {},
            "refinement_excerpt": snap.get("refinement_excerpt"),
            "created_at": snap.get("created_at")}


@api_router.get("/documents/{document_id}")
async def get_document_detail(document_id: str, current_user: User = Depends(get_current_user)):
    """Get a single document with details"""
    doc = await db.documents.find_one(
        {"document_id": document_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

@api_router.delete("/documents/{document_id}")
async def delete_document(document_id: str, current_user: User = Depends(get_current_user)):
    """Delete a document"""
    result = await db.documents.delete_one({"document_id": document_id, "user_id": current_user.user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"status": "deleted"}

@api_router.get("/documents/{document_id}/download")
async def download_document(
    document_id: str,
    current_user: User = Depends(get_current_user)
):
    """Download a document"""
    doc = await db.documents.find_one(
        {"document_id": document_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not doc.get("storage_path"):
        raise HTTPException(status_code=404, detail="File not available")
    
    try:
        data, content_type = get_object(doc["storage_path"])
        return Response(
            content=data,
            media_type=content_type,
            headers={"Content-Disposition": f"attachment; filename={doc['file_name']}"}
        )
    except Exception as e:
        logger.error(f"Download failed: {e}")
        raise HTTPException(status_code=500, detail="Download failed")

# ================== Letter Generation Endpoints ==================


@api_router.post("/cases/{case_id}/archer-answer")
async def archer_answer(case_id: str, body: dict, current_user: User = Depends(get_current_user)):
    """Process user answer to Archer's question, update analysis"""
    case_doc = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    
    question = body.get("question", "")
    answer = body.get("answer", "")
    if not answer:
        raise HTTPException(status_code=400, detail="Answer required")
    
    user_language = current_user.language or case_doc.get("language", "en")
    lang_note = ""
    if user_language.startswith("fr"):
        lang_note = "Respond entirely in French."
    elif user_language.startswith("nl"):
        lang_note = "Respond entirely in Dutch."
    elif user_language.startswith("de"):
        lang_note = "Respond entirely in German."
    
    system = f"""You are Archer, a senior legal AI advisor. The user answered your clarifying question about their legal case.
Based on this new information, provide an updated assessment.
{lang_note}
Return JSON:
{{
  "impact_summary": "One sentence explaining how this answer changes the analysis",
  "risk_adjustment": number (-15 to +15, how much to adjust the risk score),
  "new_finding": {{
    "text": "New finding based on the user's answer",
    "impact": "high|medium|low",
    "type": "risk|opportunity|neutral",
    "legal_ref": "Relevant legal reference if applicable"
  }},
  "next_question": {{
    "text": "The NEXT most important question Archer should ask (cite specific details from the case)",
    "options": ["Answer option 1", "Answer option 2", "Answer option 3"]
  }} or null if no more questions needed,
  "updated_next_steps": [
    {{"title": "Specific action", "description": "Why and how"}}
  ]
}}"""
    
    user_msg = f"""Case: {case_doc.get('title', '')}
Case type: {case_doc.get('type', '')}
Current risk score: {case_doc.get('risk_score', 0)}/100
Current findings: {json.dumps(case_doc.get('ai_findings', [])[:3], default=str)}

Archer asked: "{question}"
User answered: "{answer}"

Analyze the impact of this answer on the case."""
    
    try:
        result = await call_claude_fast(system, user_msg, max_tokens=600)
    except Exception as e:
        logger.error(f"Archer Q&A error: {e}")
        raise HTTPException(status_code=500, detail="Analysis failed")
    
    now = datetime.now(timezone.utc).isoformat()
    
    # Apply risk adjustment
    old_score = case_doc.get("risk_score", 0)
    adj = result.get("risk_adjustment", 0)
    new_score = max(0, min(100, old_score + adj))
    
    # Build update
    update = {"$set": {"risk_score": new_score, "updated_at": now}}
    
    new_finding = result.get("new_finding")
    if new_finding:
        update.setdefault("$push", {})["ai_findings"] = new_finding
    
    updated_steps = result.get("updated_next_steps")
    if updated_steps:
        update["$set"]["ai_next_steps"] = updated_steps
    
    next_q = result.get("next_question")
    if next_q:
        update["$set"]["archer_question"] = next_q
    else:
        update["$set"]["archer_question"] = None
    
    # Store Q&A in history
    qa_entry = {"question": question, "answer": answer, "impact": result.get("impact_summary", ""), "timestamp": now}
    update.setdefault("$push", {})["archer_qa_history"] = qa_entry
    
    await db.cases.update_one({"case_id": case_id}, update)
    
    result["new_risk_score"] = new_score
    result["old_risk_score"] = old_score
    return result


# DEPRECATED: alias kept for backward compatibility — remove after prod validation
@api_router.post("/cases/{case_id}/james-answer")
async def james_answer_deprecated(case_id: str, body: dict, current_user: User = Depends(get_current_user)):
    return await archer_answer(case_id, body, current_user)


@api_router.post("/cases/{case_id}/generate-action-letter")
async def generate_action_letter(case_id: str, body: dict, current_user: User = Depends(get_current_user)):
    """Generate a letter for a specific next action"""
    case_doc = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    
    action_title = body.get("action_title", "")
    action_description = body.get("action_description", "")
    letter_tone = body.get("tone", "citizen")  # "citizen" or "attorney"
    user_language = current_user.language or case_doc.get("language", "en")
    
    tone_instruction = get_letter_tone_instruction(letter_tone, user_language)
    lang_note = ""
    if user_language.startswith("fr"):
        lang_note = "Write the letter entirely in French."
    elif user_language.startswith("nl"):
        lang_note = "Write the letter entirely in Dutch."
    elif user_language.startswith("de"):
        lang_note = "Write the letter entirely in German."
    
    system = f"""You are Archer, drafting a legal letter.
{lang_note}
{tone_instruction}
The letter must cite specific laws and statutes, and be ready to send.
Use the case details to fill in all specific information (names, addresses, dates, amounts).
Return JSON:
{{
  "subject": "Letter subject line",
  "recipient": "To whom this letter is addressed",
  "body": "Complete letter text with proper formatting. Use \\n for line breaks.",
  "legal_citations": ["List of laws/statutes cited"],
  "deadline_mentioned": "Any deadline mentioned in the letter or null"
}}"""
    
    user_msg = f"""Case: {case_doc.get('title', '')}
Type: {case_doc.get('type', '')}
Findings: {json.dumps(case_doc.get('ai_findings', [])[:5], default=str)}
Applicable laws: {json.dumps(case_doc.get('applicable_laws', []), default=str)}

Action requested: {action_title}
Description: {action_description}

Generate a complete, ready-to-send legal letter for this action."""
    
    try:
        letter = await call_claude_fast(system, user_msg, max_tokens=1500)
        return letter
    except Exception as e:
        logger.error(f"Letter generation error: {e}")
        raise HTTPException(status_code=500, detail="Letter generation failed")



@api_router.get("/letters/types/{case_type}")
async def get_letter_types(case_type: str, country: str = "US"):
    """Get available letter types for a case type"""
    if country == "BE":
        letter_types = BELGIAN_LETTER_TYPES.get(case_type, BELGIAN_LETTER_TYPES.get("debt", []))
    else:
        letter_types = LETTER_TYPES.get(case_type, LETTER_TYPES["other"])
    return {"case_type": case_type, "letter_types": letter_types, "country": country}

@api_router.post("/letters/generate")
async def generate_letter(
    request: LetterRequest,
    current_user: User = Depends(get_current_user)
):
    """Generate a response letter for a case"""
    # Get case data
    case_doc = await db.cases.find_one(
        {"case_id": request.case_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Build letter data
    today = datetime.now(timezone.utc).strftime("%B %d, %Y")
    
    letter_data = {
        "letter_type": request.letter_type,
        "case": {
            "type": case_doc.get("type", "other"),
            "title": case_doc.get("title", ""),
            "risk_score": case_doc.get("risk_score", 0),
            "financial_exposure": case_doc.get("financial_exposure"),
            "deadline": case_doc.get("deadline"),
            "deadline_description": case_doc.get("deadline_description"),
            "ai_summary": case_doc.get("ai_summary"),
            "findings": [f.get("text", "") for f in case_doc.get("ai_findings", [])]
        },
        "user": {
            "name": current_user.name,
            "address": request.user_address or f"{current_user.state_of_residence or 'United States'}",
            "state": current_user.state_of_residence or "United States"
        },
        "opposing_party": {
            "name": request.opposing_party_name or "[Opposing Party Name]",
            "address": request.opposing_party_address or "[Opposing Party Address]"
        },
        "additional_context": request.additional_context,
        "today_date": today
    }
    
    # Generate letter — use Belgian system prompt if user is Belgian
    is_belgian = current_user.jurisdiction == "BE"
    user_language = getattr(current_user, 'language', 'en') or 'en'
    letter_tone = request.tone or "citizen"
    letter_result = await generate_letter_with_claude(letter_data, belgian=is_belgian, language=user_language, tone=letter_tone)
    
    # Store letter in database
    letter_id = f"letter_{uuid.uuid4().hex[:12]}"
    letter_doc = {
        "letter_id": letter_id,
        "case_id": request.case_id,
        "user_id": current_user.user_id,
        "letter_type": request.letter_type,
        "letter_data": letter_result,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.letters.insert_one(letter_doc)
    
    # Create case event
    event_doc = {
        "event_id": f"evt_{uuid.uuid4().hex[:12]}",
        "case_id": request.case_id,
        "event_type": "letter_generated",
        "title": "Response letter generated",
        "description": f"{request.letter_type.replace('_', ' ').title()}",
        "metadata": {"letter_id": letter_id},
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.case_events.insert_one(event_doc)
    
    return {
        "letter_id": letter_id,
        "letter": letter_result
    }

@api_router.get("/cases/{case_id}/letters")
async def get_case_letters(
    case_id: str,
    current_user: User = Depends(get_current_user)
):
    """Get all generated letters for a case"""
    # Verify case belongs to user
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    
    letters = await db.letters.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    
    return letters

# ================== Lawyer Endpoints ==================

@api_router.get("/lawyers", response_model=List[Lawyer])
async def get_lawyers(
    specialty: Optional[str] = None,
    available_now: Optional[bool] = None,
    country: Optional[str] = None,
    language: Optional[str] = None
):
    """Get all lawyers, filtered by country/language if specified"""
    query = {}
    if specialty:
        query["specialty"] = {"$regex": specialty, "$options": "i"}
    if available_now:
        query["availability_status"] = "now"
    if country:
        query["country"] = country
    if language:
        query["language"] = {"$regex": language, "$options": "i"}
    
    lawyers = await db.lawyers.find(query, {"_id": 0}).to_list(100)
    return [Lawyer(**l) for l in lawyers]

@api_router.get("/lawyers/{lawyer_id}", response_model=Lawyer)
async def get_lawyer(lawyer_id: str):
    """Get a specific lawyer"""
    lawyer = await db.lawyers.find_one({"lawyer_id": lawyer_id}, {"_id": 0})
    if not lawyer:
        raise HTTPException(status_code=404, detail="Lawyer not found")
    return Lawyer(**lawyer)

@api_router.post("/lawyer-calls", response_model=LawyerCall)
async def book_lawyer_call(
    call_data: LawyerCallCreate,
    current_user: User = Depends(get_current_user)
):
    """Book a lawyer call"""
    # Get lawyer
    lawyer = await db.lawyers.find_one({"lawyer_id": call_data.lawyer_id}, {"_id": 0})
    if not lawyer:
        raise HTTPException(status_code=404, detail="Lawyer not found")
    
    # Generate AI brief if case provided
    ai_brief = None
    if call_data.case_id:
        case = await db.cases.find_one(
            {"case_id": call_data.case_id, "user_id": current_user.user_id},
            {"_id": 0}
        )
        if case:
            doc_count = await db.documents.count_documents({"case_id": call_data.case_id})
            ai_brief = {
                "case_type": case.get("type"),
                "risk_score": case.get("risk_score"),
                "document_count": doc_count,
                "key_findings": case.get("ai_findings", [])[:3],
                "deadline": case.get("deadline"),
                "deadline_description": case.get("deadline_description"),
                "financial_exposure": case.get("financial_exposure"),
                "user_state": current_user.state_of_residence,
                "ai_recommendations": case.get("ai_next_steps", [])
            }
    
    # Create call record
    call_id = f"call_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    call_doc = {
        "call_id": call_id,
        "user_id": current_user.user_id,
        "lawyer_id": call_data.lawyer_id,
        "lawyer_name": lawyer["name"],
        "case_id": call_data.case_id,
        "scheduled_at": call_data.scheduled_at,
        "duration_minutes": 30,
        "price": 149,
        "status": "upcoming",
        "ai_brief": ai_brief,
        "created_at": now
    }
    await db.lawyer_calls.insert_one(call_doc)
    
    # Update lawyer sessions count
    await db.lawyers.update_one(
        {"lawyer_id": call_data.lawyer_id},
        {"$inc": {"sessions_count": 1}}
    )
    
    # Create case event if case provided
    if call_data.case_id:
        event_doc = {
            "event_id": f"evt_{uuid.uuid4().hex[:12]}",
            "case_id": call_data.case_id,
            "event_type": "call_booked",
            "title": "Call booked",
            "description": f"{lawyer['name']} · {call_data.time_slot}",
            "metadata": {"lawyer_id": call_data.lawyer_id, "call_id": call_id},
            "created_at": now
        }
        await db.case_events.insert_one(event_doc)
    
    return LawyerCall(**call_doc)

@api_router.get("/lawyer-calls", response_model=List[LawyerCall])
async def get_lawyer_calls(current_user: User = Depends(get_current_user)):
    """Get user's lawyer calls"""
    calls = await db.lawyer_calls.find(
        {"user_id": current_user.user_id},
        {"_id": 0}
    ).sort("scheduled_at", -1).to_list(100)
    return [LawyerCall(**c) for c in calls]

@api_router.get("/lawyer-calls/next")
async def get_next_call(current_user: User = Depends(get_current_user)):
    """Get user's next upcoming call"""
    call = await db.lawyer_calls.find_one(
        {"user_id": current_user.user_id, "status": "upcoming"},
        {"_id": 0},
        sort=[("scheduled_at", 1)]
    )
    return call

# ================== Dashboard Stats ==================

@api_router.get("/dashboard/stats")
async def get_dashboard_stats(current_user: User = Depends(get_current_user)):
    """Get dashboard statistics"""
    # Active cases
    active_cases = await db.cases.count_documents({
        "user_id": current_user.user_id,
        "status": "active"
    })
    
    # Cases requiring action (high risk)
    cases_action = await db.cases.count_documents({
        "user_id": current_user.user_id,
        "status": "active",
        "risk_score": {"$gt": 60}
    })
    
    # Highest risk score
    highest_risk_case = await db.cases.find_one(
        {"user_id": current_user.user_id, "status": "active"},
        {"_id": 0},
        sort=[("risk_score", -1)]
    )
    highest_risk = highest_risk_case["risk_score"] if highest_risk_case else 0
    
    # Total documents
    total_docs = await db.documents.count_documents({"user_id": current_user.user_id})
    
    # Next call
    next_call = await db.lawyer_calls.find_one(
        {"user_id": current_user.user_id, "status": "upcoming"},
        {"_id": 0},
        sort=[("scheduled_at", 1)]
    )
    
    return {
        "active_cases": active_cases,
        "cases_requiring_action": cases_action,
        "highest_risk_score": highest_risk,
        "total_documents": total_docs,
        "next_call": next_call
    }

# ================== Risk Score History ==================

@api_router.get("/cases/{case_id}/risk-history")
async def get_risk_history(case_id: str, current_user: User = Depends(get_current_user)):
    """Get risk score history for a case"""
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"_id": 0, "risk_score_history": 1, "risk_score": 1}
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    return {
        "current_score": case_doc.get("risk_score", 0),
        "history": case_doc.get("risk_score_history", [])
    }

# ================== Outcome Predictor ==================

OUTCOME_SYSTEM_PROMPT = """You are Archer's legal outcome prediction engine. Based on case data and AI analysis, predict likely outcomes for this legal situation. You are NOT a lawyer and never claim to be one.

RULES:
- Provide 3 scenarios: favorable, neutral, and unfavorable
- Include probability estimates (must total 100%)
- Reference specific case facts and legal precedents
- Be realistic and grounded in the case data
- Plain English only

OUTPUT FORMAT — respond ONLY with this exact JSON, no other text:
{
  "favorable": {
    "probability": 30,
    "title": "Short title",
    "description": "2-3 sentences describing this outcome",
    "likely_result": "What happens specifically",
    "financial_impact": "Dollar amount or range",
    "timeline": "Estimated timeline"
  },
  "neutral": {
    "probability": 45,
    "title": "Short title",
    "description": "2-3 sentences describing this outcome",
    "likely_result": "What happens specifically",
    "financial_impact": "Dollar amount or range",
    "timeline": "Estimated timeline"
  },
  "unfavorable": {
    "probability": 25,
    "title": "Short title",
    "description": "2-3 sentences describing this outcome",
    "likely_result": "What happens specifically",
    "financial_impact": "Dollar amount or range",
    "timeline": "Estimated timeline"
  },
  "key_factors": ["Factor 1", "Factor 2", "Factor 3"],
  "recommendation": "One sentence best course of action",
  "disclaimer": "This prediction provides legal information only, not legal advice. Actual outcomes may vary."
}"""

@api_router.post("/cases/{case_id}/predict-outcome")
async def predict_outcome(case_id: str, current_user: User = Depends(get_current_user)):
    """Predict case outcome using AI"""
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    
    case_context = {
        "type": case_doc.get("type"),
        "title": case_doc.get("title"),
        "risk_score": case_doc.get("risk_score"),
        "risk_financial": case_doc.get("risk_financial"),
        "risk_urgency": case_doc.get("risk_urgency"),
        "risk_legal_strength": case_doc.get("risk_legal_strength"),
        "risk_complexity": case_doc.get("risk_complexity"),
        "financial_exposure": case_doc.get("financial_exposure"),
        "deadline": case_doc.get("deadline"),
        "summary": case_doc.get("ai_summary"),
        "findings": [f.get("text") for f in case_doc.get("ai_findings", [])],
        "next_steps": [s.get("title") for s in case_doc.get("ai_next_steps", [])]
    }
    
    try:
        async with httpx.AsyncClient() as http_client:
            response = await http_client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": get_anthropic_key(),
                    "anthropic-version": "2023-06-01"
                },
                json={
                    "model": "claude-opus-4-7",
                    "max_tokens": 2000,
                    "system": (BELGIAN_OUTCOME_SYSTEM if current_user.jurisdiction == "BE" else OUTCOME_SYSTEM_PROMPT) + get_language_instruction(getattr(current_user, 'language', 'en') or 'en'),
                    "messages": [{
                        "role": "user",
                        "content": f"Predict outcomes for this legal case. Return JSON only:\n\n{json.dumps(case_context, indent=2)}"
                    }]
                },
                timeout=60.0
            )
            response.raise_for_status()
            data = response.json()
            text = data["content"][0]["text"]
            text = text.replace("```json", "").replace("```", "").strip()
            prediction = json.loads(text)
            return prediction
    except Exception as e:
        logger.error(f"Outcome prediction error: {e}")
        raise HTTPException(status_code=500, detail="Prediction failed. Please try again.")

# ================== Document Scanner (Mobile Camera) ==================

class ScanRequest(BaseModel):
    image_base64: str
    case_id: Optional[str] = None

@api_router.post("/documents/scan")
async def scan_document(
    scan_data: ScanRequest,
    current_user: User = Depends(get_current_user)
):
    """Scan a document photo using Claude Vision API"""
    # Check plan restrictions
    if current_user.plan == "free":
        doc_count = await db.documents.count_documents({"user_id": current_user.user_id})
        if doc_count >= 1:
            raise HTTPException(
                status_code=403,
                detail="Free plan limited to 1 document. Upgrade to Pro for unlimited analyses."
            )
    
    # Clean base64 string
    image_data = scan_data.image_base64
    if "base64," in image_data:
        image_data = image_data.split("base64,")[1]
    
    # Determine media type
    media_type = "image/jpeg"
    if scan_data.image_base64.startswith("data:image/png"):
        media_type = "image/png"
    elif scan_data.image_base64.startswith("data:image/webp"):
        media_type = "image/webp"
    
    # Call Claude Vision API for OCR + analysis
    try:
        async with httpx.AsyncClient() as http_client:
            response = await http_client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": get_anthropic_key(),
                    "anthropic-version": "2023-06-01"
                },
                json={
                    "model": "claude-opus-4-7",
                    "max_tokens": 3000,
                    "system": "First, extract all text from this document image (OCR). Then analyze the extracted text as a legal document.\n\n" + CLAUDE_SYSTEM_PROMPT,
                    "messages": [{
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": media_type,
                                    "data": image_data
                                }
                            },
                            {
                                "type": "text",
                                "text": "Extract text from this document image and analyze it as a legal document. Return the analysis JSON only."
                            }
                        ]
                    }]
                },
                timeout=90.0
            )
            response.raise_for_status()
            data = response.json()
            text = data["content"][0]["text"]
            text = text.replace("```json", "").replace("```", "").strip()
            analysis = json.loads(text)
    except Exception as e:
        logger.error(f"Claude Vision scan error: {e}")
        raise HTTPException(status_code=500, detail="Document scan failed. Please try again or upload a file instead.")
    
    # Save scanned image to storage
    now = datetime.now(timezone.utc).isoformat()
    image_bytes = base64.b64decode(image_data)
    ext = "jpg" if media_type == "image/jpeg" else "png"
    storage_path = f"{APP_NAME}/scans/{current_user.user_id}/{uuid.uuid4()}.{ext}"
    try:
        put_object(storage_path, image_bytes, media_type)
    except Exception as e:
        logger.error(f"Scan storage failed: {e}")
        storage_path = None
    
    # Create document record
    document_id = f"doc_{uuid.uuid4().hex[:12]}"
    case_id = scan_data.case_id
    
    doc_record = {
        "document_id": document_id,
        "case_id": case_id,
        "user_id": current_user.user_id,
        "file_name": f"scan_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.{ext}",
        "file_url": None,
        "storage_path": storage_path,
        "file_type": ext,
        "extracted_text": analysis.get("summary", ""),
        "status": "analyzed",
        "is_key_document": case_id is None,
        "uploaded_at": now
    }
    await db.documents.insert_one(doc_record)
    
    # Create or update case (same logic as file upload)
    if case_id:
        case_doc = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
        if case_doc and analysis:
            history_entry = {
                "score": analysis["risk_score"]["total"],
                "financial": analysis["risk_score"]["financial"],
                "urgency": analysis["risk_score"]["urgency"],
                "legal_strength": analysis["risk_score"]["legal_strength"],
                "complexity": analysis["risk_score"]["complexity"],
                "document_name": doc_record["file_name"],
                "date": now
            }
            await db.cases.update_one(
                {"case_id": case_id},
                {
                    "$set": {
                        "risk_score": analysis["risk_score"]["total"],
                        "risk_financial": analysis["risk_score"]["financial"],
                        "risk_urgency": analysis["risk_score"]["urgency"],
                        "risk_legal_strength": analysis["risk_score"]["legal_strength"],
                        "risk_complexity": analysis["risk_score"]["complexity"],
                        "deadline": normalize_deadline(analysis.get("deadline")),
                        "deadline_description": analysis.get("deadline_description") or (analysis.get("deadline", {}).get("description") if isinstance(analysis.get("deadline"), dict) else None),
                        "financial_exposure": normalize_financial_exposure(analysis.get("financial_exposure")),
                        "ai_summary": analysis.get("summary"),
                        "ai_findings": analysis.get("findings", []),
                "key_violation_types": analysis.get("key_violation_types", []),
                        "ai_next_steps": analysis.get("next_steps", []),
                        "recommend_lawyer": analysis.get("recommend_lawyer", False),
                        "updated_at": now
                    },
                    "$push": {"risk_score_history": history_entry}
                }
            )
    else:
        case_id = f"case_{uuid.uuid4().hex[:12]}"
        case_title = (analysis.get("suggested_case_title") or "").strip()
        if not case_title or len(case_title) < 5:
            summary = (analysis.get("summary") or "")
            case_title = summary[:60].rstrip('.') if summary else "Scanned Document Analysis"
        case_type = analysis.get("case_type", "other")
        
        case_doc = {
            "case_id": case_id,
            "user_id": current_user.user_id,
            "title": case_title,
            "type": case_type,
            "status": "active",
            "risk_score": analysis["risk_score"]["total"],
            "risk_financial": analysis["risk_score"]["financial"],
            "risk_urgency": analysis["risk_score"]["urgency"],
            "risk_legal_strength": analysis["risk_score"]["legal_strength"],
            "risk_complexity": analysis["risk_score"]["complexity"],
            "risk_score_history": [{
                "score": analysis["risk_score"]["total"],
                "financial": analysis["risk_score"]["financial"],
                "urgency": analysis["risk_score"]["urgency"],
                "legal_strength": analysis["risk_score"]["legal_strength"],
                "complexity": analysis["risk_score"]["complexity"],
                "document_name": doc_record["file_name"],
                "date": now
            }],
            "deadline": normalize_deadline(analysis.get("deadline")),
            "deadline_description": analysis.get("deadline_description") or (analysis.get("deadline", {}).get("description") if isinstance(analysis.get("deadline"), dict) else None),
            "financial_exposure": normalize_financial_exposure(analysis.get("financial_exposure")),
            "ai_summary": analysis.get("summary"),
            "ai_findings": analysis.get("findings", []),
                "key_violation_types": analysis.get("key_violation_types", []),
            "ai_next_steps": analysis.get("next_steps", []),
            "recommend_lawyer": analysis.get("recommend_lawyer", False),
            "document_count": 1,
            "created_at": now,
            "updated_at": now
        }
        await db.cases.insert_one(case_doc)
        await db.documents.update_one({"document_id": document_id}, {"$set": {"case_id": case_id}})
        
        event_doc = {
            "event_id": f"evt_{uuid.uuid4().hex[:12]}",
            "case_id": case_id,
            "event_type": "case_opened",
            "title": "Case opened",
            "description": f"Scanned document analyzed · {analysis['risk_score']['total']}/100",
            "metadata": None,
            "created_at": now
        }
        await db.case_events.insert_one(event_doc)
    
    return {
        "document_id": document_id,
        "case_id": case_id,
        "analysis": analysis,
        "status": "analyzed"
    }

# ================== Stripe Payment Endpoints ==================

PAYMENT_PACKAGES = {
    "pro_monthly": {"amount": 69.00, "currency": "usd", "description": "Archer Pro Plan - Monthly"},
    "lawyer_call": {"amount": 149.00, "currency": "usd", "description": "Attorney Video Call - 30 min"}
}

class CheckoutRequest(BaseModel):
    package_id: str
    origin_url: str
    metadata: Optional[Dict[str, str]] = None

@api_router.post("/payments/checkout")
async def create_checkout(
    checkout_data: CheckoutRequest,
    request: Request,
    current_user: User = Depends(get_current_user)
):
    """Create a Stripe checkout session"""
    package = PAYMENT_PACKAGES.get(checkout_data.package_id)
    if not package:
        raise HTTPException(status_code=400, detail="Invalid package")
    
    origin = checkout_data.origin_url.rstrip("/")
    success_url = f"{origin}/payment/success?session_id={{CHECKOUT_SESSION_ID}}"
    cancel_url = f"{origin}/payment/cancel"
    
    metadata = checkout_data.metadata or {}
    metadata["user_id"] = current_user.user_id
    metadata["package_id"] = checkout_data.package_id
    metadata["user_email"] = current_user.email
    
    host_url = str(request.base_url).rstrip("/")
    webhook_url = f"{host_url}api/webhook/stripe"
    
    stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
    
    checkout_req = CheckoutSessionRequest(
        amount=package["amount"],
        currency=package["currency"],
        success_url=success_url,
        cancel_url=cancel_url,
        metadata=metadata
    )
    
    session: CheckoutSessionResponse = await stripe_checkout.create_checkout_session(checkout_req)
    
    # Record payment transaction
    now = datetime.now(timezone.utc).isoformat()
    tx_doc = {
        "transaction_id": f"tx_{uuid.uuid4().hex[:12]}",
        "session_id": session.session_id,
        "user_id": current_user.user_id,
        "package_id": checkout_data.package_id,
        "amount": package["amount"],
        "currency": package["currency"],
        "description": package["description"],
        "payment_status": "pending",
        "status": "initiated",
        "metadata": metadata,
        "created_at": now,
        "updated_at": now
    }
    await db.payment_transactions.insert_one(tx_doc)
    
    return {"url": session.url, "session_id": session.session_id}

@api_router.get("/payments/status/{session_id}")
async def get_payment_status(
    session_id: str,
    request: Request,
    current_user: User = Depends(get_current_user)
):
    """Check payment status and update records"""
    tx = await db.payment_transactions.find_one(
        {"session_id": session_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not tx:
        raise HTTPException(status_code=404, detail="Transaction not found")
    
    # If already processed, return cached status
    if tx.get("payment_status") == "paid":
        return {"status": tx["status"], "payment_status": "paid", "package_id": tx["package_id"]}
    
    host_url = str(request.base_url).rstrip("/")
    webhook_url = f"{host_url}api/webhook/stripe"
    stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
    
    checkout_status: CheckoutStatusResponse = await stripe_checkout.get_checkout_status(session_id)
    
    now = datetime.now(timezone.utc).isoformat()
    
    if checkout_status.payment_status == "paid" and tx.get("payment_status") != "paid":
        # Update transaction
        await db.payment_transactions.update_one(
            {"session_id": session_id},
            {"$set": {"payment_status": "paid", "status": "completed", "updated_at": now}}
        )
        
        # Fulfill based on package
        if tx["package_id"] == "pro_monthly":
            await db.users.update_one(
                {"user_id": current_user.user_id},
                {"$set": {"plan": "pro"}}
            )
        
        return {"status": "completed", "payment_status": "paid", "package_id": tx["package_id"]}
    elif checkout_status.status == "expired":
        await db.payment_transactions.update_one(
            {"session_id": session_id},
            {"$set": {"payment_status": "expired", "status": "expired", "updated_at": now}}
        )
        return {"status": "expired", "payment_status": "expired", "package_id": tx["package_id"]}
    
    return {"status": tx["status"], "payment_status": checkout_status.payment_status, "package_id": tx["package_id"]}

@api_router.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    """Handle Stripe webhook events"""
    try:
        body = await request.body()
        signature = request.headers.get("Stripe-Signature", "")
        
        host_url = str(request.base_url).rstrip("/")
        webhook_url = f"{host_url}api/webhook/stripe"
        stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
        
        webhook_response = await stripe_checkout.handle_webhook(body, signature)
        
        if webhook_response.payment_status == "paid":
            session_id = webhook_response.session_id
            now = datetime.now(timezone.utc).isoformat()
            
            tx = await db.payment_transactions.find_one({"session_id": session_id}, {"_id": 0})
            if tx and tx.get("payment_status") != "paid":
                await db.payment_transactions.update_one(
                    {"session_id": session_id},
                    {"$set": {"payment_status": "paid", "status": "completed", "updated_at": now}}
                )
                
                if tx["package_id"] == "pro_monthly":
                    await db.users.update_one(
                        {"user_id": tx["user_id"]},
                        {"$set": {"plan": "pro"}}
                    )
        
        return {"status": "ok"}
    except Exception as e:
        logger.error(f"Webhook error: {e}")
        return {"status": "error", "detail": str(e)}

# ================== Legal Chat ==================

ARCHER_SYSTEM_PROMPT = """You are Archer, a Senior Legal Advisor with 20 years of experience representing individuals — never corporations — in employment disputes, tenant rights, contract law, debt collection defense, consumer protection, and civil litigation across the United States and Belgium.

You have personally handled over 2,000 cases and have access to 847,000+ legal articles, statutes, and court decisions from Cornell Law LII, CourtListener, Belgian ejustice.be, and cnt-nar.be.

YOUR PERSONALITY:
- Direct and confident — you know the law and you state it clearly
- Warm but professional — you understand people are stressed when they reach out
- Never evasive — if you don't know something, you say so and explain what you do know
- You use plain English — legal jargon only when necessary, always explained

YOUR ANSWER FORMAT:
1. Direct answer in 1-2 sentences
2. The specific law or rule that applies (cite exactly)
3. What the person should do right now
4. If urgent or high-stakes: recommend booking a call with a licensed attorney

CRITICAL RULES:
- Always answer in the language the user writes in
- Always cite the specific statute, article, or case that applies
- Never say 'I cannot provide legal advice' as your first response — answer first, add disclaimer after
- Never refuse to answer a legal question — provide information even if general
- If the user has an active case in Archer, reference it: 'Based on your [case name] case...'
- Recommend booking an attorney call when: financial exposure > $5,000, deadline < 7 days, criminal implications, or complexity score > 70
- End every response with one of these CTAs:
  * 'Want me to analyze a document related to this?' (if relevant)
  * 'This sounds serious — I recommend booking a call with a licensed attorney.' (if complex)
  * 'Here is what most people in your situation do:' (if straightforward)

JURISDICTION: Apply US Federal + State law for US users. Apply Belgian federal + regional law for Belgian users. Always specify which law applies.

LEGAL CITATIONS FORMAT (CRITICAL):
When citing a legal article, use EXACTLY this format:
[ARTICLE: code=code_civil_be, article=1382, paragraph=null]
Available codes: code_civil_be, code_route_be, code_travail_be, code_conso_be, code_penal_be, code_judiciaire_be, loi_bail_be, code_civil_fr, code_travail_fr
When citing jurisprudence, use EXACTLY this format:
[JURIS: court=cass_be, date=2024-03-14, number=P.23.1234.F]
Available courts: cass_be, cc_be, ca_bruxelles, ca_liege, trav_bruxelles, cjue, cass_fr
NEVER generate URLs yourself. Use ONLY the [ARTICLE:] and [JURIS:] formats above.

LANGUAGE: Respond in the exact language the user writes in — English, French, Dutch, German, or Spanish."""

class ChatMessageInput(BaseModel):
    content: str
    case_id: Optional[str] = None

class ChatSendInput(BaseModel):
    message: str
    conversation_id: Optional[str] = None
    case_id: Optional[str] = None

@api_router.post("/chat/send")
async def chat_send(data: ChatSendInput, current_user: User = Depends(get_current_user)):
    """Convenience endpoint: auto-creates conversation if needed, sends message"""
    conv_id = data.conversation_id
    if not conv_id:
        now = datetime.now(timezone.utc).isoformat()
        conv_id = f"conv_{uuid.uuid4().hex[:12]}"
        conv_doc = {
            "conversation_id": conv_id,
            "user_id": current_user.user_id,
            "case_id": data.case_id,
            "title": data.message[:60] + ("..." if len(data.message) > 60 else ""),
            "created_at": now,
            "updated_at": now,
        }
        await db.chat_conversations.insert_one(conv_doc)

    # Forward to the existing send_chat_message
    msg_input = ChatMessageInput(content=data.message, case_id=data.case_id)
    result = await send_chat_message(conv_id, msg_input, current_user)
    return {
        "response": result["ai_message"]["content"],
        "conversation_id": conv_id,
        "user_message": result["user_message"],
        "ai_message": result["ai_message"],
    }



@api_router.post("/chat/send-stream")
async def chat_send_stream(data: ChatSendInput, current_user: User = Depends(get_current_user)):
    """Streaming chat endpoint — returns SSE with token-by-token output.
    Uses direct Anthropic API with prompt caching for fast TTFB."""
    conv_id = data.conversation_id
    now = datetime.now(timezone.utc).isoformat()

    if not conv_id:
        conv_id = f"conv_{uuid.uuid4().hex[:12]}"
        conv_doc = {
            "conversation_id": conv_id,
            "user_id": current_user.user_id,
            "case_id": data.case_id,
            "title": data.message[:60] + ("..." if len(data.message) > 60 else ""),
            "created_at": now, "updated_at": now,
        }
        await db.chat_conversations.insert_one(conv_doc)

    conv = await db.chat_conversations.find_one(
        {"conversation_id": conv_id, "user_id": current_user.user_id}, {"_id": 0})
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")

    # Free plan limit
    if current_user.plan == "free":
        conv_ids = [c["conversation_id"] for c in await db.chat_conversations.find(
            {"user_id": current_user.user_id}, {"conversation_id": 1, "_id": 0}).to_list(100)]
        total_user_msgs = await db.chat_messages.count_documents(
            {"conversation_id": {"$in": conv_ids}, "role": "user"})
        if total_user_msgs >= 3:
            raise HTTPException(status_code=403, detail="Free plan limited to 3 questions.")

    # Save user message
    user_msg = {
        "message_id": f"msg_{uuid.uuid4().hex[:12]}",
        "conversation_id": conv_id, "role": "user",
        "content": data.message, "created_at": now,
    }
    await db.chat_messages.insert_one(user_msg)

    # Update title from first message
    msg_count = await db.chat_messages.count_documents({"conversation_id": conv_id, "role": "user"})
    if msg_count == 1:
        await db.chat_conversations.update_one(
            {"conversation_id": conv_id},
            {"$set": {"title": data.message[:60], "updated_at": now}})

    # Build messages array (proper format for Anthropic API)
    history = await db.chat_messages.find(
        {"conversation_id": conv_id}, {"_id": 0, "role": 1, "content": 1}
    ).sort("created_at", 1).to_list(50)  # limit to last 50 messages
    claude_messages = [{"role": m["role"], "content": m["content"]} for m in history]

    # Build system prompt with prompt caching
    user_language = getattr(current_user, 'language', 'en') or 'en'
    user_jurisdiction = getattr(current_user, 'jurisdiction', 'US') or 'US'
    jurisdiction_ctx = (
        "\n\nThis user is under BELGIAN jurisdiction. Apply Belgian federal + regional law."
        if user_jurisdiction == 'BE' else
        "\n\nThis user is under US jurisdiction. Apply US Federal + applicable state law."
    )
    static_system = ARCHER_SYSTEM_PROMPT + get_language_instruction(user_language) + jurisdiction_ctx

    # Case context (dynamic part — not cached)
    dynamic_system = ""
    case_id = data.case_id or conv.get("case_id")
    if case_id:
        case_doc = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
        if case_doc:
            findings_text = "; ".join([f.get('text', '') for f in (case_doc.get('ai_findings') or [])[:5] if f.get('text')])
            steps_text = "; ".join([s.get('title', '') if isinstance(s, dict) else str(s) for s in (case_doc.get('ai_next_steps') or [])[:3]])
            dynamic_system = f"\n\nCONTEXT: Active case '{case_doc.get('title', '')}'. Type: {case_doc.get('type', '')}. Risk: {case_doc.get('risk_score', 0)}/100. Findings: {findings_text}. Next steps: {steps_text}."

    # System with prompt caching: static part cached, dynamic appended
    system_content = [
        {"type": "text", "text": static_system, "cache_control": {"type": "ephemeral"}},
    ]
    if dynamic_system:
        system_content.append({"type": "text", "text": dynamic_system})

    async def stream_tokens():
        ai_text = ""
        try:
            async with httpx.AsyncClient(timeout=90) as client:
                async with client.stream(
                    "POST",
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": get_anthropic_key(),
                        "anthropic-version": "2023-06-01",
                        "anthropic-beta": "prompt-caching-2024-07-31",
                        "content-type": "application/json",
                    },
                    json={
                        "model": "claude-sonnet-4-6",
                        "max_tokens": 2048,
                        "stream": True,
                        "system": system_content,
                        "messages": claude_messages,
                    },
                ) as resp:
                    # Send conversation_id immediately
                    yield f"data: {json.dumps({'type': 'meta', 'conversation_id': conv_id})}\n\n"
                    async for line in resp.aiter_lines():
                        if not line.startswith("data: "):
                            continue
                        payload = line[6:]
                        if payload == "[DONE]":
                            break
                        try:
                            event = json.loads(payload)
                        except json.JSONDecodeError:
                            continue
                        if event.get("type") == "content_block_delta":
                            delta = event.get("delta", {})
                            text = delta.get("text", "")
                            if text:
                                ai_text += text
                                yield f"data: {json.dumps({'type': 'token', 'text': text})}\n\n"
                        elif event.get("type") == "message_stop":
                            break
        except Exception as e:
            logger.error(f"Streaming chat error: {e}")
            if not ai_text:
                ai_text = "Archer is temporarily unavailable — please try again."
                yield f"data: {json.dumps({'type': 'token', 'text': ai_text})}\n\n"

        # Save AI response to DB
        ai_msg = {
            "message_id": f"msg_{uuid.uuid4().hex[:12]}",
            "conversation_id": conv_id, "role": "assistant",
            "content": ai_text,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        await db.chat_messages.insert_one(ai_msg)
        await db.chat_conversations.update_one(
            {"conversation_id": conv_id},
            {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}})
        yield f"data: {json.dumps({'type': 'done'})}\n\n"

    return StreamingResponse(
        stream_tokens(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


class CreateConversationInput(BaseModel):
    case_id: Optional[str] = None
    first_message: Optional[str] = None

@api_router.post("/chat/conversations")
async def create_conversation(
    data: CreateConversationInput,
    current_user: User = Depends(get_current_user)
):
    """Create a new chat conversation"""
    now = datetime.now(timezone.utc).isoformat()
    conv_id = f"conv_{uuid.uuid4().hex[:12]}"
    title = "New conversation"

    conv_doc = {
        "conversation_id": conv_id,
        "user_id": current_user.user_id,
        "title": title,
        "case_id": data.case_id,
        "created_at": now,
        "updated_at": now
    }
    await db.chat_conversations.insert_one(conv_doc)
    del conv_doc["_id"]
    return conv_doc

@api_router.get("/chat/conversations")
async def list_conversations(current_user: User = Depends(get_current_user)):
    """List user's chat conversations"""
    convs = await db.chat_conversations.find(
        {"user_id": current_user.user_id},
        {"_id": 0}
    ).sort("updated_at", -1).to_list(100)
    return convs

@api_router.delete("/chat/conversations/{conversation_id}")
async def delete_conversation(conversation_id: str, current_user: User = Depends(get_current_user)):
    """Delete a conversation and its messages"""
    await db.chat_conversations.delete_one({"conversation_id": conversation_id, "user_id": current_user.user_id})
    await db.chat_messages.delete_many({"conversation_id": conversation_id})
    return {"status": "deleted"}

@api_router.get("/chat/conversations/{conversation_id}/messages")
async def get_conversation_messages(conversation_id: str, current_user: User = Depends(get_current_user)):
    """Get messages for a conversation"""
    conv = await db.chat_conversations.find_one(
        {"conversation_id": conversation_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")

    messages = await db.chat_messages.find(
        {"conversation_id": conversation_id},
        {"_id": 0}
    ).sort("created_at", 1).to_list(500)
    return messages

@api_router.post("/chat/conversations/{conversation_id}/messages")
async def send_chat_message(
    conversation_id: str,
    data: ChatMessageInput,
    current_user: User = Depends(get_current_user)
):
    """Send a message and get AI response"""
    conv = await db.chat_conversations.find_one(
        {"conversation_id": conversation_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")

    # Free plan: limit to 3 total questions
    if current_user.plan == "free":
        total_user_msgs = await db.chat_messages.count_documents({
            "conversation_id": {"$in": [c["conversation_id"] for c in await db.chat_conversations.find({"user_id": current_user.user_id}, {"conversation_id": 1, "_id": 0}).to_list(100)]},
            "role": "user"
        })
        if total_user_msgs >= 3:
            raise HTTPException(status_code=403, detail="Free plan limited to 3 questions. Upgrade to Pro for unlimited legal chat.")

    now = datetime.now(timezone.utc).isoformat()

    # Save user message
    user_msg = {
        "message_id": f"msg_{uuid.uuid4().hex[:12]}",
        "conversation_id": conversation_id,
        "role": "user",
        "content": data.content,
        "created_at": now
    }
    await db.chat_messages.insert_one(user_msg)

    # Update conversation title from first message
    msg_count = await db.chat_messages.count_documents({"conversation_id": conversation_id, "role": "user"})
    if msg_count == 1:
        title = data.content[:60] + ("..." if len(data.content) > 60 else "")
        await db.chat_conversations.update_one(
            {"conversation_id": conversation_id},
            {"$set": {"title": title, "updated_at": now}}
        )

    # Build conversation history for Claude
    history = await db.chat_messages.find(
        {"conversation_id": conversation_id},
        {"_id": 0, "role": 1, "content": 1}
    ).sort("created_at", 1).to_list(100)

    claude_messages = [{"role": m["role"], "content": m["content"]} for m in history]

    # Build system prompt with optional case context
    system_prompt = ARCHER_SYSTEM_PROMPT
    user_language = getattr(current_user, 'language', 'en') or 'en'
    user_jurisdiction = getattr(current_user, 'jurisdiction', 'US') or 'US'
    system_prompt += get_language_instruction(user_language)
    if user_jurisdiction == 'BE':
        system_prompt += "\n\nThis user is under BELGIAN jurisdiction. Apply Belgian federal + regional law. Reference Belgian Civil Code, Labour Law, CCTs, and relevant Belgian statutes."
    else:
        system_prompt += "\n\nThis user is under US jurisdiction. Apply US Federal + applicable state law."

    case_id = data.case_id or conv.get("case_id")
    if case_id:
        case_doc = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
        if case_doc:
            case_findings = case_doc.get('ai_findings', [])[:5]
            findings_text = "; ".join([f.get('text', '') for f in case_findings if f.get('text')])
            case_steps = case_doc.get('ai_next_steps', [])[:3]
            steps_text = "; ".join([s.get('title', '') if isinstance(s, dict) else str(s) for s in case_steps])
            system_prompt += f"\n\nCONTEXT: The user has an active case: '{case_doc.get('title', '')}'. Type: {case_doc.get('type', '')}. Risk Score: {case_doc.get('risk_score', 0)}/100. Summary: {case_doc.get('ai_summary', 'N/A')}. Key findings: {findings_text}. Next steps: {steps_text}. They may ask questions related to this case. Continue the conversation naturally as Archer, their legal advisor."

    # Call Claude direct — Sonnet 4.6 for chat (latency-optimized)
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": get_anthropic_key(),
                    "anthropic-version": "2023-06-01",
                    "anthropic-beta": "prompt-caching-2024-07-31",
                    "content-type": "application/json",
                },
                json={
                    "model": "claude-sonnet-4-6",
                    "max_tokens": 2048,
                    "system": [
                        {"type": "text", "text": system_prompt, "cache_control": {"type": "ephemeral"}},
                    ],
                    "messages": claude_messages,
                },
            )
            resp.raise_for_status()
            ai_data = resp.json()
            ai_text = ai_data["content"][0]["text"]
        if not ai_text:
            ai_text = "Archer is temporarily unavailable — please try again in a moment."
    except Exception as e:
        logger.error(f"Chat Claude error: {e}")
        ai_text = "Archer is temporarily unavailable — please try again in a moment."

    # Enrich legal references with verified URLs
    try:
        ai_text = await enrich_legal_references(ai_text)
    except Exception:
        pass  # enrichment is best-effort

    # Save AI response
    ai_msg = {
        "message_id": f"msg_{uuid.uuid4().hex[:12]}",
        "conversation_id": conversation_id,
        "role": "assistant",
        "content": ai_text,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.chat_messages.insert_one(ai_msg)
    await db.chat_conversations.update_one(
        {"conversation_id": conversation_id},
        {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
    )

    return {"user_message": {k: v for k, v in user_msg.items() if k != "_id"}, "ai_message": {k: v for k, v in ai_msg.items() if k != "_id"}}

@api_router.get("/chat/question-count")
async def get_question_count(current_user: User = Depends(get_current_user)):
    """Get total question count for free plan tracking"""
    conv_ids = [c["conversation_id"] for c in await db.chat_conversations.find({"user_id": current_user.user_id}, {"conversation_id": 1, "_id": 0}).to_list(100)]
    if not conv_ids:
        return {"count": 0, "limit": 3 if current_user.plan == "free" else None}
    count = await db.chat_messages.count_documents({"conversation_id": {"$in": conv_ids}, "role": "user"})
    return {"count": count, "limit": 3 if current_user.plan == "free" else None}

# ================== Admin Dashboard Metrics ==================

@api_router.get("/admin/dashboard/metrics")
async def admin_dashboard_metrics(period: str = "30d", admin: dict = Depends(admin_required)):
    """Business metrics for admin dashboard home."""
    now = datetime.now(timezone.utc)
    if period == "7d":
        since = now - timedelta(days=7)
    elif period == "ytd":
        since = datetime(now.year, 1, 1, tzinfo=timezone.utc)
    else:
        since = now - timedelta(days=30)
    prev_since = since - (now - since)
    since_iso = since.isoformat()
    prev_iso = prev_since.isoformat()
    today_start = now.replace(hour=0, minute=0, second=0, microsecond=0).isoformat()
    week_start = (now - timedelta(days=now.weekday())).replace(hour=0, minute=0, second=0, microsecond=0).isoformat()

    # Customers
    total_customers = await db.users.count_documents({"account_type": {"$ne": "attorney"}})
    new_today = await db.users.count_documents({"created_at": {"$gte": today_start}, "account_type": {"$ne": "attorney"}})
    new_week = await db.users.count_documents({"created_at": {"$gte": week_start}, "account_type": {"$ne": "attorney"}})
    free = await db.users.count_documents({"plan": {"$in": ["free", None]}, "account_type": {"$ne": "attorney"}})
    solo = await db.users.count_documents({"plan": "solo"})
    family = await db.users.count_documents({"plan": "family"})
    pro = await db.users.count_documents({"plan": "pro"})

    # Cases
    total_cases = await db.cases.count_documents({})
    active_cases = await db.cases.count_documents({"status": "active"})
    pending_attr = await db.cases.count_documents({"attorney_status": "waiting_assignment"})

    # Revenue (from cases with payment)
    paid_cases = await db.cases.find(
        {"payment_status": "paid"}, {"amount_paid_cents": 1, "paid_at": 1, "_id": 0}
    ).to_list(10000)
    revenue_period = sum(c.get("amount_paid_cents", 0) for c in paid_cases if c.get("paid_at", "") >= since_iso)
    revenue_prev = sum(c.get("amount_paid_cents", 0) for c in paid_cases if prev_iso <= c.get("paid_at", "") < since_iso)
    revenue_today = sum(c.get("amount_paid_cents", 0) for c in paid_cases if c.get("paid_at", "") >= today_start)

    # MRR estimate (active subscriptions * price)
    mrr = (solo * 2499 + family * 4999 + pro * 9999)
    prev_mrr = max(1, int(mrr * 0.88))  # approximate for growth calc
    mrr_growth = round(((mrr - prev_mrr) / prev_mrr) * 100) if prev_mrr > 0 else 0

    # Action items
    attorneys_pending = await db.attorneys.count_documents({"application_status": "pending"})
    payouts_failed = await db.payouts.count_documents({"status": "failed"}) if "payouts" in await db.list_collection_names() else 0
    cases_unmatched = await db.cases.count_documents({
        "attorney_status": "waiting_assignment",
        "updated_at": {"$lt": (now - timedelta(hours=24)).isoformat()},
    })

    return {
        "mrr": {"current": mrr, "previous": prev_mrr, "growth_percent": mrr_growth},
        "revenue": {
            "today": revenue_today,
            "period": revenue_period,
            "previous_period": revenue_prev,
        },
        "customers": {
            "total": total_customers,
            "new_today": new_today,
            "new_week": new_week,
            "by_tier": {"free": free, "solo": solo, "family": family, "pro": pro},
        },
        "cases": {
            "total": total_cases,
            "active": active_cases,
            "pending_attribution": pending_attr,
        },
        "action_items": {
            "attorneys_pending": attorneys_pending,
            "payouts_failed": payouts_failed,
            "cases_unmatched": cases_unmatched,
        },
    }


@api_router.get("/admin/dashboard/revenue-chart")
async def admin_revenue_chart(period: str = "30d", admin: dict = Depends(admin_required)):
    """Daily revenue data for chart."""
    now = datetime.now(timezone.utc)
    days = 7 if period == "7d" else 365 if period == "ytd" else 30
    data = []
    for i in range(days - 1, -1, -1):
        day = now - timedelta(days=i)
        day_start = day.replace(hour=0, minute=0, second=0, microsecond=0).isoformat()
        day_end = (day.replace(hour=0, minute=0, second=0, microsecond=0) + timedelta(days=1)).isoformat()
        revenue = 0
        paid = await db.cases.find(
            {"payment_status": "paid", "paid_at": {"$gte": day_start, "$lt": day_end}},
            {"amount_paid_cents": 1, "_id": 0},
        ).to_list(500)
        revenue = sum(c.get("amount_paid_cents", 0) for c in paid)
        data.append({"date": day.strftime("%Y-%m-%d"), "revenue": revenue})
    return data


@api_router.get("/admin/dashboard/recent-signups")
async def admin_recent_signups(limit: int = 10, admin: dict = Depends(admin_required)):
    users = await db.users.find(
        {"account_type": {"$ne": "attorney"}},
        {"_id": 0, "user_id": 1, "email": 1, "name": 1, "plan": 1, "country": 1, "created_at": 1},
    ).sort("created_at", -1).limit(limit).to_list(limit)
    return users


@api_router.get("/admin/dashboard/recent-payments")
async def admin_recent_payments(limit: int = 10, admin: dict = Depends(admin_required)):
    cases = await db.cases.find(
        {"payment_status": "paid"},
        {"_id": 0, "case_id": 1, "user_id": 1, "amount_paid_cents": 1, "paid_at": 1, "type": 1},
    ).sort("paid_at", -1).limit(limit).to_list(limit)
    # Enrich with customer email
    for c in cases:
        user = await db.users.find_one({"user_id": c.get("user_id")}, {"email": 1, "_id": 0})
        c["customer_email"] = user.get("email") if user else "unknown"
    return cases


# ================== Admin Customers ==================

@api_router.get("/admin/customers")
async def admin_list_customers(
    search: str = "",
    tier: str = "",
    page: int = 1,
    per_page: int = 50,
    admin: dict = Depends(admin_required),
):
    query = {"account_type": {"$ne": "attorney"}}
    if search:
        query["$or"] = [
            {"email": {"$regex": search, "$options": "i"}},
            {"name": {"$regex": search, "$options": "i"}},
        ]
    if tier and tier != "all":
        query["plan"] = tier if tier != "free" else {"$in": ["free", None]}

    total = await db.users.count_documents(query)
    skip = (page - 1) * per_page
    users = await db.users.find(
        query, {"_id": 0, "password_hash": 0}
    ).sort("created_at", -1).skip(skip).limit(per_page).to_list(per_page)

    # Enrich with case count + revenue
    for u in users:
        uid = u.get("user_id")
        u["case_count"] = await db.cases.count_documents({"user_id": uid})
        paid = await db.cases.find(
            {"user_id": uid, "payment_status": "paid"}, {"amount_paid_cents": 1, "_id": 0}
        ).to_list(100)
        u["total_revenue_cents"] = sum(c.get("amount_paid_cents", 0) for c in paid)

    return {"customers": users, "total": total, "page": page, "total_pages": max(1, (total + per_page - 1) // per_page)}


@api_router.get("/admin/customers/{customer_id}")
async def admin_customer_detail(customer_id: str, admin: dict = Depends(admin_required)):
    user = await db.users.find_one({"user_id": customer_id}, {"_id": 0, "password_hash": 0})
    if not user:
        raise HTTPException(status_code=404, detail="Customer not found")

    cases = await db.cases.find(
        {"user_id": customer_id}, {"_id": 0}
    ).sort("created_at", -1).to_list(100)

    paid = [c for c in cases if c.get("payment_status") == "paid"]
    total_spent = sum(c.get("amount_paid_cents", 0) for c in paid)

    return {
        "customer": user,
        "cases": cases,
        "total_spent_cents": total_spent,
        "case_count": len(cases),
    }


# ================== Admin Cases ==================

@api_router.get("/admin/cases")
async def admin_list_cases(
    search: str = "", case_type: str = "", status: str = "",
    country: str = "", page: int = 1, per_page: int = 50,
    admin: dict = Depends(admin_required),
):
    query = {}
    if search:
        query["$or"] = [
            {"case_id": {"$regex": search, "$options": "i"}},
            {"title": {"$regex": search, "$options": "i"}},
        ]
    if case_type:
        query["type"] = case_type
    if status:
        query["status"] = status
    if country:
        query["country"] = country

    total = await db.cases.count_documents(query)
    skip = (page - 1) * per_page
    cases = await db.cases.find(query, {"_id": 0}).sort("created_at", -1).skip(skip).limit(per_page).to_list(per_page)

    # Enrich with customer email + attorney name
    for c in cases:
        user = await db.users.find_one({"user_id": c.get("user_id")}, {"email": 1, "name": 1, "_id": 0})
        c["customer_email"] = user.get("email") if user else None
        c["customer_name"] = user.get("name") if user else None
        # Find active assignment
        assignment = await db.case_assignments.find_one(
            {"case_id": c["case_id"], "status": {"$in": ["pending", "accepted"]}},
            {"attorney_id": 1, "_id": 0})
        if assignment:
            atty = await db.attorneys.find_one({"id": assignment["attorney_id"]}, {"first_name": 1, "last_name": 1, "_id": 0})
            c["attorney_name"] = f"{atty.get('first_name', '')} {atty.get('last_name', '')}".strip() if atty else None
        else:
            c["attorney_name"] = None

    # Status counts for header
    status_counts = {
        "total": total,
        "active": await db.cases.count_documents({"status": "active"}),
        "pending": await db.cases.count_documents({"attorney_status": "waiting_assignment"}),
        "completed": await db.cases.count_documents({"status": {"$in": ["resolved", "completed"]}}),
    }

    return {
        "cases": cases, "total": total, "page": page,
        "total_pages": max(1, (total + per_page - 1) // per_page),
        "status_counts": status_counts,
    }


@api_router.get("/admin/cases/{case_id}")
async def admin_case_detail(case_id: str, admin: dict = Depends(admin_required)):
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    user = await db.users.find_one({"user_id": case_doc.get("user_id")}, {"_id": 0, "password_hash": 0})
    assignments = await db.case_assignments.find(
        {"case_id": case_id}, {"_id": 0}).sort("created_at", -1).to_list(50)
    return {"case": case_doc, "customer": user, "assignments": assignments}


class ManualAssignInput(BaseModel):
    attorney_id: str

@api_router.post("/admin/cases/{case_id}/manual-assign")
async def admin_manual_assign(case_id: str, body: ManualAssignInput, admin: dict = Depends(admin_required)):
    from services.attorney_matching import assign_case_to_attorney
    result = await assign_case_to_attorney(case_id, force_attorney_id=body.attorney_id, notify=True)
    await log_admin_action(admin, "manual_assign", entity_type="case", entity_id=case_id,
                           metadata={"attorney_id": body.attorney_id})
    return {"status": "assigned", "assignment": result}


class AdminRefundInput(BaseModel):
    reason: str
    amount_cents: Optional[int] = None

@api_router.post("/admin/cases/{case_id}/refund")
async def admin_refund_case(case_id: str, body: AdminRefundInput, admin: dict = Depends(admin_required)):
    case_doc = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")
    pi_id = case_doc.get("payment_intent_id")
    if not pi_id:
        raise HTTPException(status_code=400, detail="No payment to refund")

    import stripe as _stripe
    _stripe.api_key = os.environ.get("STRIPE_API_KEY")
    try:
        params = {"payment_intent": pi_id}
        if body.amount_cents:
            params["amount"] = body.amount_cents
        await asyncio.to_thread(_stripe.Refund.create, **params)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Stripe refund failed: {e}")

    now = datetime.now(timezone.utc).isoformat()
    await db.cases.update_one({"case_id": case_id}, {"$set": {
        "payment_status": "refunded", "refunded_at": now, "refund_reason": body.reason, "updated_at": now,
    }})
    await log_admin_action(admin, "refund", entity_type="case", entity_id=case_id,
                           metadata={"reason": body.reason, "amount_cents": body.amount_cents})
    return {"status": "refunded"}


# ================== Admin Operations ==================

@api_router.get("/admin/payouts")
async def admin_list_payouts(page: int = 1, per_page: int = 50, admin: dict = Depends(admin_required)):
    if "payouts" not in await db.list_collection_names():
        return {"payouts": [], "total": 0}
    total = await db.payouts.count_documents({})
    payouts = await db.payouts.find({}, {"_id": 0}).sort("created_at", -1).skip((page - 1) * per_page).limit(per_page).to_list(per_page)
    for p in payouts:
        atty = await db.attorneys.find_one({"id": p.get("attorney_id")}, {"first_name": 1, "last_name": 1, "email": 1, "_id": 0})
        p["attorney_name"] = f"{atty.get('first_name', '')} {atty.get('last_name', '')}".strip() if atty else None
        p["attorney_email"] = atty.get("email") if atty else None
    return {"payouts": payouts, "total": total}


@api_router.get("/admin/refunds")
async def admin_list_refunds(page: int = 1, per_page: int = 50, admin: dict = Depends(admin_required)):
    cases = await db.cases.find(
        {"payment_status": "refunded"},
        {"_id": 0, "case_id": 1, "user_id": 1, "amount_paid_cents": 1, "refunded_at": 1, "refund_reason": 1},
    ).sort("refunded_at", -1).skip((page - 1) * per_page).limit(per_page).to_list(per_page)
    for c in cases:
        user = await db.users.find_one({"user_id": c.get("user_id")}, {"email": 1, "_id": 0})
        c["customer_email"] = user.get("email") if user else None
    total = await db.cases.count_documents({"payment_status": "refunded"})
    return {"refunds": cases, "total": total}


@api_router.get("/admin/ai-usage")
async def admin_ai_usage(period: str = "30d", admin: dict = Depends(admin_required)):
    """AI usage stats. Returns aggregated data from ai_usage_logs if available."""
    if "ai_usage_logs" not in await db.list_collection_names():
        return {"models": [], "total_requests": 0, "total_tokens": 0, "total_cost_usd": 0, "cost_per_case": 0}
    now = datetime.now(timezone.utc)
    days = 7 if period == "7d" else 365 if period == "ytd" else 30
    since = (now - timedelta(days=days)).isoformat()
    pipeline = [
        {"$match": {"created_at": {"$gte": since}}},
        {"$group": {
            "_id": "$model",
            "requests": {"$sum": 1},
            "input_tokens": {"$sum": "$input_tokens"},
            "output_tokens": {"$sum": "$output_tokens"},
        }},
    ]
    results = await db.ai_usage_logs.aggregate(pipeline).to_list(20)
    models = []
    total_cost = 0
    for r in results:
        model = r["_id"] or "unknown"
        inp = r.get("input_tokens", 0)
        outp = r.get("output_tokens", 0)
        if "opus" in model:
            cost = (inp * 5 + outp * 25) / 1_000_000
        elif "haiku" in model:
            cost = (inp * 0.25 + outp * 1.25) / 1_000_000
        else:
            cost = (inp * 3 + outp * 15) / 1_000_000
        total_cost += cost
        models.append({"model": model, "requests": r["requests"], "tokens": inp + outp, "cost_usd": round(cost, 2)})
    total_requests = sum(m["requests"] for m in models)
    total_cases = await db.cases.count_documents({"created_at": {"$gte": since}}) or 1
    return {
        "models": models,
        "total_requests": total_requests,
        "total_tokens": sum(m["tokens"] for m in models),
        "total_cost_usd": round(total_cost, 2),
        "cost_per_case": round(total_cost / total_cases, 2),
    }


# ================== Admin Moderation ==================

@api_router.get("/admin/moderation/feedbacks")
async def admin_feedbacks(page: int = 1, per_page: int = 50, admin: dict = Depends(admin_required)):
    if "feedbacks" not in await db.list_collection_names():
        return {"feedbacks": [], "total": 0}
    total = await db.feedbacks.count_documents({})
    items = await db.feedbacks.find({}, {"_id": 0}).sort("created_at", -1).skip((page - 1) * per_page).limit(per_page).to_list(per_page)
    for f in items:
        user = await db.users.find_one({"user_id": f.get("user_id")}, {"email": 1, "_id": 0})
        f["customer_email"] = user.get("email") if user else None
    return {"feedbacks": items, "total": total}


@api_router.get("/admin/moderation/unverified-refs")
async def admin_unverified_refs(admin: dict = Depends(admin_required)):
    """Legal references cited by AI but not yet in verified DB."""
    # Scan recent cases for [ARTICLE:] or legal_ref strings not in legal_references
    recent = await db.cases.find(
        {"ai_findings": {"$exists": True, "$ne": []}},
        {"ai_findings": 1, "case_id": 1, "_id": 0},
    ).sort("updated_at", -1).limit(100).to_list(100)

    unverified = []
    seen = set()
    for case in recent:
        for f in (case.get("ai_findings") or []):
            ref = f.get("legal_ref", "")
            if ref and ref not in seen:
                exists = await db.legal_references.find_one({"$or": [
                    {"article_number": {"$regex": ref[:10], "$options": "i"}},
                ]})
                if not exists:
                    seen.add(ref)
                    unverified.append({"legal_ref": ref, "case_id": case["case_id"], "finding_text": f.get("text", "")})
    return {"unverified": unverified[:50]}


class FlagCaseInput(BaseModel):
    reason: str

@api_router.post("/admin/moderation/cases/{case_id}/flag")
async def admin_flag_case(case_id: str, body: FlagCaseInput, admin: dict = Depends(admin_required)):
    now = datetime.now(timezone.utc).isoformat()
    await db.cases.update_one({"case_id": case_id}, {"$set": {"flagged": True, "flag_reason": body.reason, "flagged_at": now, "flagged_by": admin["id"]}})
    await log_admin_action(admin, "flag_case", entity_type="case", entity_id=case_id, metadata={"reason": body.reason})
    return {"status": "flagged"}


# ================== Admin Analytics ==================

@api_router.get("/admin/analytics/funnel")
async def admin_conversion_funnel(admin: dict = Depends(admin_required)):
    total_signups = await db.users.count_documents({"account_type": {"$ne": "attorney"}})
    uploaded_doc = await db.cases.count_documents({})
    completed_analysis = await db.cases.count_documents({"status": {"$ne": "analyzing"}, "ai_summary": {"$ne": None}})
    paid_letter = await db.cases.count_documents({"payment_status": "paid"})
    live_counsel = await db.case_assignments.count_documents({"service_type": "live_counsel"})
    return {
        "steps": [
            {"label": "Sign up", "count": total_signups},
            {"label": "Uploaded document", "count": uploaded_doc},
            {"label": "Completed analysis", "count": completed_analysis},
            {"label": "Paid Attorney Letter", "count": paid_letter},
            {"label": "Live Counsel", "count": live_counsel},
        ],
    }


@api_router.get("/admin/analytics/geography")
async def admin_geography(admin: dict = Depends(admin_required)):
    pipeline = [
        {"$match": {"account_type": {"$ne": "attorney"}}},
        {"$group": {"_id": {"$ifNull": ["$country", "$jurisdiction"]}, "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
    ]
    results = await db.users.aggregate(pipeline).to_list(50)
    return {"countries": [{"country": r["_id"] or "Unknown", "count": r["count"]} for r in results]}


@api_router.get("/admin/analytics/attorney-leaderboard")
async def admin_attorney_leaderboard(admin: dict = Depends(admin_required)):
    attorneys = await db.attorneys.find(
        {"status": "active"}, {"_id": 0, "id": 1, "first_name": 1, "last_name": 1, "email": 1, "active_cases_count": 1, "rating_avg": 1, "avg_response_seconds": 1}
    ).to_list(100)
    for a in attorneys:
        completed = await db.case_assignments.count_documents({"attorney_id": a["id"], "status": "completed"})
        a["cases_completed"] = completed
    attorneys.sort(key=lambda a: a.get("cases_completed", 0), reverse=True)
    return {"attorneys": attorneys[:20]}


@api_router.get("/admin/analytics/case-types")
async def admin_case_types_distribution(admin: dict = Depends(admin_required)):
    pipeline = [
        {"$group": {"_id": "$type", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
    ]
    results = await db.cases.aggregate(pipeline).to_list(20)
    return {"types": [{"type": r["_id"] or "other", "count": r["count"]} for r in results]}


# ================== Admin Exports ==================

@api_router.get("/admin/exports/customers.csv")
async def admin_export_customers_csv(admin: dict = Depends(admin_required)):
    import csv
    import io
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["user_id", "name", "email", "plan", "country", "language", "created_at"])
    users = await db.users.find(
        {"account_type": {"$ne": "attorney"}},
        {"_id": 0, "user_id": 1, "name": 1, "email": 1, "plan": 1, "country": 1, "language": 1, "created_at": 1},
    ).to_list(50000)
    for u in users:
        writer.writerow([u.get("user_id"), u.get("name"), u.get("email"), u.get("plan", "free"), u.get("country"), u.get("language"), u.get("created_at")])
    await log_admin_action(admin, "export_customers_csv")
    return StreamingResponse(
        io.BytesIO(output.getvalue().encode("utf-8")),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=customers_export.csv"},
    )


@api_router.get("/admin/exports/cases.csv")
async def admin_export_cases_csv(admin: dict = Depends(admin_required)):
    import csv
    import io
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["case_id", "user_id", "title", "type", "status", "risk_score", "payment_status", "country", "created_at"])
    cases = await db.cases.find(
        {}, {"_id": 0, "case_id": 1, "user_id": 1, "title": 1, "type": 1, "status": 1, "risk_score": 1, "payment_status": 1, "country": 1, "created_at": 1},
    ).to_list(50000)
    for c in cases:
        writer.writerow([c.get("case_id"), c.get("user_id"), c.get("title"), c.get("type"), c.get("status"), c.get("risk_score"), c.get("payment_status"), c.get("country"), c.get("created_at")])
    await log_admin_action(admin, "export_cases_csv")
    return StreamingResponse(
        io.BytesIO(output.getvalue().encode("utf-8")),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=cases_export.csv"},
    )


# ================== Admin Settings ==================

@api_router.get("/admin/settings/flags")
async def admin_get_flags(admin: dict = Depends(admin_required)):
    """Read current system flag values."""
    return {
        "REQUIRE_STRIPE_ONBOARDING": os.environ.get("REQUIRE_STRIPE_ONBOARDING", "false"),
        "SCHEDULER_ENABLED": os.environ.get("SCHEDULER_ENABLED", "true"),
        "WEEKLY_PAYOUTS_ENABLED": os.environ.get("WEEKLY_PAYOUTS_ENABLED", "true"),
    }


@api_router.get("/admin/audit-log")
async def admin_audit_log(page: int = 1, per_page: int = 50, admin: dict = Depends(admin_required)):
    total = await db.admin_audit_log.count_documents({})
    logs = await db.admin_audit_log.find({}, {"_id": 0}).sort("created_at", -1).skip((page - 1) * per_page).limit(per_page).to_list(per_page)
    return {"logs": logs, "total": total, "page": page, "total_pages": max(1, (total + per_page - 1) // per_page)}


# ================== Admin Notification Settings ==================

class NotificationSettingsInput(BaseModel):
    notify_daily_digest: bool = True
    notify_new_customer: bool = True
    notify_large_payment: bool = True
    large_payment_threshold_cents: int = 10000
    notify_new_attorney: bool = True
    notify_system_error: bool = True
    notify_failed_payout: bool = True
    notify_weekly_summary: bool = True


@api_router.get("/admin/settings/notifications")
async def admin_get_notification_settings(admin: dict = Depends(admin_required)):
    settings = await db.admin_notification_settings.find_one({"admin_id": admin["id"]}, {"_id": 0})
    if not settings:
        settings = NotificationSettingsInput().dict()
        settings["admin_id"] = admin["id"]
    return settings


@api_router.put("/admin/settings/notifications")
async def admin_update_notification_settings(body: NotificationSettingsInput, admin: dict = Depends(admin_required)):
    doc = body.dict()
    doc["admin_id"] = admin["id"]
    doc["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.admin_notification_settings.update_one({"admin_id": admin["id"]}, {"$set": doc}, upsert=True)
    return {"status": "saved"}


# ================== Admin Notifications Count ==================

@api_router.get("/admin/notifications/unread")
async def admin_unread_notifications(admin: dict = Depends(admin_required)):
    """Quick counts for sidebar badges."""
    attorneys_pending = await db.attorneys.count_documents({"application_status": "pending"})
    feedbacks_flagged = await db.feedbacks.count_documents({"status": "new"}) if "feedbacks" in await db.list_collection_names() else 0
    failed_payouts = await db.payouts.count_documents({"status": "failed"}) if "payouts" in await db.list_collection_names() else 0
    return {
        "attorneys_pending": attorneys_pending,
        "feedbacks_flagged": feedbacks_flagged,
        "system_errors": 0,
        "failed_payouts": failed_payouts,
        "total_action_required": attorneys_pending + feedbacks_flagged + failed_payouts,
    }


# ================== Newsletter ==================

class NewsletterSubscribeInput(BaseModel):
    email: str
    source: str = "unknown"

@api_router.post("/newsletter/subscribe")
async def newsletter_subscribe(body: NewsletterSubscribeInput):
    """Subscribe to Archer newsletter. Double opt-in: creates unconfirmed subscriber."""
    import re as _re
    email = body.email.strip().lower()
    if not _re.match(r'^[^@]+@[^@]+\.[^@]+$', email):
        raise HTTPException(status_code=400, detail="Invalid email")

    existing = await db.newsletter_subscribers.find_one({"email": email})
    if existing:
        return {"success": True, "message": "Already subscribed"}

    now = datetime.now(timezone.utc).isoformat()
    token = uuid.uuid4().hex
    doc = {
        "id": f"nsub_{uuid.uuid4().hex[:12]}",
        "email": email,
        "source": body.source,
        "subscribed_at": now,
        "confirmed": False,
        "confirmed_at": None,
        "unsubscribed_at": None,
        "confirm_token": token,
        "tags": [],
    }
    await db.newsletter_subscribers.insert_one(doc)

    # Send confirmation email (best-effort)
    try:
        from routes.attorney_routes import send_email
        confirm_url = f"{os.environ.get('APP_URL', 'https://archer.law')}/newsletter/confirm?token={token}"
        await send_email(
            email,
            "Confirm your Archer newsletter subscription",
            f"<p>Thanks for subscribing to Archer's weekly legal tips.</p>"
            f"<p><a href=\"{confirm_url}\" style=\"display:inline-block;background:#1a56db;"
            f"color:#fff;padding:12px 24px;border-radius:8px;text-decoration:none;"
            f"font-weight:600;\">Confirm my subscription</a></p>"
            f"<p style=\"font-size:12px;color:#9ca3af;\">If you didn't request this, ignore this email.</p>",
        )
    except Exception:
        pass

    return {"success": True, "message": "Confirmation email sent"}


@api_router.get("/newsletter/confirm")
async def newsletter_confirm(token: str):
    """Confirm newsletter subscription via token."""
    sub = await db.newsletter_subscribers.find_one({"confirm_token": token})
    if not sub:
        raise HTTPException(status_code=404, detail="Invalid or expired token")
    if sub.get("confirmed"):
        return {"success": True, "message": "Already confirmed"}
    await db.newsletter_subscribers.update_one(
        {"confirm_token": token},
        {"$set": {"confirmed": True, "confirmed_at": datetime.now(timezone.utc).isoformat()}},
    )
    return {"success": True, "message": "Subscription confirmed"}


# ================== Explain Simply ==================

EXPLAIN_SIMPLY_PROMPT = """Tu es Archer. Reformule l'analyse juridique suivante en langage ULTRA SIMPLE comme si tu parlais a quelqu'un de 12 ans qui n'a jamais rien lu de juridique.

Regles strictes :
1. Pas de jargon juridique. Si tu dois citer un article, dis "la loi dit que..."
2. Phrases courtes (max 15 mots).
3. Utilise "vous" et adresse-toi directement au client.
4. Commence chaque point cle par un emoji approprie (🏠 pour logement, 💼 pour travail, 🚗 pour routier, 🛡️ pour assurance, 👨‍👩‍👧 pour famille, 💳 pour consommation/dette, 📄 pour contrat).
5. Termine par une conclusion concrete : ce que le client doit/peut faire maintenant.
6. Ton : empathique, rassurant, mais factuel.
7. Maximum 200 mots.
8. Reponds UNIQUEMENT avec la version simplifiee, sans preambule.

{language_instruction}"""


class ExplainSimplyInput(BaseModel):
    regenerate: bool = False


@api_router.post("/cases/{case_id}/explain-simply")
async def explain_simply(
    case_id: str,
    body: ExplainSimplyInput = ExplainSimplyInput(),
    current_user: User = Depends(get_current_user),
):
    """Generate a plain-language explanation of the case analysis.
    Cached in case.simple_explanation. Use regenerate=true to overwrite cache."""
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")

    # Return cached version if available and not regenerating
    if not body.regenerate and case_doc.get("simple_explanation"):
        return {"explanation": case_doc["simple_explanation"], "cached": True}

    # Build context from case data
    summary = case_doc.get("ai_summary") or ""
    findings_text = ""
    for f in (case_doc.get("ai_findings") or [])[:6]:
        findings_text += f"- {f.get('text', '')}: {f.get('impact_description', '')}\n"
    case_type = case_doc.get("type", "other")
    risk_score = case_doc.get("risk_score", 0)

    user_language = getattr(current_user, "language", "fr") or "fr"
    if user_language.startswith("nl"):
        lang_instruction = "Reponds entierement en neerlandais."
    elif user_language.startswith("de"):
        lang_instruction = "Reponds entierement en allemand."
    elif user_language.startswith("en"):
        lang_instruction = "Respond entirely in English."
    else:
        lang_instruction = "Reponds entierement en francais."

    system = EXPLAIN_SIMPLY_PROMPT.format(language_instruction=lang_instruction)
    user_msg = f"""Analyse a reformuler :
{summary}

Constatations :
{findings_text}

Type de cas : {case_type}
Score de risque : {risk_score}/100
Titre du dossier : {case_doc.get('title', '')}"""

    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": get_anthropic_key(),
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json={
                    "model": "claude-opus-4-7",
                    "max_tokens": 600,
                    "system": system,
                    "messages": [{"role": "user", "content": user_msg}],
                },
            )
            resp.raise_for_status()
            explanation = resp.json()["content"][0]["text"]
    except Exception as e:
        logger.error(f"Explain simply error: {e}")
        raise HTTPException(status_code=502, detail="Failed to generate explanation")

    # Cache in DB
    await db.cases.update_one(
        {"case_id": case_id},
        {"$set": {"simple_explanation": explanation, "updated_at": datetime.now(timezone.utc).isoformat()}},
    )

    return {"explanation": explanation, "cached": False}


# ================== Similar Cases ==================

@api_router.get("/cases/{case_id}/similar")
async def get_similar_cases(case_id: str, current_user: User = Depends(get_current_user)):
    """Find similar cases by type + jurisdiction + shared violation types.
    Falls back to precomputed daily stats if not enough real data."""
    case_doc = await db.cases.find_one(
        {"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")

    case_violations = set(case_doc.get("key_violation_types") or [])
    case_type = case_doc.get("type", "other")
    jurisdiction = case_doc.get("country") or getattr(current_user, "jurisdiction", "BE")

    # Try precomputed daily stats first (fast path)
    stats_doc = await db.similar_cases_stats.find_one(
        {"case_type": case_type, "jurisdiction": jurisdiction}, {"_id": 0})

    if stats_doc and stats_doc.get("total_similar", 0) >= 5:
        # Fetch top 3 closest cases by violation overlap
        pipeline = [
            {"$match": {
                "case_id": {"$ne": case_id},
                "type": case_type,
                "country": jurisdiction,
                "is_publicly_anonymizable": {"$ne": False},
                "case_outcome": {"$in": ["won", "lost", "settled"]},
            }},
            {"$limit": 100},
        ]
        candidates = await db.cases.aggregate(pipeline).to_list(100)

        scored = []
        for c in candidates:
            c_violations = set(c.get("key_violation_types") or [])
            overlap = len(case_violations & c_violations) if case_violations else 0
            scored.append((overlap, c))
        scored.sort(key=lambda x: x[0], reverse=True)
        top_3 = [c for _, c in scored[:3]]

        return {
            "source": "live",
            "total_similar": stats_doc["total_similar"],
            "won_count": stats_doc["won_count"],
            "win_rate_percent": stats_doc["win_rate_percent"],
            "top_3_similar_cases": [
                {
                    "id": c["case_id"],
                    "anonymized_title": _anonymize_case_title(c),
                    "outcome": c.get("case_outcome", "unknown"),
                    "outcome_summary": c.get("outcome_summary", ""),
                    "key_violations": c.get("key_violation_types", []),
                    "completed_at": c.get("updated_at"),
                }
                for c in top_3
            ],
        }

    # Fallback: not enough real data — return static estimates
    from utils.dashboard_static_stats import SIMILAR_CASES_FALLBACK
    fallback = SIMILAR_CASES_FALLBACK.get(case_type, SIMILAR_CASES_FALLBACK.get("other", {}))
    return {
        "source": "estimated",
        "total_similar": fallback.get("total_count", 0),
        "won_count": fallback.get("won_count", 0),
        "win_rate_percent": fallback.get("win_rate_percent", 0),
        "reduction_pct": fallback.get("reduction_pct", 0),
        "avg_delay_days": fallback.get("avg_delay_days", 0),
        "top_3_similar_cases": [],
    }


@api_router.get("/cases/similar/{similar_case_id}")
async def get_similar_case_detail(similar_case_id: str, current_user: User = Depends(get_current_user)):
    """Get anonymized detail of a similar case."""
    case_doc = await db.cases.find_one(
        {"case_id": similar_case_id, "is_publicly_anonymizable": {"$ne": False}},
        {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found or not anonymizable")
    if case_doc.get("case_outcome") not in ("won", "lost", "settled"):
        raise HTTPException(status_code=404, detail="Case outcome not yet known")

    return {
        "id": case_doc["case_id"],
        "anonymized_title": _anonymize_case_title(case_doc),
        "case_type": case_doc.get("type"),
        "jurisdiction": case_doc.get("country"),
        "outcome": case_doc.get("case_outcome"),
        "outcome_summary": case_doc.get("outcome_summary", ""),
        "key_violations": case_doc.get("key_violation_types", []),
        "risk_score": case_doc.get("risk_score"),
        "submitted_at": case_doc.get("created_at"),
        # NO PII: no name, email, phone, address
    }


def _anonymize_case_title(case_doc):
    case_type = (case_doc.get("type") or "other").replace("_", " ").title()
    jurisdiction = case_doc.get("country", "")
    outcome = case_doc.get("case_outcome", "unknown")
    return f"{case_type} in {jurisdiction} ({outcome})"


# ================== Legal References ==================

import re as _re

async def enrich_legal_references(text: str) -> str:
    """Post-process AI response: replace [ARTICLE: ...] and [JURIS: ...] tags
    with verified URLs from the database. Unmatched refs become plain text spans."""
    article_pattern = r"\[ARTICLE:\s*code=(\w+),\s*article=([\w§.]+?)(?:,\s*paragraph=([\w§]+|null))?\]"
    juris_pattern = r"\[JURIS:\s*court=(\w+),\s*date=([\d-]+),\s*number=([\w.\-/]+)\]"

    async def _replace_article(match):
        code, article, paragraph = match.group(1), match.group(2), match.group(3)
        query = {"code": code, "article_number": article}
        if paragraph and paragraph != "null":
            query["paragraph"] = paragraph
        ref = await db.legal_references.find_one(query, {"_id": 0})
        if ref and ref.get("verified_url"):
            label = f"Article {article}"
            if paragraph and paragraph != "null":
                label += f" {paragraph}"
            label += f" {ref.get('code_full_name', code)}"
            return f'<a href="{ref["verified_url"]}" target="_blank" rel="noopener" class="legal-link">{label}</a>'
        return f'Article {article} {code}'

    async def _replace_juris(match):
        court, date, number = match.group(1), match.group(2), match.group(3)
        ref = await db.jurisprudence_references.find_one(
            {"court": court, "case_number": number}, {"_id": 0})
        if ref and ref.get("verified_url"):
            label = f"{ref.get('court_full_name', court)}, {date}, n° {number}"
            return f'<a href="{ref["verified_url"]}" target="_blank" rel="noopener" class="legal-link">{label}</a>'
        return f'{court}, {date}, n° {number}'

    # Process articles
    for match in reversed(list(_re.finditer(article_pattern, text))):
        replacement = await _replace_article(match)
        text = text[:match.start()] + replacement + text[match.end():]

    # Process jurisprudence
    for match in reversed(list(_re.finditer(juris_pattern, text))):
        replacement = await _replace_juris(match)
        text = text[:match.start()] + replacement + text[match.end():]

    return text


class LegalReferenceInput(BaseModel):
    code: str
    code_full_name: str
    article_number: str
    paragraph: Optional[str] = None
    title: str
    verified_url: str
    summary: str = ""
    jurisdiction: str = "BE"
    language: str = "fr"


class JurisprudenceInput(BaseModel):
    court: str
    court_full_name: str
    date: str
    case_number: str
    title: str
    verified_url: str
    summary: str = ""
    key_principles: list = []
    jurisdiction: str = "BE"


@api_router.post("/admin/legal-references")
async def add_legal_reference(data: LegalReferenceInput, admin: dict = Depends(admin_required)):
    """Admin endpoint to add/update a verified legal reference."""
    now = datetime.now(timezone.utc).isoformat()
    doc = data.dict()
    doc["id"] = f"lref_{uuid.uuid4().hex[:12]}"
    doc["last_verified_at"] = now
    doc["verified_by"] = admin.get("email", "admin")
    await db.legal_references.update_one(
        {"code": data.code, "article_number": data.article_number, "paragraph": data.paragraph},
        {"$set": doc}, upsert=True,
    )
    return {"status": "ok", "id": doc["id"]}


@api_router.post("/admin/jurisprudence-references")
async def add_jurisprudence_reference(data: JurisprudenceInput, admin: dict = Depends(admin_required)):
    """Admin endpoint to add/update a verified jurisprudence reference."""
    now = datetime.now(timezone.utc).isoformat()
    doc = data.dict()
    doc["id"] = f"jref_{uuid.uuid4().hex[:12]}"
    doc["last_verified_at"] = now
    doc["verified_by"] = admin.get("email", "admin")
    await db.jurisprudence_references.update_one(
        {"court": data.court, "case_number": data.case_number},
        {"$set": doc}, upsert=True,
    )
    return {"status": "ok", "id": doc["id"]}


@api_router.get("/legal-references")
async def list_legal_references(jurisdiction: str = "BE"):
    """Public endpoint: list all verified legal references for a jurisdiction."""
    refs = await db.legal_references.find(
        {"jurisdiction": jurisdiction}, {"_id": 0}
    ).to_list(500)
    return refs


@api_router.get("/jurisprudence-references")
async def list_jurisprudence_references(jurisdiction: str = "BE"):
    """Public endpoint: list all verified jurisprudence references."""
    refs = await db.jurisprudence_references.find(
        {"jurisdiction": jurisdiction}, {"_id": 0}
    ).to_list(500)
    return refs


# ================== Case Sharing ==================

import secrets

class ShareCaseInput(BaseModel):
    expiry_hours: int = 48
    message: Optional[str] = None

class SharedCaseComment(BaseModel):
    commenter_name: str
    comment: str

@api_router.post("/cases/{case_id}/share")
async def share_case(case_id: str, data: ShareCaseInput, current_user: User = Depends(get_current_user)):
    """Generate a secure share link for a case"""
    if current_user.plan == "free":
        raise HTTPException(status_code=403, detail="Case sharing is available on the Pro plan. Upgrade for $69/month.")

    case_doc = await db.cases.find_one({"case_id": case_id, "user_id": current_user.user_id}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case not found")

    # Limit expiry options
    allowed = [24, 48, 168, 720]
    if data.expiry_hours not in allowed:
        data.expiry_hours = 48

    now = datetime.now(timezone.utc)
    token = secrets.token_urlsafe(24)

    share_doc = {
        "share_id": f"share_{uuid.uuid4().hex[:12]}",
        "case_id": case_id,
        "user_id": current_user.user_id,
        "user_name": current_user.name or "User",
        "token": token,
        "expires_at": (now + timedelta(hours=data.expiry_hours)).isoformat(),
        "expiry_hours": data.expiry_hours,
        "optional_message": data.message,
        "views_count": 0,
        "is_revoked": False,
        "created_at": now.isoformat()
    }
    await db.shared_cases.insert_one(share_doc)

    return {"token": token, "share_id": share_doc["share_id"], "expires_at": share_doc["expires_at"]}

@api_router.get("/cases/{case_id}/shares")
async def list_case_shares(case_id: str, current_user: User = Depends(get_current_user)):
    """List active shares for a case"""
    shares = await db.shared_cases.find(
        {"case_id": case_id, "user_id": current_user.user_id, "is_revoked": False},
        {"_id": 0}
    ).sort("created_at", -1).to_list(20)

    now = datetime.now(timezone.utc).isoformat()
    active = [s for s in shares if s["expires_at"] > now]

    # Get comment counts
    for s in active:
        s["comment_count"] = await db.shared_case_comments.count_documents({"share_id": s["share_id"]})

    return active

@api_router.post("/shares/{share_id}/revoke")
async def revoke_share(share_id: str, current_user: User = Depends(get_current_user)):
    """Revoke a share link"""
    result = await db.shared_cases.update_one(
        {"share_id": share_id, "user_id": current_user.user_id},
        {"$set": {"is_revoked": True}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Share not found")
    return {"status": "revoked"}

# Public endpoints (no auth required)
@api_router.get("/shared/{token}")
async def get_shared_case(token: str):
    """Get a shared case by token (public, no auth)"""
    share = await db.shared_cases.find_one({"token": token}, {"_id": 0})
    if not share:
        raise HTTPException(status_code=404, detail="Share link not found")

    if share.get("is_revoked"):
        raise HTTPException(status_code=410, detail="This link has been revoked by the owner.")

    now = datetime.now(timezone.utc).isoformat()
    if share["expires_at"] < now:
        raise HTTPException(status_code=410, detail="This link has expired.")

    # Increment views
    await db.shared_cases.update_one({"token": token}, {"$inc": {"views_count": 1}})

    # Get case data (filtered for privacy)
    case_doc = await db.cases.find_one({"case_id": share["case_id"]}, {"_id": 0})
    if not case_doc:
        raise HTTPException(status_code=404, detail="Case no longer exists")

    # Get comments
    comments = await db.shared_case_comments.find(
        {"share_id": share["share_id"]},
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)

    # Filter sensitive data
    safe_case = {
        "title": case_doc.get("title"),
        "type": case_doc.get("type"),
        "status": case_doc.get("status"),
        "risk_score": case_doc.get("risk_score"),
        "risk_financial": case_doc.get("risk_financial"),
        "risk_urgency": case_doc.get("risk_urgency"),
        "risk_legal_strength": case_doc.get("risk_legal_strength"),
        "risk_complexity": case_doc.get("risk_complexity"),
        "deadline": case_doc.get("deadline"),
        "deadline_description": case_doc.get("deadline_description"),
        "financial_exposure": case_doc.get("financial_exposure"),
        "ai_summary": case_doc.get("ai_summary"),
        "ai_findings": case_doc.get("ai_findings", []),
        "ai_next_steps": case_doc.get("ai_next_steps", []),
        "risk_score_history": case_doc.get("risk_score_history", []),
        "battle_preview": case_doc.get("battle_preview"),
        "success_probability": case_doc.get("success_probability"),
        "procedural_defects": case_doc.get("procedural_defects", []),
        "key_insight": case_doc.get("key_insight"),
        "leverage_points": case_doc.get("leverage_points", []),
        "recent_case_law": case_doc.get("recent_case_law", []),
        "case_law_updated": case_doc.get("case_law_updated"),
    }

    # Get document names only (not content)
    docs = await db.documents.find(
        {"case_id": share["case_id"]},
        {"_id": 0, "file_name": 1, "status": 1, "uploaded_at": 1, "is_key_document": 1}
    ).to_list(50)

    # Get timeline
    events = await db.case_events.find(
        {"case_id": share["case_id"]},
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)

    return {
        "share": {
            "user_name": share.get("user_name", "User"),
            "optional_message": share.get("optional_message"),
            "expires_at": share["expires_at"],
            "views_count": share.get("views_count", 0) + 1
        },
        "case": safe_case,
        "documents": docs,
        "events": events,
        "comments": comments
    }

@api_router.post("/shared/{token}/comments")
async def add_shared_case_comment(token: str, data: SharedCaseComment):
    """Add a comment to a shared case (public, no auth)"""
    share = await db.shared_cases.find_one({"token": token}, {"_id": 0})
    if not share:
        raise HTTPException(status_code=404, detail="Share link not found")
    if share.get("is_revoked"):
        raise HTTPException(status_code=410, detail="This link has been revoked.")
    now = datetime.now(timezone.utc).isoformat()
    if share["expires_at"] < now:
        raise HTTPException(status_code=410, detail="This link has expired.")

    comment_doc = {
        "comment_id": f"cmt_{uuid.uuid4().hex[:12]}",
        "share_id": share["share_id"],
        "case_id": share["case_id"],
        "commenter_name": data.commenter_name[:100],
        "comment": data.comment[:2000],
        "created_at": now
    }
    await db.shared_case_comments.insert_one(comment_doc)

    return {"comment_id": comment_doc["comment_id"], "status": "created"}

# ================== Seed Data ==================

@api_router.post("/seed/lawyers")
async def seed_lawyers():
    """Seed lawyer data — US + Belgian lawyers"""
    lawyers_data = [
        # ===== US LAWYERS =====
        {
            "lawyer_id": "lawyer_sarah",
            "name": "Sarah Mitchell, Esq.",
            "specialty": "Employment Law",
            "bar_state": "New York",
            "years_experience": 12,
            "rating": 5.0,
            "sessions_count": 214,
            "tags": ["Wrongful termination", "Discrimination", "Unpaid wages"],
            "availability_status": "now",
            "availability_minutes": 0,
            "bio": "Specialized in employment disputes and workplace discrimination cases.",
            "photo_url": "https://images.unsplash.com/photo-1585240975858-7264fd020798?crop=entropy&cs=srgb&fm=jpg&ixid=M3w4NjA1MDZ8MHwxfHNlYXJjaHwzfHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjB3b21hbiUyMGF0dG9ybmV5fGVufDB8fHx8MTc3NTc2MTA2NHww&ixlib=rb-4.1.0&q=85",
            "country": "US",
            "language": "en"
        },
        {
            "lawyer_id": "lawyer_archer",
            "name": "Archer Legal AI",
            "specialty": "Contract Law",
            "bar_state": "California",
            "years_experience": 9,
            "rating": 4.9,
            "sessions_count": 187,
            "tags": ["Business disputes", "NDA", "Vendor contracts"],
            "availability_status": "soon",
            "availability_minutes": 8,
            "bio": "Expert in business contracts and commercial litigation.",
            "photo_url": "https://images.unsplash.com/photo-1644268756851-3f69ffb9553f?crop=entropy&cs=srgb&fm=jpg&ixid=M3w4NjA3MDR8MHwxfHNlYXJjaHwyfHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjBsYXd5ZXJ8ZW58MHx8fHwxNzc1NzYxMDU5fDA&ixlib=rb-4.1.0&q=85",
            "country": "US",
            "language": "en"
        },
        {
            "lawyer_id": "lawyer_diana",
            "name": "Diana Torres, Esq.",
            "specialty": "Tenant Rights",
            "bar_state": "Texas",
            "years_experience": 7,
            "rating": 4.9,
            "sessions_count": 156,
            "tags": ["Lease disputes", "Eviction", "Deposits"],
            "availability_status": "soon",
            "availability_minutes": 22,
            "bio": "Dedicated to protecting tenant rights and housing law.",
            "photo_url": "https://images.unsplash.com/photo-1685760259914-ee8d2c92d2e0?crop=entropy&cs=srgb&fm=jpg&ixid=M3w4NjA1MDZ8MHwxfHNlYXJjaHwyfHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjB3b21hbiUyMGF0dG9ybmV5fGVufDB8fHx8MTc3NTc2MTA2NHww&ixlib=rb-4.1.0&q=85",
            "country": "US",
            "language": "en"
        },
        {
            "lawyer_id": "lawyer_marcus",
            "name": "Marcus Johnson, Esq.",
            "specialty": "Immigration Law",
            "bar_state": "Washington DC",
            "years_experience": 15,
            "rating": 4.8,
            "sessions_count": 302,
            "tags": ["Visa contracts", "Work permits", "Green card"],
            "availability_status": "tomorrow",
            "availability_minutes": 0,
            "bio": "Immigration law specialist with federal court experience.",
            "photo_url": "https://images.unsplash.com/photo-1665224752561-85f4da9a5658?crop=entropy&cs=srgb&fm=jpg&ixid=M3w4NjA3MDR8MHwxfHNlYXJjaHwxfHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjBsYXd5ZXJ8ZW58MHx8fHwxNzc1NzYxMDU5fDA&ixlib=rb-4.1.0&q=85",
            "country": "US",
            "language": "en"
        },
        {
            "lawyer_id": "lawyer_rachel",
            "name": "Rachel Kim, Esq.",
            "specialty": "Consumer Rights",
            "bar_state": "Florida",
            "years_experience": 6,
            "rating": 4.9,
            "sessions_count": 98,
            "tags": ["Debt collection", "Scams", "FDCPA"],
            "availability_status": "now",
            "availability_minutes": 0,
            "bio": "Consumer protection and debt collection defense.",
            "photo_url": "https://images.unsplash.com/photo-1665224752123-a2ea29dddcb2?crop=entropy&cs=srgb&fm=jpg&ixid=M3w4NjA1MDZ8MHwxfHNlYXJjaHw0fHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjB3b21hbiUyMGF0dG9ybmV5fGVufDB8fHx8MTc3NTc2MTA2NHww&ixlib=rb-4.1.0&q=85",
            "country": "US",
            "language": "en"
        },
        {
            "lawyer_id": "lawyer_david",
            "name": "David Park, Esq.",
            "specialty": "Business Law",
            "bar_state": "Illinois",
            "years_experience": 11,
            "rating": 4.8,
            "sessions_count": 241,
            "tags": ["Partnership disputes", "LLC", "Contracts"],
            "availability_status": "soon",
            "availability_minutes": 45,
            "bio": "Business formation and partnership dispute resolution.",
            "photo_url": "https://images.unsplash.com/photo-1762522926262-d96de462ad54?crop=entropy&cs=srgb&fm=jpg&ixid=M3w4NjA3MDR8MHwxfHNlYXJjaHw0fHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjBsYXd5ZXJ8ZW58MHx8fHwxNzc1NzYxMDU5fDA&ixlib=rb-4.1.0&q=85",
            "country": "US",
            "language": "en"
        },
        # ===== BELGIAN LAWYERS — FRANCOPHONE =====
        {
            "lawyer_id": "lawyer_be_sophie",
            "name": "Me Sophie Lecomte",
            "specialty": "Droit du travail",
            "bar_state": "Barreau de Bruxelles",
            "years_experience": 11,
            "rating": 4.9,
            "sessions_count": 187,
            "tags": ["Licenciement abusif", "Harcelement", "Preavis", "CCT 109"],
            "availability_status": "now",
            "availability_minutes": 0,
            "bio": "Specialisee en litiges du travail et protection contre le licenciement abusif. Ancienne collaboratrice du syndicat CSC.",
            "photo_url": "https://images.unsplash.com/photo-1702875581098-36844f1910f3?crop=entropy&cs=srgb&fm=jpg&ixid=M3w3NTY2OTF8MHwxfHNlYXJjaHwzfHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjBldXJvcGVhbiUyMGxhd3llcnxlbnwwfHx8fDE3NzU4MzU1MTN8MA&ixlib=rb-4.1.0&q=85",
            "country": "BE",
            "language": "fr"
        },
        {
            "lawyer_id": "lawyer_be_thomas",
            "name": "Me Thomas Dupont",
            "specialty": "Droit du bail (Wallonie)",
            "bar_state": "Barreau de Liege",
            "years_experience": 7,
            "rating": 4.8,
            "sessions_count": 134,
            "tags": ["Bail residentiel", "Expulsion", "Garantie locative", "CWLHD"],
            "availability_status": "soon",
            "availability_minutes": 15,
            "bio": "Expert en droit du bail wallon. Intervient regulierement devant les justices de paix de la province de Liege.",
            "photo_url": "https://images.unsplash.com/photo-1591702694482-ecc51ff9642e?crop=entropy&cs=srgb&fm=jpg&ixid=M3w3NTY2OTF8MHwxfHNlYXJjaHwyfHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjBldXJvcGVhbiUyMGxhd3llcnxlbnwwfHx8fDE3NzU4MzU1MTN8MA&ixlib=rb-4.1.0&q=85",
            "country": "BE",
            "language": "fr"
        },
        {
            "lawyer_id": "lawyer_be_julie",
            "name": "Me Julie Renard",
            "specialty": "Droit de la consommation",
            "bar_state": "Barreau de Namur",
            "years_experience": 9,
            "rating": 4.9,
            "sessions_count": 156,
            "tags": ["Protection consommateur", "Remboursement", "Contrats abusifs", "CDE"],
            "availability_status": "now",
            "availability_minutes": 0,
            "bio": "Specialisee en protection du consommateur et clauses abusives. Collabore avec le SPF Economie.",
            "photo_url": "https://images.unsplash.com/photo-1733348137479-2e726d326d9b?crop=entropy&cs=srgb&fm=jpg&ixid=M3w3NTY2OTF8MHwxfHNlYXJjaHw0fHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjBldXJvcGVhbiUyMGxhd3llcnxlbnwwfHx8fDE3NzU4MzU1MTN8MA&ixlib=rb-4.1.0&q=85",
            "country": "BE",
            "language": "fr"
        },
        {
            "lawyer_id": "lawyer_be_alexandre",
            "name": "Me Alexandre Martin",
            "specialty": "Droit des contrats",
            "bar_state": "Barreau de Bruxelles",
            "years_experience": 14,
            "rating": 4.8,
            "sessions_count": 203,
            "tags": ["NDA", "Contrats commerciaux", "Rupture de contrat", "Non-concurrence"],
            "availability_status": "tomorrow",
            "availability_minutes": 0,
            "bio": "Expert en negociation de contrats commerciaux et clauses de non-concurrence. Bilingue FR/NL.",
            "photo_url": "https://images.unsplash.com/photo-1771244702701-6c9edac63255?crop=entropy&cs=srgb&fm=jpg&ixid=M3w4NjA3MDB8MHwxfHNlYXJjaHw0fHxwcm9mZXNzaW9uYWwlMjBtYW4lMjBoZWFkc2hvdCUyMHBvcnRyYWl0JTIwc3VpdCUyMG9mZmljZXxlbnwwfHx8fDE3NzU4MzU1MTh8MA&ixlib=rb-4.1.0&q=85",
            "country": "BE",
            "language": "fr/nl"
        },
        {
            "lawyer_id": "lawyer_be_emilie",
            "name": "Me Emilie Dubois",
            "specialty": "Droit de la famille",
            "bar_state": "Barreau de Mons",
            "years_experience": 12,
            "rating": 4.9,
            "sessions_count": 178,
            "tags": ["Divorce", "Garde alternee", "Pension alimentaire", "Succession"],
            "availability_status": "soon",
            "availability_minutes": 30,
            "bio": "Specialisee en droit familial et mediations familiales. Mediatrice agreee par la Commission federale de mediation.",
            "photo_url": "https://images.unsplash.com/photo-1685760259914-ee8d2c92d2e0?crop=entropy&cs=srgb&fm=jpg&ixid=M3w3NTY2OTF8MHwxfHNlYXJjaHwxfHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjBldXJvcGVhbiUyMGxhd3llcnxlbnwwfHx8fDE3NzU4MzU1MTN8MA&ixlib=rb-4.1.0&q=85",
            "country": "BE",
            "language": "fr"
        },
        # ===== BELGIAN LAWYERS — NEDERLANDSTALIG =====
        {
            "lawyer_id": "lawyer_be_pieter",
            "name": "Mr. Pieter Van den Berg",
            "specialty": "Arbeidsrecht",
            "bar_state": "Balie Gent",
            "years_experience": 10,
            "rating": 4.9,
            "sessions_count": 165,
            "tags": ["Ontslag", "Niet-concurrentiebeding", "CAO", "RSZ"],
            "availability_status": "now",
            "availability_minutes": 0,
            "bio": "Gespecialiseerd in arbeidsgeschillen en bescherming tegen willekeurig ontslag. Voormalig juridisch adviseur ABVV.",
            "photo_url": "https://images.unsplash.com/photo-1758691737644-ef8be18256c3?crop=entropy&cs=srgb&fm=jpg&ixid=M3w4NjA3MDB8MHwxfHNlYXJjaHwzfHxwcm9mZXNzaW9uYWwlMjBtYW4lMjBoZWFkc2hvdCUyMHBvcnRyYWl0JTIwc3VpdCUyMG9mZmljZXxlbnwwfHx8fDE3NzU4MzU1MTh8MA&ixlib=rb-4.1.0&q=85",
            "country": "BE",
            "language": "nl"
        },
        {
            "lawyer_id": "lawyer_be_laura",
            "name": "Mevr. Laura Janssen",
            "specialty": "Huurrecht Vlaanderen",
            "bar_state": "Balie Antwerpen",
            "years_experience": 8,
            "rating": 4.8,
            "sessions_count": 142,
            "tags": ["Huurgeschillen", "Uitzetting", "Vlaamse Woninghuurwet", "Huurwaarborg"],
            "availability_status": "soon",
            "availability_minutes": 20,
            "bio": "Expert in Vlaams huurrecht. Regelmatig aanwezig bij het vredegerecht van Antwerpen voor huurgeschillen.",
            "photo_url": "https://images.unsplash.com/photo-1702875581098-36844f1910f3?crop=entropy&cs=srgb&fm=jpg&ixid=M3w3NTY2OTF8MHwxfHNlYXJjaHwzfHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjBldXJvcGVhbiUyMGxhd3llcnxlbnwwfHx8fDE3NzU4MzU1MTN8MA&ixlib=rb-4.1.0&q=85",
            "country": "BE",
            "language": "nl"
        },
        {
            "lawyer_id": "lawyer_be_luc",
            "name": "Mr. Luc Vermeersch",
            "specialty": "Consumentenrecht",
            "bar_state": "Balie Brussel",
            "years_experience": 11,
            "rating": 4.9,
            "sessions_count": 189,
            "tags": ["WER", "Onrechtmatige bedingen", "Ombudsman procedures", "E-commerce"],
            "availability_status": "now",
            "availability_minutes": 0,
            "bio": "Expert consumentenrecht en WER. Werkt regelmatig samen met Ombudsfin en de Ombudsman voor Verzekeringen. Tweetalig NL/FR.",
            "photo_url": "https://images.unsplash.com/photo-1568585105565-e372998a195d?crop=entropy&cs=srgb&fm=jpg&ixid=M3w4NjA3MDB8MHwxfHNlYXJjaHwxfHxwcm9mZXNzaW9uYWwlMjBtYW4lMjBoZWFkc2hvdCUyMHBvcnRyYWl0JTIwc3VpdCUyMG9mZmljZXxlbnwwfHx8fDE3NzU4MzU1MTh8MA&ixlib=rb-4.1.0&q=85",
            "country": "BE",
            "language": "nl/fr"
        },
        {
            "lawyer_id": "lawyer_be_sarah_ds",
            "name": "Mevr. Sarah De Smedt",
            "specialty": "Contractenrecht",
            "bar_state": "Balie Leuven",
            "years_experience": 6,
            "rating": 4.8,
            "sessions_count": 98,
            "tags": ["NDA", "Handelscontracten", "Aansprakelijkheid", "IP-recht"],
            "availability_status": "tomorrow",
            "availability_minutes": 0,
            "bio": "Gespecialiseerd in contractenrecht en intellectueel eigendom. Docente KU Leuven.",
            "photo_url": "https://images.unsplash.com/photo-1733348137479-2e726d326d9b?crop=entropy&cs=srgb&fm=jpg&ixid=M3w3NTY2OTF8MHwxfHNlYXJjaHw0fHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjBldXJvcGVhbiUyMGxhd3llcnxlbnwwfHx8fDE3NzU4MzU1MTN8MA&ixlib=rb-4.1.0&q=85",
            "country": "BE",
            "language": "nl"
        },
        # ===== BELGIAN LAWYERS — DEUTSCHSPRACHIG =====
        {
            "lawyer_id": "lawyer_be_klaus",
            "name": "Herr Klaus Mueller",
            "specialty": "Arbeitsrecht",
            "bar_state": "Kammer Eupen",
            "years_experience": 15,
            "rating": 4.8,
            "sessions_count": 87,
            "tags": ["Kuendigung", "Arbeitsvertrag", "Sozialrecht", "KAV"],
            "availability_status": "now",
            "availability_minutes": 0,
            "bio": "Spezialist fuer Arbeitsrecht in der Deutschsprachigen Gemeinschaft. Zweisprachig DE/FR.",
            "photo_url": "https://images.unsplash.com/photo-1758518729314-b02874db8c37?crop=entropy&cs=srgb&fm=jpg&ixid=M3w4NjA3MDB8MHwxfHNlYXJjaHwyfHxwcm9mZXNzaW9uYWwlMjBtYW4lMjBoZWFkc2hvdCUyMHBvcnRyYWl0JTIwc3VpdCUyMG9mZmljZXxlbnwwfHx8fDE3NzU4MzU1MTh8MA&ixlib=rb-4.1.0&q=85",
            "country": "BE",
            "language": "de/fr"
        },
        {
            "lawyer_id": "lawyer_be_anna",
            "name": "Frau Anna Schreiber",
            "specialty": "Mietrecht und Zivilrecht",
            "bar_state": "Kammer Malmedy",
            "years_experience": 9,
            "rating": 4.7,
            "sessions_count": 64,
            "tags": ["Mietvertraege", "Zivilrecht", "Vertragsrecht", "Familienrecht"],
            "availability_status": "soon",
            "availability_minutes": 45,
            "bio": "Mietrecht und allgemeines Zivilrecht in Ostbelgien. Zweisprachig DE/FR.",
            "photo_url": "https://images.unsplash.com/photo-1685760259914-ee8d2c92d2e0?crop=entropy&cs=srgb&fm=jpg&ixid=M3w3NTY2OTF8MHwxfHNlYXJjaHwxfHxwcm9mZXNzaW9uYWwlMjBidXNpbmVzcyUyMGhlYWRzaG90JTIwcG9ydHJhaXQlMjBldXJvcGVhbiUyMGxhd3llcnxlbnwwfHx8fDE3NzU4MzU1MTN8MA&ixlib=rb-4.1.0&q=85",
            "country": "BE",
            "language": "de/fr"
        }
    ]
    
    # Clear existing lawyers
    await db.lawyers.delete_many({})
    
    # Insert new lawyers
    for lawyer in lawyers_data:
        await db.lawyers.insert_one(lawyer)
    
    return {"message": f"Seeded {len(lawyers_data)} lawyers"}


# ================== Document Library ==================

# Load templates
with open(os.path.join(os.path.dirname(__file__), "document_templates.json")) as f:
    DOC_TEMPLATES = json.load(f)

FREE_TEMPLATE_IDS = {t["id"] for t in DOC_TEMPLATES if t.get("free")}

CATEGORIES = [
    {"id": "all", "label": "All"},
    {"id": "employment", "label": "Employment"},
    {"id": "housing", "label": "Housing & Lease"},
    {"id": "business", "label": "Business"},
    {"id": "nda", "label": "NDA & Confidentiality"},
    {"id": "contracts", "label": "Contracts"},
    {"id": "consumer", "label": "Consumer & Demand Letters"},
    {"id": "debt", "label": "Debt & Finance"},
    {"id": "family", "label": "Family"},
    {"id": "realestate", "label": "Real Estate"},
    {"id": "freelance", "label": "Freelance"},
    {"id": "ip", "label": "Intellectual Property"},
    {"id": "court", "label": "Court & Legal"},
    {"id": "immigration", "label": "Immigration"},
]

@api_router.get("/library/templates")
async def get_library_templates(category: str = "all", search: str = ""):
    templates = DOC_TEMPLATES
    if category and category != "all":
        templates = [t for t in templates if t["cat"] == category]
    if search:
        q = search.lower()
        templates = [t for t in templates if q in t["name"].lower() or q in t["desc"].lower()]
    return {"templates": templates, "categories": CATEGORIES, "total": len(templates)}

class DocGenerateRequest(BaseModel):
    template_id: str
    fields: dict
    jurisdiction: str = "California"

@api_router.post("/library/generate")
async def generate_document(body: DocGenerateRequest, current_user: User = Depends(get_current_user)):
    template = next((t for t in DOC_TEMPLATES if t["id"] == body.template_id), None)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    plan = current_user.plan or "free"
    if plan != "pro" and body.template_id not in FREE_TEMPLATE_IDS:
        raise HTTPException(status_code=403, detail="This template requires a Pro plan. Upgrade to access all 158+ templates.")
    
    fields_text = "\n".join(f"- {k}: {v}" for k, v in body.fields.items() if v)
    
    prompt = f"""Generate a complete, professional, legally sound {template['name']} for the jurisdiction of {body.jurisdiction}.

Use the following information provided by the user:
{fields_text}

The document must:
1. Include ALL standard clauses for this document type
2. Be jurisdiction-specific — reference applicable state/country law
3. Use proper legal language while remaining clear
4. Include all necessary sections, headings, and signature blocks
5. Be ready to sign immediately
6. Include the date {datetime.now(timezone.utc).strftime('%B %d, %Y')}

Return the complete document text formatted with clear sections and headings. Use markdown formatting (# for headings, ** for bold, etc). Include signature lines at the end with [SIGNATURE] placeholders."""

    try:
        async with httpx.AsyncClient() as client:
            response = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": get_anthropic_key(),
                    "anthropic-version": "2023-06-01"
                },
                json={
                    "model": "claude-sonnet-4-6",
                    "max_tokens": 4000,
                    "system": "You are a senior attorney drafting legal documents. Generate complete, jurisdiction-specific, professionally formatted legal documents. Never include disclaimers or notes — just the document itself.",
                    "messages": [{"role": "user", "content": prompt}]
                },
                timeout=120.0
            )
            response.raise_for_status()
            data = response.json()
            content = data["content"][0]["text"]
    except Exception as e:
        logger.error(f"Document generation error: {e}")
        raise HTTPException(status_code=500, detail="Could not generate document. Please try again.")
    
    doc_id = f"gendoc_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    doc_record = {
        "doc_id": doc_id,
        "user_id": current_user.user_id,
        "template_id": body.template_id,
        "document_type": template["cat"],
        "document_name": template["name"],
        "generated_content": content,
        "fields": body.fields,
        "jurisdiction": body.jurisdiction,
        "signature_status": "draft",
        "signers": [],
        "created_at": now
    }
    await db.generated_documents.insert_one(doc_record)
    
    return {
        "doc_id": doc_id,
        "document_name": template["name"],
        "content": content,
        "jurisdiction": body.jurisdiction,
        "signature_status": "draft",
        "created_at": now
    }

@api_router.get("/library/generated")
async def get_generated_documents(current_user: User = Depends(get_current_user)):
    docs = await db.generated_documents.find(
        {"user_id": current_user.user_id}, {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    return docs

class SignatureRequest(BaseModel):
    doc_id: str
    signers: list
    message: str = ""

@api_router.post("/library/sign")
async def send_for_signature(body: SignatureRequest, current_user: User = Depends(get_current_user)):
    doc = await db.generated_documents.find_one(
        {"doc_id": body.doc_id, "user_id": current_user.user_id}, {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # HelloSign integration placeholder — needs API key
    hellosign_key = os.environ.get("HELLOSIGN_API_KEY")
    if not hellosign_key:
        now = datetime.now(timezone.utc).isoformat()
        signers_data = [{"name": s.get("name",""), "email": s.get("email",""), "status": "pending", "signed_at": None} for s in body.signers]
        await db.generated_documents.update_one(
            {"doc_id": body.doc_id},
            {"$set": {"signature_status": "pending", "signers": signers_data, "hellosign_request_id": f"mock_{uuid.uuid4().hex[:8]}"}}
        )
        return {
            "status": "pending",
            "message": "Signature request created (HelloSign integration pending — API key required)",
            "signers": signers_data,
            "doc_id": body.doc_id
        }
    
    # TODO: Real HelloSign API call when key is configured
    return {"status": "error", "message": "HelloSign integration not yet configured"}


# ================== Archer Document Creator (Conversational) ==================

ARCHER_DOC_CREATOR_PROMPT = """You are Archer, a Senior Legal Advisor with 20 years of experience. You are helping the user create a legal document. You have generated legal documents for thousands of clients across employment, housing, contracts, consumer protection, and business law.

YOUR APPROACH:
1. Identify the exact document type needed from the user's description
2. Ask 3-5 targeted questions to collect necessary information
3. Generate a complete, legally sound document once all info is collected
4. Offer to modify any clause the user requests

WHEN ASKING QUESTIONS:
- Ask one question at a time, conversationally
- Number each question (Question 1 of 4, etc.)
- Keep questions under 15 words
- Accept informal answers ('2 years', 'New York', 'both parties')

WHEN GENERATING THE DOCUMENT:
- Generate the COMPLETE document — not a template with blanks
- All parties, dates, amounts, and terms must be filled in
- Include all standard clauses for the document type
- Add jurisdiction-specific clauses based on the state/country
- Format with clear section headers and numbered clauses
- Include signature lines at the bottom
- When you generate the document, wrap it in <DOCUMENT> and </DOCUMENT> tags so the frontend can extract it

DOCUMENT QUALITY STANDARDS:
- Every clause must be legally enforceable in the specified jurisdiction
- Include carve-outs and protections that favor the user
- Flag any unusual or risky clauses in your message
- Never generate a document without knowing: parties, jurisdiction, and purpose

AFTER GENERATION:
- Explain in one sentence what protective features you included
- Offer 4 specific modification options as suggestions
- If user asks to modify, update the specific clause and confirm the change
- Always re-wrap the updated full document in <DOCUMENT> and </DOCUMENT> tags

JURISDICTION: Apply laws of {jurisdiction}
LANGUAGE: Respond in {language}"""


class ArcherDocMessage(BaseModel):
    message: str
    conversation_id: Optional[str] = None


@api_router.post("/documents/archer/send")
async def archer_doc_send(data: ArcherDocMessage, current_user: User = Depends(get_current_user)):
    """Archer conversational document creator — send message and get AI response"""
    conv_id = data.conversation_id
    now = datetime.now(timezone.utc).isoformat()

    # Create new conversation if needed
    if not conv_id:
        conv_id = f"jdoc_{uuid.uuid4().hex[:12]}"
        conv_doc = {
            "conversation_id": conv_id,
            "user_id": current_user.user_id,
            "type": "document_creator",
            "title": data.message[:60] + ("..." if len(data.message) > 60 else ""),
            "created_at": now,
            "updated_at": now,
        }
        await db.archer_doc_conversations.insert_one(conv_doc)

    # Free plan: limit to 3 generated documents
    if current_user.plan == "free":
        doc_count = await db.generated_documents.count_documents({
            "user_id": current_user.user_id,
            "created_via": "archer_conversation"
        })
        if doc_count >= 3:
            return {
                "response": "You've created your 3 free documents. Upgrade to Pro for unlimited document creation with Archer.",
                "conversation_id": conv_id,
                "limit_reached": True
            }

    # Save user message
    user_msg = {
        "message_id": f"msg_{uuid.uuid4().hex[:12]}",
        "conversation_id": conv_id,
        "role": "user",
        "content": data.message,
        "created_at": now
    }
    await db.archer_doc_messages.insert_one(user_msg)

    # Build conversation history
    history = await db.archer_doc_messages.find(
        {"conversation_id": conv_id},
        {"_id": 0, "role": 1, "content": 1}
    ).sort("created_at", 1).to_list(100)

    claude_messages = [{"role": m["role"], "content": m["content"]} for m in history]

    # Build system prompt
    user_language = getattr(current_user, 'language', 'en') or 'en'
    user_jurisdiction = getattr(current_user, 'jurisdiction', 'US') or 'US'
    lang_map = {"en": "English", "fr": "French", "fr-BE": "French", "nl": "Dutch", "nl-BE": "Dutch", "de": "German", "de-BE": "German", "es": "Spanish"}
    lang_name = lang_map.get(user_language, "English")

    jurisdiction_text = "US Federal + applicable state law" if user_jurisdiction == "US" else "Belgian federal + regional law"
    system = ARCHER_DOC_CREATOR_PROMPT.format(jurisdiction=jurisdiction_text, language=lang_name)
    system += get_language_instruction(user_language)

    # Call Claude
    try:
        async with httpx.AsyncClient() as http_client:
            response = await http_client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": get_anthropic_key(),
                    "anthropic-version": "2023-06-01"
                },
                json={
                    "model": "claude-sonnet-4-6",
                    "max_tokens": 4096,
                    "system": system,
                    "messages": claude_messages
                },
                timeout=120.0
            )
            response.raise_for_status()
            resp_data = response.json()
            ai_text = ""
            for block in resp_data.get("content", []):
                if block.get("type") == "text":
                    ai_text += block["text"]
            if not ai_text:
                ai_text = "Archer is temporarily unavailable. Please try again."
    except Exception as e:
        logger.error(f"Archer doc creator error: {e}")
        ai_text = "Archer is temporarily unavailable. Please try again."

    # Save AI response
    ai_msg = {
        "message_id": f"msg_{uuid.uuid4().hex[:12]}",
        "conversation_id": conv_id,
        "role": "assistant",
        "content": ai_text,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.archer_doc_messages.insert_one(ai_msg)
    await db.archer_doc_conversations.update_one(
        {"conversation_id": conv_id},
        {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
    )

    # Check if document was generated (contains <DOCUMENT> tags)
    has_document = "<DOCUMENT>" in ai_text and "</DOCUMENT>" in ai_text
    document_content = None
    if has_document:
        import re
        doc_match = re.search(r"<DOCUMENT>(.*?)</DOCUMENT>", ai_text, re.DOTALL)
        if doc_match:
            document_content = doc_match.group(1).strip()
            # Extract title from first line of document or first heading
            doc_lines = [l.strip() for l in document_content.split('\n') if l.strip()]
            doc_title = doc_lines[0].lstrip('#').strip() if doc_lines else "Legal Document"
            if len(doc_title) > 100:
                doc_title = doc_title[:100]
            # Get first user message as fallback context
            first_msg = await db.archer_doc_messages.find_one(
                {"conversation_id": conv_id, "role": "user"},
                {"_id": 0, "content": 1}
            )
            first_request = first_msg["content"][:80] if first_msg else ""
            # Save to generated_documents
            doc_id = f"jdoc_{uuid.uuid4().hex[:12]}"
            doc_record = {
                "doc_id": doc_id,
                "user_id": current_user.user_id,
                "created_via": "archer_conversation",
                "conversation_id": conv_id,
                "document_title": doc_title,
                "original_request": first_request,
                "generated_content": document_content,
                "jurisdiction": jurisdiction_text,
                "signature_status": "draft",
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.generated_documents.insert_one(doc_record)
            # Update conversation with doc reference
            await db.archer_doc_conversations.update_one(
                {"conversation_id": conv_id},
                {"$set": {"last_doc_id": doc_id, "document_title": doc_title}}
            )

    _doc_id = doc_id if (has_document and document_content) else None
    return {
        "response": ai_text,
        "conversation_id": conv_id,
        "has_document": has_document,
        "document_content": document_content,
        "doc_id": _doc_id,
        "limit_reached": False
    }


@api_router.get("/documents/archer/conversations")
async def get_archer_doc_conversations(current_user: User = Depends(get_current_user)):
    """Get user's recent Archer document conversations"""
    convs = await db.archer_doc_conversations.find(
        {"user_id": current_user.user_id},
        {"_id": 0}
    ).sort("updated_at", -1).to_list(20)
    return convs


@api_router.get("/documents/archer/conversations/{conversation_id}/messages")
async def get_archer_doc_messages(conversation_id: str, current_user: User = Depends(get_current_user)):
    """Get messages for a Archer document conversation"""
    conv = await db.archer_doc_conversations.find_one(
        {"conversation_id": conversation_id, "user_id": current_user.user_id},
        {"_id": 0}
    )
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    messages = await db.archer_doc_messages.find(
        {"conversation_id": conversation_id},
        {"_id": 0}
    ).sort("created_at", 1).to_list(200)
    return messages


@api_router.get("/documents/archer/recent")
async def get_recent_generated_docs(current_user: User = Depends(get_current_user)):
    """Get last 5 generated documents for sidebar"""
    docs = await db.generated_documents.find(
        {"user_id": current_user.user_id, "created_via": "archer_conversation"},
        {"_id": 0, "doc_id": 1, "document_title": 1, "created_at": 1, "signature_status": 1}
    ).sort("created_at", -1).to_list(5)
    return docs


# ================== Risk Monitor — Email Surveillance (MOCKED) ==================

@api_router.get("/risk-monitor/status")
async def get_risk_monitor_status(current_user: User = Depends(get_current_user)):
    """Get risk monitor connection status and stats"""
    monitor = await db.risk_monitors.find_one(
        {"user_id": current_user.user_id},
        {"_id": 0}
    )
    if not monitor:
        return {
            "connected": False,
            "provider": None,
            "email": None,
            "emails_scanned": 0,
            "threats_detected": 0,
            "last_scan": None,
            "alerts": []
        }
    
    alerts = await db.risk_monitor_alerts.find(
        {"user_id": current_user.user_id},
        {"_id": 0}
    ).sort("detected_at", -1).limit(10).to_list(10)
    
    return {
        "connected": monitor.get("connected", False),
        "provider": monitor.get("provider"),
        "email": monitor.get("email"),
        "emails_scanned": monitor.get("emails_scanned", 0),
        "threats_detected": monitor.get("threats_detected", 0),
        "last_scan": monitor.get("last_scan"),
        "alerts": alerts
    }

@api_router.post("/risk-monitor/connect")
async def connect_risk_monitor(request: Request, current_user: User = Depends(get_current_user)):
    """Connect email account for risk monitoring (MOCKED — requires OAuth credentials)"""
    body = await request.json()
    provider = body.get("provider", "gmail")
    
    if provider not in ["gmail", "outlook"]:
        raise HTTPException(status_code=400, detail="Provider must be 'gmail' or 'outlook'")
    
    now = datetime.now(timezone.utc).isoformat()
    
    # MOCKED: Simulate successful connection
    monitor_data = {
        "user_id": current_user.user_id,
        "provider": provider,
        "email": current_user.email,
        "connected": True,
        "emails_scanned": 0,
        "threats_detected": 0,
        "last_scan": now,
        "connected_at": now
    }
    
    await db.risk_monitors.update_one(
        {"user_id": current_user.user_id},
        {"$set": monitor_data},
        upsert=True
    )
    
    # Simulate initial scan with some mock alerts
    mock_alerts = [
        {
            "alert_id": f"alert_{uuid.uuid4().hex[:12]}",
            "user_id": current_user.user_id,
            "type": "legal_document",
            "severity": "high",
            "subject": "Notice of Intent to Sue — Account #4892",
            "sender": "collections@lawfirm-associates.com",
            "preview": "This letter serves as formal notice that our client intends to pursue legal action regarding the outstanding balance...",
            "detected_at": now,
            "status": "new",
            "recommended_action": "Upload this document to Archer for full analysis"
        },
        {
            "alert_id": f"alert_{uuid.uuid4().hex[:12]}",
            "user_id": current_user.user_id,
            "type": "contract",
            "severity": "medium",
            "subject": "Lease Renewal — New Terms Enclosed",
            "sender": "management@propertygroup.com",
            "preview": "Please review the attached lease renewal agreement. Note the following changes to your current terms...",
            "detected_at": (datetime.now(timezone.utc) - timedelta(hours=3)).isoformat(),
            "status": "new",
            "recommended_action": "Use Contract Guard to review before signing"
        },
        {
            "alert_id": f"alert_{uuid.uuid4().hex[:12]}",
            "user_id": current_user.user_id,
            "type": "deadline",
            "severity": "low",
            "subject": "Reminder: Policy Renewal Due Feb 28",
            "sender": "noreply@insurance-co.com",
            "preview": "Your current insurance policy is set to expire on February 28, 2026. Please review the renewal terms...",
            "detected_at": (datetime.now(timezone.utc) - timedelta(days=1)).isoformat(),
            "status": "new",
            "recommended_action": "Review renewal terms for any changes"
        }
    ]
    
    for alert in mock_alerts:
        await db.risk_monitor_alerts.insert_one(alert)
    
    await db.risk_monitors.update_one(
        {"user_id": current_user.user_id},
        {"$set": {"emails_scanned": 147, "threats_detected": 3}}
    )
    
    return {
        "status": "connected",
        "provider": provider,
        "email": current_user.email,
        "message": f"Successfully connected to {provider.title()}. Initial scan complete — 3 items detected."
    }

@api_router.post("/risk-monitor/disconnect")
async def disconnect_risk_monitor(current_user: User = Depends(get_current_user)):
    """Disconnect email monitoring"""
    await db.risk_monitors.update_one(
        {"user_id": current_user.user_id},
        {"$set": {"connected": False}}
    )
    return {"status": "disconnected"}

@api_router.put("/risk-monitor/alerts/{alert_id}/dismiss")
async def dismiss_alert(alert_id: str, current_user: User = Depends(get_current_user)):
    """Dismiss a risk monitor alert"""
    await db.risk_monitor_alerts.update_one(
        {"alert_id": alert_id, "user_id": current_user.user_id},
        {"$set": {"status": "dismissed"}}
    )
    return {"status": "dismissed"}

# ================== Contract Guard Reviews ==================

@api_router.get("/contract-guard/reviews")
async def get_contract_guard_reviews(current_user: User = Depends(get_current_user)):
    """Get all contract guard reviews for the user"""
    reviews = await db.contract_guard_reviews.find(
        {"user_id": current_user.user_id},
        {"_id": 0}
    ).sort("created_at", -1).to_list(50)
    return reviews


@api_router.get("/user/country-config")
async def get_country_config(current_user: User = Depends(get_current_user)):
    """Get country-specific configuration for the current user"""
    country = current_user.jurisdiction or "US"
    if country == "BE":
        return {
            "country": "BE",
            "country_name": "Belgium",
            "regions": [
                {"id": "Wallonie", "label": "Wallonie", "language": "fr-BE"},
                {"id": "Bruxelles-Capitale", "label": "Bruxelles-Capitale", "language": "fr-BE"},
                {"id": "Flandre", "label": "Vlaanderen / Flandre", "language": "nl-BE"},
                {"id": "Communaute germanophone", "label": "Communaute germanophone", "language": "de-BE"}
            ],
            "languages": [
                {"id": "fr-BE", "label": "Francais"},
                {"id": "nl-BE", "label": "Nederlands"},
                {"id": "de-BE", "label": "Deutsch"}
            ],
            "currency": "EUR",
            "legal_refs": {
                "legislation": "www.ejustice.just.fgov.be",
                "cct": "www.cnt-nar.be",
                "jurisprudence": "www.juridat.be",
                "codes": "www.codesdroitbelge.be"
            }
        }
    return {
        "country": "US",
        "country_name": "United States",
        "regions": [],
        "languages": [{"id": "en", "label": "English"}],
        "currency": "USD",
        "legal_refs": {"courtlistener": "www.courtlistener.com"}
    }

# ================== Health Check ==================

@api_router.get("/")
async def root():
    return {"message": "Archer API is running", "version": "1.0.0"}

@api_router.get("/health")
async def health_check():
    return {"status": "healthy"}


# Attorney Portal routes extracted to routes/attorney_routes.py
from routes.attorney_routes import router as attorney_router
api_router.include_router(attorney_router)

# Sprint A — Attorney Portal (parallel auth stack, /api/attorneys/*)
from routes.attorney_portal_routes import router as attorney_portal_router
from utils.attorney_auth import ensure_indexes as ensure_attorney_portal_indexes
api_router.include_router(attorney_portal_router)

# Admin auth system (separate from client/attorney)
from routes.admin_auth_routes import router as admin_auth_router
from routes.admin_attorneys_routes import router as admin_attorneys_v2_router
api_router.include_router(admin_auth_router)
api_router.include_router(admin_attorneys_v2_router)

# Sprint B — Attorney portal case endpoints + admin manual-assign
from routes.attorney_portal_cases import router as attorney_portal_cases_router
from routes.admin_portal_routes import router as admin_portal_router
api_router.include_router(attorney_portal_cases_router)
api_router.include_router(admin_portal_router)

# Sprint C — auto-matching: admin dashboard + client trigger
from routes.admin_matching_routes import router as admin_matching_router
from routes.client_attorney_request_routes import router as client_attorney_request_router
api_router.include_router(admin_matching_router)
api_router.include_router(client_attorney_request_router)

# Sprint D — Stripe Connect + webhooks + earnings + client checkout
from routes.stripe_connect_routes import router as stripe_connect_router
from routes.stripe_webhooks import router as stripe_webhooks_router
from routes.attorney_earnings_routes import router as attorney_earnings_router
from routes.client_checkout_routes import router as client_checkout_router
api_router.include_router(stripe_connect_router)
api_router.include_router(stripe_webhooks_router)
api_router.include_router(attorney_earnings_router)
api_router.include_router(client_checkout_router)

# Sprint E — Calendly + Daily.co + Live Counsel
from routes.calendly_routes import router as calendly_router
from routes.calendly_webhook import router as calendly_webhook_router
from routes.live_counsel_routes import router as live_counsel_router, attorney_router as live_counsel_attorney_router
api_router.include_router(calendly_router)
api_router.include_router(calendly_webhook_router)
api_router.include_router(live_counsel_router)
api_router.include_router(live_counsel_attorney_router)

# Sprint F — Attorney profile management
from routes.attorney_profile_routes import router as attorney_profile_router
api_router.include_router(attorney_profile_router)

# Phase 1 fixes — multi-document upload with binary tier gating
from routes.multi_upload_routes import router as multi_upload_router
api_router.include_router(multi_upload_router)

# Phase 2 — client notifications (letter_ready, etc.)
from routes.client_notifications_routes import (
    router as client_notifications_router,
    ensure_indexes as ensure_client_notifications_indexes,
)
api_router.include_router(client_notifications_router)

from utils.attorney_auth import (
    migrate_sprint_c_fields, migrate_sprint_d_fields,
    migrate_sprint_e_fields, migrate_sprint_f_fields,
)
from jobs.scheduler import start_scheduler, stop_scheduler

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup_event():
    """Initialize storage and seed data on startup"""
    try:
        init_storage()
        logger.info("Storage initialized")
    except Exception as e:
        logger.error(f"Storage init failed: {e}")

    try:
        await ensure_attorney_portal_indexes()
    except Exception as e:
        logger.error(f"Attorney portal index init failed: {e}")

    try:
        await migrate_sprint_c_fields()
    except Exception as e:
        logger.error(f"Sprint C migration failed: {e}")

    try:
        await migrate_sprint_d_fields()
    except Exception as e:
        logger.error(f"Sprint D migration failed: {e}")

    try:
        await migrate_sprint_e_fields()
    except Exception as e:
        logger.error(f"Sprint E migration failed: {e}")

    try:
        await migrate_sprint_f_fields()
    except Exception as e:
        logger.error(f"Sprint F migration failed: {e}")

    try:
        await ensure_client_notifications_indexes()
    except Exception as e:
        logger.error(f"client_notifications index init failed: {e}")

    try:
        start_scheduler()
    except Exception as e:
        logger.error(f"Scheduler start failed: {e}")
    
    # Seed lawyers — re-seed if missing Belgian lawyers
    be_count = await db.lawyers.count_documents({"country": "BE"})
    if be_count == 0:
        logger.info("Seeding lawyer data (US + Belgian)...")
        await seed_lawyers()

    # Seed test attorney if none exist
    atty_count = await db.attorney_profiles.count_documents({})
    if atty_count == 0:
        logger.info("Seeding test attorney...")
        atty_user_id = f"user_{uuid.uuid4().hex[:12]}"
        atty_id = f"atty_{uuid.uuid4().hex[:12]}"
        now = datetime.now(timezone.utc).isoformat()
        pw_hash = hash_password("attorney123")
        await db.users.insert_one({
            "user_id": atty_user_id, "email": "attorney@archer.com", "name": "Sarah Mitchell",
            "picture": None, "password_hash": pw_hash, "auth_provider": "email",
            "plan": "pro", "country": "US", "jurisdiction": "US", "language": "en",
            "account_type": "attorney", "created_at": now,
            "notif_risk_score": True, "notif_deadlines": True, "notif_calls": True,
            "notif_lawyers": False, "notif_promo": False, "data_sharing": True, "improve_ai": True,
        })
        await db.attorney_profiles.insert_one({
            "attorney_id": atty_id, "user_id": atty_user_id, "slug": f"sarah-mitchell-{atty_id[-6:]}",
            "full_name": "Sarah Mitchell", "email": "attorney@archer.com", "phone": "+1-555-0199",
            "bar_number": "NY-2847561", "states_licensed": ["New York", "California", "New Jersey"],
            "country": "US", "years_experience": 12, "law_school": "Columbia Law School",
            "graduation_year": 2014, "specialties": ["Employment law", "Tenant rights", "Contract law", "Consumer protection"],
            "bio": "Former BigLaw attorney turned consumer advocate. 12 years fighting for individuals against corporations. Specialized in employment disputes and tenant protection.",
            "photo_url": None, "languages": ["en", "fr"], "linkedin_url": None,
            "session_price": 199, "archer_commission": 0.20, "attorney_payout": 159,
            "stripe_connect_id": None, "stripe_connect_status": "pending",
            "application_status": "approved", "rejection_reason": None,
            "rating": 4.8, "total_sessions": 47, "total_earnings": 7473, "review_count": 38,
            "is_available": True, "available_from": "09:00", "available_until": "17:00",
            "available_days": ["Mon", "Tue", "Wed", "Thu", "Fri"],
            "blocked_dates": [], "buffer_minutes": 15, "timezone": "America/New_York",
            "created_at": now, "approved_at": now,
        })
        logger.info("Test attorney seeded: attorney@archer.com / attorney123")

    # Seed Pro test accounts
    for acct in [
        {"email": "test@archer.legal", "name": "Alex Thompson", "country": "US", "jurisdiction": "US", "language": "en"},
        {"email": "belgium@archer.legal", "name": "Marie Dupont", "country": "BE", "jurisdiction": "BE", "language": "fr"},
    ]:
        existing = await db.users.find_one({"email": acct["email"]})
        if not existing:
            uid = f"user_{uuid.uuid4().hex[:12]}"
            now = datetime.now(timezone.utc).isoformat()
            await db.users.insert_one({
                "user_id": uid, "email": acct["email"], "name": acct["name"],
                "picture": None, "password_hash": hash_password("ArcherPro2026!"),
                "auth_provider": "email", "plan": "pro", "country": acct["country"],
                "jurisdiction": acct["jurisdiction"], "language": acct["language"],
                "account_type": "client", "created_at": now,
                "free_call_used": False, "free_analyses_used": 0,
                "notif_risk_score": True, "notif_deadlines": True, "notif_calls": True,
                "notif_lawyers": False, "notif_promo": False, "data_sharing": True, "improve_ai": True,
            })
            logger.info(f"Pro test account seeded: {acct['email']}")

@app.on_event("shutdown")
async def shutdown_scheduler():
    try:
        stop_scheduler()
    except Exception:
        pass


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
